"""Generate Changes in Review State Analysis report for 2026-03-03."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def make_chg_url(chg):
    return f"https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number={chg}"


def styled_header_row(table, headers):
    for i, t in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(t)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")


def add_data_row(table, chg, rest, shade=None):
    row = table.add_row()
    cell0 = row.cells[0]
    cell0.text = ""
    p = cell0.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_hyperlink(p, chg, make_chg_url(chg))
    if shade:
        set_cell_shading(cell0, shade)
    for idx, val in enumerate(rest):
        cell = row.cells[idx + 1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        if shade:
            set_cell_shading(cell, shade)
    return row


def add_simple_row(table, cells, shade=None, bold_first=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
        if shade:
            set_cell_shading(cell, shade)
    return row


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    title = doc.add_heading("Changes in Review State \u2014 Analysis & Audit", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Report Date: March 3, 2026  |  Scope: Normal & Emergency Changes  |  State: Review")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared by: Dan Fallon, Change Manager")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This report audits 9 change requests currently in Review state in ServiceNow "
        "as of March 3, 2026 (8 Normal, 1 Emergency). Per ISMS-WI-11.01-06 \u00a74.1.2.4\u20135, "
        "changes in Review state should have post-implementation validation documented "
        "in work notes before the Change Manager reviews and closes the ticket."
    )

    p = doc.add_paragraph()
    run = p.add_run("CRITICAL FINDING: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.add_run(
        "All 9 changes have planned implementation dates that have already passed "
        "(02/12\u201302/27). These changes are either implemented but not closed, or were "
        "never executed and require rescheduling. This backlog represents a process gap "
        "in change lifecycle management."
    )

    doc.add_paragraph()

    # ── Summary Dashboard ─────────────────────────────────────────────────
    doc.add_heading("Summary Dashboard", level=1)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    styled_header_row(tbl, ["Metric", "Count"])

    dashboard = [
        ("Total in Review", "9"),
        ("Normal Changes", "8"),
        ("Emergency Changes", "1"),
        ("All Planned Dates Past", "9 (100%)"),
        ("Retroactive Emergency", "1"),
        ("No Rollback Plan", "1"),
        ("SAR Question Missing/Blank", "2"),
        ("Risk Assessment Incomplete", "1"),
        ("Tickets >90 Days Old", "1"),
        ("High Severity Findings", "3"),
        ("Medium Severity Findings", "3"),
    ]
    for label, val in dashboard:
        row = tbl.add_row()
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        run = p.add_run(label)
        run.font.size = Pt(9)
        row.cells[1].text = ""
        p = row.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(val)
        run.font.size = Pt(9)
        run.bold = True

    doc.add_paragraph()

    # ── Audit Findings Summary ────────────────────────────────────────────
    doc.add_heading("Audit Findings Summary", level=1)

    findings_tbl = doc.add_table(rows=1, cols=4)
    findings_tbl.style = "Table Grid"
    findings_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    styled_header_row(findings_tbl, ["#", "Finding", "Severity", "Changes Affected"])

    findings = [
        ("1", "All planned dates in the past \u2014 changes stuck in Review", "High", "All 9"),
        ("2", "Retroactive emergency change \u2014 created after implementation", "High", "CHG0039327"),
        ("3", "No rollback plan on destructive data purge with PHI", "High", "CHG0039351"),
        ("4", "SAR unanswered despite PHI/PII involvement", "Medium", "CHG0039377, CHG0039351"),
        ("5", "Risk assessment incomplete \u2014 missing answers", "Medium", "CHG0039327"),
        ("6", "Stale change \u2014 6 months old in pipeline", "Medium", "CHG0038691"),
        ("7", "Thin backout plan for 409-record data modification", "Low", "CHG0039359"),
        ("8", "PHI/PII flag questionable on account disablement", "Low", "CHG0039356"),
    ]

    sev_colors = {"High": "FFC7CE", "Medium": "FFEB9C", "Low": "C6EFCE"}
    for num, finding, sev, affected in findings:
        row = add_simple_row(findings_tbl, [num, finding, sev, affected])
        set_cell_shading(row.cells[2], sev_colors[sev])

    doc.add_paragraph()

    # ── Aging Analysis ────────────────────────────────────────────────────
    doc.add_heading("Aging Analysis \u2014 Days Past Planned End", level=1)
    doc.add_paragraph(
        "All 9 changes have planned dates before today (03/03/2026). The table below "
        "shows each change sorted by how far past its planned end date."
    )

    aging_tbl = doc.add_table(rows=1, cols=6)
    aging_tbl.style = "Table Grid"
    aging_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    styled_header_row(aging_tbl, ["CHG", "Short Description", "Type", "Planned End", "Days Past", "Severity"])

    aging_data = [
        ("CHG0039327", "New AP: MOD-AP03", "Emergency", "02/12/2026", "19 days", "CRITICAL", "FFC7CE"),
        ("CHG0039379", "Update MS Defender scan freq (Mac)", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0039377", "Disable FTP accounts - Jan 2026 review", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0039375", "Geoblock Cuba and Syria in Okta", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0039359", "409 Provider Errors (Cred Event)", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0039356", "Service Account disable - Jan 2026", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0039351", "2026 Annual Purge SFDC", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0038691", "Preboarding Checklist - Specialist field", "Normal", "02/26/2026", "5 days", "Moderate", "FFEB9C"),
        ("CHG0039376", "Enable Block malicious IPs in ZIA", "Normal", "02/27/2026", "4 days", "Moderate", "FFEB9C"),
    ]

    for chg, desc, typ, pend, days, sev, shade in aging_data:
        add_data_row(aging_tbl, chg, [desc, typ, pend, days, sev], shade)

    doc.add_paragraph()

    # ── Change-by-Change Analysis ─────────────────────────────────────────
    doc.add_heading("Change-by-Change Analysis", level=1)

    changes = [
        {
            "number": "CHG0039379",
            "desc": "Update scanning frequency for MS Defender (Mac)",
            "type": "Normal", "risk": "Low",
            "group": "Service Delivery Optimization", "assignee": "Jason Nguyen",
            "planned": "02/26/2026 16:00\u201316:30 PST",
            "impl": "OK \u2014 discrete steps (modify plist, deploy via Jamf, verify logs)",
            "backout": "OK \u2014 revert to previous profile via Jamf",
            "test": "OK \u2014 tested in IT test group",
            "risk_assess": "Complete",
            "findings": "Schedule overdue. Otherwise well-documented.",
            "severity": "Low",
        },
        {
            "number": "CHG0039377",
            "desc": "Disable FTP accounts - Access review Jan 2026",
            "type": "Normal", "risk": "Low",
            "group": "Enterprise Systems", "assignee": "Philip Weiss",
            "planned": "02/26/2026 14:00\u201314:20 PST",
            "impl": "OK \u2014 7 accounts listed with manager approvals",
            "backout": "OK \u2014 re-enable accounts",
            "test": "OK \u2014 verify account cannot access FTP",
            "risk_assess": "SAR field blank (unanswered)",
            "findings": "Schedule overdue. SAR question left blank \u2014 needs response.",
            "severity": "Medium",
        },
        {
            "number": "CHG0039376",
            "desc": "Enable Block malicious IPs and domains in ZIA",
            "type": "Normal", "risk": "Low",
            "group": "Infosec", "assignee": "Jeffrey How",
            "planned": "02/27/2026 09:00\u201310:00 PST",
            "impl": "OK \u2014 4 discrete steps in ZIA Admin",
            "backout": "OK \u2014 reverse steps to disable",
            "test": "OK \u2014 check ZIA logs post-implementation",
            "risk_assess": "Complete",
            "findings": "Schedule overdue. No subset testing possible; vendor recommendation mitigates risk.",
            "severity": "Low",
        },
        {
            "number": "CHG0039375",
            "desc": "Geoblock Cuba and Syria in Okta",
            "type": "Normal", "risk": "Low",
            "group": "Infosec", "assignee": "Jeffrey How",
            "planned": "02/26/2026 15:00\u201316:00 PST",
            "impl": "OK \u2014 3 discrete steps including Confluence update",
            "backout": "OK \u2014 reverse steps",
            "test": "OK \u2014 90-day log analysis shows zero logons from target countries",
            "risk_assess": "Complete",
            "findings": "Schedule overdue. Well-justified with 90-day log evidence. Clean change.",
            "severity": "Low",
        },
        {
            "number": "CHG0039359",
            "desc": "409 Provider Errors for \"Cred Event In Process\"",
            "type": "Normal", "risk": "Low",
            "group": "Enterprise Applications", "assignee": "Parvathi Arun",
            "planned": "02/26/2026 09:00\u201311:00 PST",
            "impl": "OK \u2014 very detailed pre/post steps for 409 records",
            "backout": "Thin \u2014 \"data import task will be stopped\"",
            "test": "OK \u2014 tested in lower env, 24-48hr monitoring planned",
            "risk_assess": "Complete but vendor support unavailable during change",
            "findings": "Schedule overdue. Backout plan is thin for a data modification affecting 409 records. Should confirm backup/export was taken before toggling statuses.",
            "severity": "Medium",
        },
        {
            "number": "CHG0039356",
            "desc": "Service Account disable - Access review Jan 2026",
            "type": "Normal", "risk": "Low",
            "group": "Enterprise Systems", "assignee": "Philip Weiss",
            "planned": "02/26/2026 14:00\u201314:15 PST",
            "impl": "OK \u2014 3 accounts listed with revocation evidence",
            "backout": "OK \u2014 re-enable accounts",
            "test": "OK \u2014 verify disabled status",
            "risk_assess": "Complete \u2014 notes PHI/PII impact and business-critical",
            "findings": "Schedule overdue. PHI/PII flagged as Yes for service account disablement \u2014 verify accuracy.",
            "severity": "Low",
        },
        {
            "number": "CHG0039351",
            "desc": "2026 Annual Purge SFDC",
            "type": "Normal", "risk": "Low",
            "group": "Enterprise Applications", "assignee": "Yanyan Meng",
            "planned": "02/26/2026 19:00\u201321:00 PST",
            "impl": "OK \u2014 436 contacts, 38 object types listed",
            "backout": "NONE \u2014 permanent data deletion per retention policy",
            "test": "OK \u2014 tested in Sandbox2, Racquel verified",
            "risk_assess": "SAR not completed despite PHI/PII = Yes; Business-critical = Yes",
            "findings": "Schedule overdue. No rollback by design \u2014 permanent destructive purge of PHI data. SAR not completed despite PHI/PII = Yes. Verify Sandbox2 sign-off and BO approval is attached before closing.",
            "severity": "High",
        },
        {
            "number": "CHG0038691",
            "desc": "Preboarding Checklist Parent/Child Logic - Add Specialist field",
            "type": "Normal", "risk": "Low",
            "group": "Enterprise Applications", "assignee": "Parvathi Arun",
            "planned": "02/26/2026 09:00\u201311:00 PST",
            "impl": "OK \u2014 detailed pre-deployment and production deployment steps",
            "backout": "OK \u2014 detailed rollback triggers and monitoring steps",
            "test": "OK \u2014 tested in sb2 with documentation",
            "risk_assess": "Complete",
            "findings": "Schedule overdue. Change is ~6 months old (created 09/11/2025). Best-documented change in batch. Determine if deployed or stalled.",
            "severity": "Medium",
        },
        {
            "number": "CHG0039327",
            "desc": "New AP: MOD-AP03 (EMERGENCY)",
            "type": "Emergency", "risk": "Low",
            "group": "Enterprise Networking", "assignee": "Daniel Anderson",
            "planned": "02/12/2026 16:00\u201319:51 PST",
            "impl": "OK \u2014 detailed Meraki/AP configuration steps",
            "backout": "OK \u2014 tear down WAP, delete from dashboard",
            "test": "OK \u2014 connectivity verified",
            "risk_assess": "Incomplete \u2014 business-critical question unanswered",
            "findings": "RETROACTIVE EMERGENCY. Change created at 19:54 PST \u2014 3 minutes AFTER the planned end time of 19:51 PST. Implemented before the CR was filed. Now 19 days old in Review. Risk assessment incomplete. Needs immediate closure with post-implementation review.",
            "severity": "High",
        },
    ]

    for chg in changes:
        # Heading with hyperlink
        h = doc.add_heading(level=2)
        add_hyperlink(h, chg["number"], make_chg_url(chg["number"]))
        h.add_run(f" \u2014 {chg['desc']}")

        # Detail table
        info_tbl = doc.add_table(rows=0, cols=2)
        info_tbl.style = "Table Grid"

        fields = [
            ("Type", chg["type"]),
            ("Risk", chg["risk"]),
            ("Assignment Group", chg["group"]),
            ("Assigned To", chg["assignee"]),
            ("Planned Window", chg["planned"]),
            ("Implementation Plan", chg["impl"]),
            ("Backout Plan", chg["backout"]),
            ("Test Plan", chg["test"]),
            ("Risk Assessment", chg["risk_assess"]),
        ]

        for label, value in fields:
            row = info_tbl.add_row()
            row.cells[0].text = ""
            p0 = row.cells[0].paragraphs[0]
            run = p0.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            row.cells[1].text = ""
            p1 = row.cells[1].paragraphs[0]
            run = p1.add_run(value)
            run.font.size = Pt(9)

            # Color-code issues
            val_lower = value.lower()
            if "none" in val_lower and "backout" in label.lower():
                set_cell_shading(row.cells[1], "FFC7CE")
            elif "thin" in val_lower:
                set_cell_shading(row.cells[1], "FFEB9C")
            elif "blank" in val_lower or "incomplete" in val_lower or "not completed" in val_lower:
                set_cell_shading(row.cells[1], "FFEB9C")

        # Findings
        p = doc.add_paragraph()
        run = p.add_run("Findings: ")
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(chg["findings"])
        run.font.size = Pt(10)
        if chg["severity"] == "High":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_paragraph()

    # ── Recommended Actions ───────────────────────────────────────────────
    doc.add_heading("Recommended Actions", level=1)

    actions = [
        (
            "1. Move All 9 Changes Out of Review",
            "All 9 changes have planned dates that have passed. Each should be either "
            "closed (if implemented successfully) or rescheduled with new planned dates. "
            "Implementers must add validation work notes confirming outcome before "
            "the Change Manager can close with appropriate closure code.",
        ),
        (
            "2. Post-Implementation Review for CHG0039327",
            "The retroactive emergency change (New AP: MOD-AP03) requires a post-mortem "
            "review. Daniel Anderson should document the emergency approval chain and "
            "circumstances. Complete the unanswered risk assessment question. Close after review.",
        ),
        (
            "3. SAR Compliance for CHG0039351 (SFDC Purge)",
            "This is a permanent, destructive data purge affecting PHI/PII with no rollback. "
            "Confirm whether a Security Architecture Review is required given PHI/PII = Yes. "
            "Verify Sandbox2 sign-off and Business Owner approval documentation is attached "
            "to the change record before closing.",
        ),
        (
            "4. Complete SAR Field on CHG0039377",
            "The FTP account disablement change (Philip Weiss) has the SAR question blank "
            "in the Risk Assessment. This should be answered to complete the record.",
        ),
        (
            "5. Investigate Stale Change CHG0038691",
            "This change was created on 09/11/2025 (~6 months ago) for a Salesforce Flow "
            "deployment. Determine whether it was actually deployed on 02/26/2026 as planned, "
            "or if it has been stalled. Close or cancel accordingly.",
        ),
    ]

    for title_text, body_text in actions:
        p = doc.add_paragraph()
        run = p.add_run(title_text)
        run.bold = True
        run.font.size = Pt(10)
        doc.add_paragraph(body_text, style="List Bullet")

    doc.add_paragraph()

    # ── Policy References ─────────────────────────────────────────────────
    doc.add_heading("Policy References", level=1)
    for ref in [
        "ISMS-STA-11.01-01: Enterprise Change Management Program",
        "ISMS-PROC-11.01-05: Procedure for Managing Change Request",
        "ISMS-WI-11.01-06: Work Instruction for Submitting Change Request",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Generated 2026-03-03 by EA  |  Data source: ServiceNow Production")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    output_path = (
        r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management"
        r"\CCB\2026\Changes in Review State Analysis 2026-03-03.docx"
    )
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
