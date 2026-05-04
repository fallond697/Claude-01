"""Build Concur cross-system inventory Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

SN = "https://vituity.service-now.com"


def chg_url(num): return f"{SN}/change_request.do?sysparm_query=number={num}"
def inc_url(num): return f"{SN}/incident.do?sysparm_query=number={num}"
def prb_url(num): return f"{SN}/problem.do?sysparm_query=number={num}"
def kb_url(num): return f"{SN}/kb_view.do?sysparm_article={num}"
def jira_url(key): return f"https://medamerica.atlassian.net/browse/{key}"


def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold:
        rPr.append(OxmlElement("w:b"))
    new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve")
    new_run.append(t); h.append(new_run); paragraph._p.append(h)


def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_table_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4"); b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_header_row(table, headers):
    cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cells[i].text = ""
        run = cells[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cells[i], "1F4E79")


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


doc = Document()
style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Concur Application — Cross-System Inventory")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"Sources: ServiceNow • Jira • SharePoint   |   Generated {date.today().isoformat()}   |   Owner: Dan Fallon").italic = True

# Executive Summary
add_h(doc, "Executive Summary", 1)
for label, body in [
    ("Owner", "Enterprise Applications group; admins Ramanathan Sockalingam, Parvathi Arun, Edmund Trinidad."),
    ("Function", "Vituity travel & expense (T&E) platform; SAP Concur Solutions."),
    ("Top Integrations", "HCM (multiple file feeds), NetSuite, Coupa (project codes), Snowflake (data warehouse), Workato/Moveworks/Otto, Navan (TripActions) for travel."),
    ("Active 2026 Work", "Concur ingestion pipeline for AN Invoicing (EDA-6877 done), AP Connections Security Architecture Review (PROD-12118), Phase 2 API placeholder (EDA-7003)."),
    ("Major Cancellations", "Heavy cancellation cluster around HCM→Concur file integrations (employee, approver, multi-assignment) — see Hotspots."),
    ("Documentation Home", "SharePoint sp283 (Concur hub) + sp480 + People Operations KB (16 Concur articles)."),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{label}: ").bold = True
    p.add_run(body)

# Hotspots
add_h(doc, "Cross-System Hotspots", 1)
for title_text, body, link, url in [
    ("HCM → Concur integration cluster collapsed", "8 changes in the HCM→Concur employee/approver/multi-assignment file series were canceled (CHG0037661, 37662, 37666, 37667, 37668, 38108, 39419, 31988, 31877). Multiple Jira stories Cancelled. Worth understanding why this initiative stalled.", "CHG0037668", chg_url("CHG0037668")),
    ("AP Connections Security Architecture Review open", "PROD-12118 — Security Architecture Review for AP Connections to Concur, owned by Gopalakrishna Budharam. Created 2026-04-20.", "PROD-12118", jira_url("PROD-12118")),
    ("Concur Phase 2 API ingestion placeholder", "EDA-7003 in Backlog — placeholder for Phase 2 of AN Invoicing API connection to Concur.", "EDA-7003", jira_url("EDA-7003")),
    ("Zenoti ↔ Concur spike never executed", "PROD-4704 (To Do since May 2025). Compound Pharmacy use case epic PROD-4703 also unassigned.", "PROD-4704", jira_url("PROD-4704")),
    ("PROD-4057 Blocked", "Missing-on-employees HCM file (CHG0037662) — story is Blocked since Jun 2025.", "PROD-4057", jira_url("PROD-4057")),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title_text} — ").bold = True
    p.add_run(body + " (")
    add_hyperlink(p, url, link); p.add_run(")")

# CMDB
add_h(doc, "ServiceNow CMDB", 1)
doc.add_paragraph("Three Concur-related CIs (other matches were rubygem-concurrent-ruby — unrelated software packages, excluded).")
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Name", "Class", "Status", "Support Group"])
for row in [
    ("Concur", "Configuration Item", "Operational", "—"),
    ("Concur", "Business Application", "Operational", "Enterprise Applications"),
    ("Concur", "Business Service", "Operational", "—"),
]:
    cells = t.add_row().cells
    for i, v in enumerate(row): cells[i].text = v

# Changes
add_h(doc, "ServiceNow — Change Requests (26 returned)", 1)

add_h(doc, "Notable Closed Changes", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Number", "Short Description", "Type", "Assignee"])
for num, desc, typ, who in [
    ("CHG0038629", "Revert prior Concur changes", "Normal", "Parvathi Arun"),
    ("CHG0038493", "Convert Email to VERA Ticket & Assign to Concur", "Normal", "Edmund Trinidad"),
    ("CHG0038031", "Concur profile missing", "Normal", "Parvathi Arun"),
    ("CHG0037901", "Requesting Application Configuration", "Standard", "Ray Blor"),
    ("CHG0037781", "Modification of Email address delivery", "Normal", "Ramanathan Sockalingam"),
    ("CHG0036790", "Decommission/remove Extracts from HCM", "Normal", "Ramanathan Sockalingam"),
    ("CHG0035232", "Project code field Coupa↔Concur↔NetSuite", "Normal", "Jitin Xavier"),
    ("CHG0035190", "Snowflake Warehouse Changes (optimization)", "Normal", "Avinash VedavyasPrabhu"),
    ("CHG0034775", "Add Application System Administrator", "Standard", "Edmund Trinidad"),
    ("CHG0033084", "HCM to Concur Interface Leap 21", "Normal", "Yanyan Meng"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], chg_url(num), num)
    cells[1].text = desc; cells[2].text = typ; cells[3].text = who

add_h(doc, "Cancelled Changes (HCM→Concur cluster + others)", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Number", "Short Description", "Type", "Assignee"])
for num, desc, typ, who in [
    ("CHG0039419", "Concur Integration — Truncate HCM Dept Names", "Normal", "Ramanathan Sockalingam"),
    ("CHG0038108", "Update Concur Approval Workflow Logic", "Normal", "Ramanathan Sockalingam"),
    ("CHG0037668", "DEVELOPMENT - HCM to Concur - EMPLOYEE FILE", "Standard", "Ramanathan Sockalingam"),
    ("CHG0037667", "MULTI-ASSIGNMENT FILE - Priority 2", "Standard", "Parvathi Arun"),
    ("CHG0037666", "APPROVER FILE - Priority 1 - HCM to Concur", "Standard", "Parvathi Arun"),
    ("CHG0037662", "Missing on employees HCM employee file", "Normal", "Ramanathan Sockalingam"),
    ("CHG0037661", "Weekly HCM Employee Files to upload", "Normal", "Ramanathan Sockalingam"),
    ("CHG0037415", "Quick Links on Intranet Homepage", "Standard", "Ariana Krohnfeldt"),
    ("CHG0036880", "Approver changes in HCM report request", "Standard", "Zach Zeller"),
    ("CHG0036364", "Add MoveWorks Connector to Concur", "Normal", "Beth Vanderheiden"),
    ("CHG0033572", "Add Cost Center 1105 pp-Adv", "Normal", "Unassigned"),
    ("CHG0033130", "Prepare technical documentation (NetSuite prod)", "Normal", "Jitin Xavier"),
    ("CHG0031988", "Concur POC integration w. HCM", "Normal", "Unassigned"),
    ("CHG0031877", "HCM to Concur integration", "Normal", "Subham Sawana"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], chg_url(num), num)
    cells[1].text = desc; cells[2].text = typ; cells[3].text = who

# Incidents
add_h(doc, "ServiceNow — Incidents", 1)
doc.add_paragraph("Only 4 Concur-tagged incidents in window — all resolved/closed, P3–P4. Volume is lower than Coupa, partly because access requests route through RITM.")
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Number", "Short Description", "Priority", "Assignee"])
for num, desc, pri, who in [
    ("INC1089903", "FW: Reminder Request for Item Delivery", "P4", "Ganesh Karaspalli"),
    ("INC1089272", "Need access to Concur and email account", "P3", "Harshitha Mahalingappa"),
    ("INC1087663", "Lost concur access", "P4", "Ashwini Bhamore"),
    ("INC1087490", "User needs assistance to access email and Concur", "P4", "Bharat Reddy"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], inc_url(num), num)
    cells[1].text = desc; cells[2].text = pri; cells[3].text = who

# Problems
add_h(doc, "ServiceNow — Problems", 1)
p = doc.add_paragraph(style="List Bullet")
add_hyperlink(p, prb_url("PRB0040564"), "PRB0040564")
p.add_run(" — Concur Outage (Closed/Resolved, P3-Moderate, IT Service Desk, Andrew Sanchez)")

# RITM
add_h(doc, "ServiceNow — Request Items", 1)
doc.add_paragraph(
    "50 RITMs sampled; volume dominated by code/dept/cost-center additions, expense-report help, and access requests. "
    "Most are auto-handled (Unassigned routing), with a handful escalated to Move Works."
)
doc.add_paragraph("Common request patterns:", style="Intense Quote")
for s in [
    "Cost center / department code adds (1105, 1315, 1541, 4202, PP1160, etc.)",
    "Site adds (Tillamook OR HM/EM, Queen of the Valley, Union Regional, Sutter Davis, Ambassador, Wynn Hospital, Knox Community, Northstar)",
    "Concur access (new hire, lost access, partner access)",
    "Expense report help / approval / delegate setup",
    "Mobile / iPad app issues",
]:
    doc.add_paragraph(s, style="List Bullet")

# KB
add_h(doc, "ServiceNow — Knowledge Base", 1)
doc.add_paragraph("23 articles found — Concur-specific articles below (the rest are leave/policy articles that mention Concur tangentially).")
t = doc.add_table(rows=1, cols=3); set_table_borders(t)
add_header_row(t, ["Article", "Title", "Knowledge Base"])
for num, title_text, kb in [
    ("KB0011025", "Concur - Knowledge Base Article", "People Operations"),
    ("KB0011026", "Concur - Create an Expense Report", "People Operations"),
    ("KB0011027", "Concur - Submit a Corporate Credit Card Report", "People Operations"),
    ("KB0011028", "Concur - Assign a Delegate", "People Operations"),
    ("KB0011029", "Concur - Act as a Delegate", "People Operations"),
    ("KB0011030", "Concur - Professional Expense Report", "People Operations"),
    ("KB0011031", "Concur - Using the Mobile App", "People Operations"),
    ("KB0011063", "Concur Operating System and Browser Requirements", "IT"),
    ("KB0011149", "Concur - Approve Expense Report", "People Operations"),
    ("KB0011353", "Concur - Review past PER expenses for CME Balance", "People Operations"),
    ("KB0011978", "Concur - Updating Direct Deposit banking info", "People Operations"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], kb_url(num), num)
    cells[1].text = title_text; cells[2].text = kb

# Jira
add_h(doc, "Jira (50 issues; projects: PROD, EDA, INFOSEC, SDOP, ENTAPP, WP)", 1)

add_h(doc, "Active 2026 Work", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in [
    ("PROD-12118", "AP Connections for Concur — Security Architecture Review", "To Do", "Budharam, Gopalakrishna"),
    ("EDA-6877", "AN Invoicing — Build/Test Concur Ingestion Pipeline", "Done", "Vedavyas Prabhu, Avinash"),
    ("EDA-7003", "[Placeholder] AN Invoicing API Connection — Concur Data Phase 2", "Backlog", "Vedavyas Prabhu, Avinash"),
    ("EDA-6238", "Data Discovery: AN Invoicing Calculator", "Done", "Bodoh, Leslie"),
    ("EDA-6761", "AN Data Mapping — Pull new AN tables from IDX Stage to Snowflake", "Done", "Kakulavaram, Shanthanreddy"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira_url(key), key)
    cells[1].text = title_text; cells[2].text = status; cells[3].text = owner

add_h(doc, "Vendor Risk Assessment / Security", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in [
    ("PROD-9737", "VRA | Concur Solutions — Phase 1", "Done", "Dhakad, Lokesh"),
    ("PROD-10115", "VRA | Concur Solutions — Phase 2", "Done", "Dhakad, Lokesh"),
    ("INFOSEC-6954", "Annual Re-Risk Assessment for Concur Solutions (SAP)", "Done", "Milan Kojic"),
    ("INFOSEC-11499", "Concur Email Issues", "Done", "Carter, Bill"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira_url(key), key)
    cells[1].text = title_text; cells[2].text = status; cells[3].text = owner

add_h(doc, "HCM ↔ Concur Integration (heavily Cancelled)", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in [
    ("PROD-3438", "HCM weekly files for employees [CHG0037540]", "Done", "Ramanathan Sockalingam"),
    ("PROD-4424", "HCM to Concur — EMPLOYEE FILE Priority 3 [CHG0037668]", "Done", "Ramanathan Sockalingam"),
    ("PROD-4546", "Concur profile missing [CHG0038031]", "Done", "Arun, Parvathi"),
    ("PROD-2134", "Decommission Extracts from HCM [CHG0036790]", "Done", "Ramanathan Sockalingam"),
    ("PROD-2182", "Update logic for HCM/Concur Integration File", "Done", "Arun, Parvathi"),
    ("PROD-4286", "Ensure Start Date is Present (Column O)", "Done", "—"),
    ("PROD-4057", "Missing on employees HCM file [CHG0037662]", "Blocked", "—"),
    ("PROD-3636", "Changes to Employee File HCM and Concur", "Cancelled", "Ramanathan Sockalingam"),
    ("PROD-3634", "Approver File HCM and Concur [CHG0037666]", "Cancelled", "Arun, Parvathi"),
    ("PROD-3635", "HCM to Concur — Multi-Assignment File", "Cancelled", "Arun, Parvathi"),
    ("PROD-4245", "Limitation to Transactions per file", "Cancelled", "Arun, Parvathi"),
    ("PROD-4285", "Validate Expense Approver Employee ID", "Cancelled", "—"),
    ("PROD-4650", "Update Concur Approval Workflow Logic [CHG0038108]", "Cancelled", "Arun, Parvathi"),
    ("PROD-5972", "DEVELOPMENT - HCM to Concur — EMPLOYEE FILE P3", "Cancelled", "Ramanathan Sockalingam"),
    ("PROD-10885", "Truncate HCM Dept Names >60 chars [CHG0039419]", "Cancelled", "Ramanathan Sockalingam"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira_url(key), key)
    cells[1].text = title_text; cells[2].text = status; cells[3].text = owner

add_h(doc, "Open / Backlog (Zenoti, Pharmacy, Other)", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in [
    ("PROD-4703", "Concur Use Case — Compound Pharmacy Purchasing (Epic)", "To Do", "UNASSIGNED"),
    ("PROD-4704", "Spike: Feasibility Zenoti ↔ Concur Integration", "To Do", "Kaparthi, Shirisha"),
    ("PROD-4287", "Change File Format (linked to PROD-3636)", "To Do", "—"),
    ("PROD-4247", "Change File Format Multiple Assignments (Concur)", "To Do", "—"),
    ("PROD-4248", "Rename Multiassignment file worksheet tab", "To Do", "—"),
    ("PROD-4249", "Validate Character length in item name field", "To Do", "—"),
    ("PROD-2215", "Approver changes in HCM report request [CHG0036880]", "To Do", "Integration Services"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira_url(key), key)
    cells[1].text = title_text; cells[2].text = status; cells[3].text = owner

add_h(doc, "Operations / Connector / Misc", 2)
t = doc.add_table(rows=1, cols=4); set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in [
    ("PROD-6106", "Convert Email to VERA Ticket & Assign to Concur Admin [CHG0038493]", "Done", "Edmund Trinidad"),
    ("SDOP-1136", "Submit Jan 2026 Credit Card Expenses in Concur", "Done", "Beth Vanderheiden"),
    ("SDOP-63", "Creator Studio — Setup Concur Connector", "Done", "Beth Vanderheiden"),
    ("PROD-3198", "Quick Links on Intranet Homepage by Population [CHG0037415]", "Done", "Balashchenko, Ekaterina"),
    ("PROD-3397", "SPIKE: Quick Links for Intranet by AD Group", "Done", "Balashchenko, Ekaterina"),
    ("PROD-2229", "[INTERRUPT] Produce test files for Concur/ADP/NetSuite", "Done", "Zeller, Zachary"),
    ("PROD-894", "Project code field Coupa/Concur → NetSuite [CHG0035232]", "Done", "Jitin Xavier"),
    ("ENTAPP-103", "Project code field Coupa/Concur → NetSuite", "Done", "Jitin Xavier"),
    ("PROD-2133", "Cost center consolidation (Epic)", "Done", "Kaparthi, Shirisha"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira_url(key), key)
    cells[1].text = title_text; cells[2].text = status; cells[3].text = owner

# SharePoint
add_h(doc, "SharePoint", 1)

add_h(doc, "Primary Sites", 2)
for name, url, note in [
    ("Concur Page (sp283)", "https://vituity.sharepoint.com/sites/sp283/SitePages/Concur.aspx", "Concur landing page (last modified 2026-04-10)"),
    ("Concur Documents (sp283)", "https://vituity.sharepoint.com/sites/sp283/Shared%20Documents/Concur", "Job aids and reference docs"),
    ("Travel Hub (sp480)", "https://vituity.sharepoint.com/sites/sp480", "Concur Connect Outlook, Travel Policy, Rewards"),
    ("Manager Resources Concur (sp373)", "https://vituity.sharepoint.com/sites/sp373/Manager%20Resources/Concur%20Best%20Practices.pdf", "Concur Best Practices PDF"),
    ("Expenses & Reimbursement (sp492)", "https://vituity.sharepoint.com/sites/sp492/Public%20Documents/All%20Project/Concur%20-%20Expenses%20%26%20Reimbursement", "Public expense reference docs"),
    ("AN Invoicing — Concur Files (sp604)", "https://vituity.sharepoint.com/sites/sp604/AN%20Invoicing%20Project/Concur%20Files", "Active — last modified 2026-03-25"),
    ("AP Documentation Concur (sp315)", "https://vituity.sharepoint.com/sites/sp315/Documentation%20and%20Process/AP/Concur", "AP process documentation"),
    ("ITS Legal & Compliance Concur", "https://vituity.sharepoint.com/sites/PM-ITSAll-DEPT/ITS%20Files%20Legal%20and%20Compliance/Concur", "Legal/compliance docs"),
    ("Leadership Dev Program Concur (sp621)", "https://vituity.sharepoint.com/sites/sp621/SiteAssets/SitePages/Leadership-Development-Program(1)(1)(1)/Concur-Tips.pdf", "Concur tips for LDP cohort"),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name, bold=True)
    p.add_run(f" — {note}")

add_h(doc, "Key Documents", 2)
for name, url in [
    ("Concur - Create an Expense Report.pdf", "https://vituity.sharepoint.com/sites/sp283/Shared%20Documents/Concur/Job%20Aids/Create%20an%20Expense%20Report%20-%20Concur%20Job%20Aid.pdf"),
    ("Concur - Approve an Expense Report.pdf", "https://vituity.sharepoint.com/sites/sp283/Shared%20Documents/Concur/Job%20Aids/Approve%20an%20Expense%20Report%20-Concur%20Job%20Aid.pdf"),
    ("Concur - Assign a Delegate.pdf", "https://vituity.sharepoint.com/sites/sp283/Shared%20Documents/Concur/Job%20Aids/Assign%20a%20Delegate%20-%20Concur%20Job%20Aid.pdf"),
    ("Concur Knowledge Base Article.pdf", "https://vituity.sharepoint.com/sites/sp283/Shared%20Documents/Concur/Job%20Aids/Concur%20Knowledge%20Base%20Article.pdf"),
    ("Concur - What Is It (KB).pdf", "https://vituity.sharepoint.com/sites/sp492/Public%20Documents/All%20Project/Concur%20-%20Expenses%20%26%20Reimbursement/Concur%20-%20What%20Is%20It_Knowledge%20Base%20Article.pdf"),
    ("Concur Best Practices.pdf", "https://vituity.sharepoint.com/sites/sp373/Manager%20Resources/Concur%20Best%20Practices.pdf"),
    ("Concur-Tips.pdf", "https://vituity.sharepoint.com/sites/sp621/SiteAssets/SitePages/Leadership-Development-Program(1)(1)(1)/Concur-Tips.pdf"),
    ("Installing the Concur Mobile App.pdf", "https://vituity.sharepoint.com/sites/sp621/SiteAssets/SitePages/Leadership-Development-Program(1)(1)(1)/Installing-the-Concur-Mobile-App---Concur-Job-Aid--1-.pdf"),
    ("Linking your email address to Concur Profile.pdf", "https://vituity.sharepoint.com/sites/sp621/SiteAssets/SitePages/Leadership-Development-Program(1)(1)(1)/Linking-your-email-address-to-Concur-Profile--Concur-Job-Aid--1-.pdf"),
    ("Concur - How to Assign Delegate.pdf (Ambassador Orientation)", "https://vituity.sharepoint.com/sites/HMAmbassadorTeamsChannel-AH/Shared%20Documents/Ambassador%20Orientation/Providers/Concur%20-%20How%20to%20Assign%20Delegate.pdf"),
    ("How to Expense.docx (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Expense%20.docx"),
    ("Finance — NetSuite/Coupa/Navan/Floqast Product Change Matrix.xlsx", "https://vituity-my.sharepoint.com/personal/fallond_vituity_com/Documents/Documents/Change%20Management/Application%20Review/Finance%20-%20Netsuite%20Coupa%20Navan%20Floqast%20Product%20Change%20Matrix.xlsx"),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name)

# Footer
doc.add_paragraph()
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run("Sources: ServiceNow CMDB, Change/Incident/Problem/RITM/KB tables; Jira (medamerica.atlassian.net); SharePoint search.").italic = True

# Save
out_dir = r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\Application Review"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, f"Concur Cross-System Inventory {date.today().isoformat()}.docx")
doc.save(out_path)
print(out_path)
