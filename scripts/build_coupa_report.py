"""Build Coupa cross-system inventory Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

SN = "https://vituity.service-now.com"


def chg_url(num: str) -> str:
    return f"{SN}/change_request.do?sysparm_query=number={num}"


def inc_url(num: str) -> str:
    return f"{SN}/incident.do?sysparm_query=number={num}"


def prb_url(num: str) -> str:
    return f"{SN}/problem.do?sysparm_query=number={num}"


def kb_url(num: str) -> str:
    return f"{SN}/kb_view.do?sysparm_article={num}"


def jira_url(key: str) -> str:
    return f"https://medamerica.atlassian.net/browse/{key}"


def add_hyperlink(paragraph, url, text, bold=False):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_header_row(table, headers):
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], "1F4E79")


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


# ── Document ────────────────────────────────────────────────────────────────
doc = Document()

# Style defaults
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Coupa Application — Cross-System Inventory")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(
    f"Sources: ServiceNow • Jira • SharePoint   |   Generated {date.today().isoformat()}   |   Owner: Dan Fallon"
).italic = True

# ── Executive Summary ───────────────────────────────────────────────────────
add_h(doc, "Executive Summary", 1)
bullets = [
    ("Owner", "Enterprise Applications group; admins Viresh Chibber, Brian Ouderkirk; upgrades Balaji Gavate."),
    ("Cadence", "Quarterly Coupa releases (R33 → R44 visible) plus recurring vendor / OS / security maintenance."),
    ("Top Integrations", "NetSuite, Concur, Workato, Zenoti (planned), Moveworks, ServiceNow, Sentinel."),
    ("Active 2026 Work", "AN Invoicing data pipelines (EDA), Moveworks API integration (SDOP), InvoiceSmash AI rollout (CHG0036978 still New)."),
    ("Documentation Home", "SharePoint sp280 + Facilities KB (KB0011046–49, KB0011366)."),
]
for label, body in bullets:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(body)

# ── Hotspots ────────────────────────────────────────────────────────────────
add_h(doc, "Cross-System Hotspots", 1)
hotspots = [
    ("CHG0036978 is zombified", "Jira Infosec phases done but CHG never closed/scheduled (~19 months in New). Worth pinging Viresh/Reina to close as superseded or cancel.", "CHG0036978", chg_url("CHG0036978")),
    ("EDA-7260 due today (2026-04-27)", "Add Coupa Name to contract master — confirm with Leslie Bodoh.", "EDA-7260", jira_url("EDA-7260")),
    ("Zenoti ↔ Coupa epic has no assignee", "PROD-4658 (MOOV–Coupa Integration) — needs an owner or formal deferral.", "PROD-4658", jira_url("PROD-4658")),
    ("EDA-4651 'HOLD'", "Wound Care / MOOV ingestion pieces orphaned; AN moved on without them.", "EDA-4651", jira_url("EDA-4651")),
]
for title_text, body, link_text, url in hotspots:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(f"{title_text} — ")
    r.bold = True
    p.add_run(body + " (")
    add_hyperlink(p, url, link_text)
    p.add_run(")")

# ── ServiceNow CMDB ─────────────────────────────────────────────────────────
add_h(doc, "ServiceNow CMDB", 1)
doc.add_paragraph("Three Configuration Items registered for Coupa.")
t = doc.add_table(rows=1, cols=4)
set_table_borders(t)
add_header_row(t, ["Name", "Class", "Status", "Support Group"])
cmdb = [
    ("Coupa", "Application", "Operational", "—"),
    ("Coupa", "Business Application", "Operational", "Enterprise Applications"),
    ("Coupa", "Business Service", "Operational", "—"),
]
for row in cmdb:
    cells = t.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

# ── ServiceNow Changes ──────────────────────────────────────────────────────
add_h(doc, "ServiceNow — Change Requests (47 returned)", 1)

add_h(doc, "Open / In-Flight", 2)
p = doc.add_paragraph()
add_hyperlink(p, chg_url("CHG0036978"), "CHG0036978", bold=True)
p.add_run(" — Enabling Coupa InvoiceSmash Scanned Invoice Extraction AI (Google Vision) — Phase 2. ")
p.add_run("State: New").bold = True
p.add_run("  |  Created 2024-09-25  |  Risk: Low  |  Assignee: Viresh Chibber")

doc.add_paragraph("CCB-readiness gaps:", style="Intense Quote")
gaps = [
    "Implementation Plan: TBD",
    "Backout / Test Plan: template text only, no actual content",
    "R&I: every question unanswered (PHI/PII, business-critical, outage hours, QA testing, vendor support, sites)",
    "No Planned Start/End — cannot be authorized",
    "No journal activity — change appears abandoned",
]
for g in gaps:
    doc.add_paragraph(g, style="List Bullet")

add_h(doc, "Recurring Patterns", 2)
patterns = [
    "Quarterly platform upgrades (R33 → R44) — all Normal, owned by Brian Ouderkirk / Viresh Chibber / Balaji Gavate",
    "Vendor / OS / Security maintenance windows — recurring Normal changes",
    "Sandbox refreshes (CHG0034269, CHG0036207)",
    "Canceled cluster (Product Mgmt — PM, Shirisha Kaparthi): CHG0038309, CHG0036828, CHG0036290, CHG0035544, CHG0035354, CHG0035276, CHG0034968, CHG0034889",
]
for pp in patterns:
    doc.add_paragraph(pp, style="List Bullet")

add_h(doc, "Notable Recent Closed Changes", 2)
t = doc.add_table(rows=1, cols=4)
set_table_borders(t)
add_header_row(t, ["Number", "Short Description", "Type", "Assignee"])
notable = [
    ("CHG0039061", "Coupa [PROD]: Upgrade to R44", "Normal", "Viresh Chibber"),
    ("CHG0038639", "Coupa [PROD]: Upgrade to R43", "Normal", "Brian Ouderkirk"),
    ("CHG0038128", "Coupa [PROD]: Upgrade to R42", "Normal", "Brian Ouderkirk"),
    ("CHG0037558", "Coupa [PROD]: Upgrade to R41", "Normal", "Viresh Chibber"),
    ("CHG0037277", "Skip Payments without Check Number", "Normal", "Jitin Xavier"),
    ("CHG0036851", "Coupa [PROD]: Upgrade to R40", "Normal", "Balaji Gavate"),
    ("CHG0036316", "Coupa [PROD]: Upgrade to R39", "Normal", "Balaji Gavate"),
    ("CHG0035570", "Coupa [PROD]: Upgrade to R38", "Normal", "Brian Ouderkirk"),
    ("CHG0035232", "Project code field Coupa↔Concur↔NetSuite", "Normal", "Jitin Xavier"),
    ("CHG0035089", "Coupa [PROD]: Vendor OS & Security", "Normal", "Brian Ouderkirk"),
    ("CHG0034552", "Coupa [PROD]: Vendor enabling MFA", "Normal", "Brian Ouderkirk"),
    ("CHG0034501", "Coupa [PROD]: Vendor OS & Security", "Normal", "Brian Ouderkirk"),
    ("CHG0034437", "Coupa [PROD]: Upgrade to R37", "Normal", "Brian Ouderkirk"),
    ("CHG0034269", "Coupa Sandbox Refresh", "Normal", "Derrick Chin"),
    ("CHG0033285", "Coupa [PROD]: Upgrade to R33", "Normal", "Brian Ouderkirk"),
]
for num, desc, typ, who in notable:
    cells = t.add_row().cells
    p0 = cells[0].paragraphs[0]
    add_hyperlink(p0, chg_url(num), num)
    cells[1].text = desc
    cells[2].text = typ
    cells[3].text = who

# ── ServiceNow Incidents ────────────────────────────────────────────────────
add_h(doc, "ServiceNow — Incidents", 1)
doc.add_paragraph(
    "47 incidents returned, all P3–P5. Volume dominated by access/permission requests (Coupa Administrator, "
    "IT Service Desk, Purchasing). No P1/P2 incidents in window."
)
doc.add_paragraph("Top assignees:", style="Intense Quote")
for who in ["Viresh Chibber", "Robert Rousseau", "Sushant Kamble", "Ashwini Bhamore", "Ganesh Karaspalli"]:
    doc.add_paragraph(who, style="List Bullet")

doc.add_paragraph("Functional issue samples:", style="Intense Quote")
inc_samples = [
    ("INC1084515", "Coupa Smash not working"),
    ("INC1084036", "Coupa Smash window not showing up"),
    ("INC1084760", "Coupa - Amazon Delivery Address"),
    ("INC1080669", "Coupa Report: Blanket Expiry - Purchasing"),
]
for num, desc in inc_samples:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, inc_url(num), num)
    p.add_run(f" — {desc}")

# ── Problems ────────────────────────────────────────────────────────────────
add_h(doc, "ServiceNow — Problems", 1)
p = doc.add_paragraph(style="List Bullet")
add_hyperlink(p, prb_url("PRB0040571"), "PRB0040571")
p.add_run(" — Coupa Outage (Closed/Resolved, P4, IT Service Desk, Andrew Sanchez)")

# ── Knowledge Base ──────────────────────────────────────────────────────────
add_h(doc, "ServiceNow — Knowledge Base", 1)
t = doc.add_table(rows=1, cols=2)
set_table_borders(t)
add_header_row(t, ["Article", "Title"])
kbs = [
    ("KB0011046", "Coupa Approver Guide"),
    ("KB0011047", "Coupa Frequently Asked Questions"),
    ("KB0011048", "Coupa Confirmation of Receipt"),
    ("KB0011049", "Coupa Overview Guide"),
    ("KB0011366", "Coupa Requestor Guide"),
]
for num, title_text in kbs:
    cells = t.add_row().cells
    p0 = cells[0].paragraphs[0]
    add_hyperlink(p0, kb_url(num), num)
    cells[1].text = title_text

# ── Jira ────────────────────────────────────────────────────────────────────
add_h(doc, "Jira (50 issues; projects: EDA, PROD, INFOSEC, SDOP, ENTAPP, TOP)", 1)

add_h(doc, "Active 2026 — AN Invoicing Pipeline (parent EDA-6237)", 2)
t = doc.add_table(rows=1, cols=5)
set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner", "Due"])
an = [
    ("EDA-7260", "Add Coupa Name to contract master", "Selected for Dev", "Bodoh, Leslie", "2026-04-27 (today)"),
    ("EDA-7266", "Coupa Shift Matching Streamlit Page", "Selected for Dev", "Tarasia, Dev", "2026-05-01"),
    ("EDA-6876", "Build/Test Coupa Ingestion Pipeline", "Done", "Vedavyas Prabhu", "—"),
    ("EDA-7002", "Coupa Supplier↔Provider Crosswalk", "Done", "Jagadeesh, Manasa", "—"),
    ("EDA-7009", "Provider Name Normalization", "Done", "Jagadeesh, Manasa", "—"),
]
for key, title_text, status, owner, due in an:
    cells = t.add_row().cells
    p0 = cells[0].paragraphs[0]
    add_hyperlink(p0, jira_url(key), key)
    cells[1].text = title_text
    cells[2].text = status
    cells[3].text = owner
    cells[4].text = due

add_h(doc, "Moveworks / Otto Provisioning", 2)
mw = [
    ("SDOP-1289", "Send Coupa API secret to Moveworks via Dilenea", "Done", "Beth Vanderheiden"),
    ("SDOP-1114", "Moveworks Coupa API Integration — Discovery & Implementation", "Done", "Beth Vanderheiden"),
    ("SDOP-1252", "Otto Software Provisioning Setup Batch 3 (Coupa + 4 others)", "In Progress", "Spengler, Dan"),
]
t = doc.add_table(rows=1, cols=4)
set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in mw:
    cells = t.add_row().cells
    p0 = cells[0].paragraphs[0]
    add_hyperlink(p0, jira_url(key), key)
    cells[1].text = title_text
    cells[2].text = status
    cells[3].text = owner

add_h(doc, "Security / Vendor Eval / Sentinel", 2)
sec = [
    ("INFOSEC-12234", "Vendor eval-Scope | Coupa | Phase 1", "Done", "Ejesieme, Brian"),
    ("INFOSEC-12299", "Vendor eval-Scope | Coupa | Phase 2", "Done", "Ejesieme, Brian"),
    ("INFOSEC-12400", "Coupa Review and Documentation", "Done", "Satyachaitanya, Chaitanya"),
    ("PROD-6013", "Sentinel Data Source Integration — Coupa Dev/Test", "Done", "Morris, Kevin"),
    ("PROD-6269", "Sentinel Data Source Integration — Coupa Production", "Done", "Morris, Kevin"),
    ("PROD-4220", "InvoiceSmash AI / Google Vision — Infosec Architecture review", "Done", "Satyachaitanya"),
    ("PROD-4575", "InvoiceSmash AI Phase 2 — Infosec POV", "Done", "Wyatt Silva"),
    ("PROD-2611", "InvoiceSmash AI Phase 2 (parent)", "Done", "Wyatt Silva"),
    ("PROD-11252", "Sentinel Data Source Integration — QGenda User & Audit Logs", "Done", "Morris, Kevin"),
    ("PROD-10289", "Coupa + ServiceNow Integration | HSR | SAR", "Done", "Budharam"),
    ("PROD-11720", "Coupa Audit Report Date Filter Remediation", "Done", "Morris, Kevin"),
    ("PROD-11329", "VRA | Coupa — Sprint 167", "Done", "Ejesieme, Brian"),
]
t = doc.add_table(rows=1, cols=4)
set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in sec:
    cells = t.add_row().cells
    p0 = cells[0].paragraphs[0]
    add_hyperlink(p0, jira_url(key), key)
    cells[1].text = title_text
    cells[2].text = status
    cells[3].text = owner

add_h(doc, "Open / Backlog (Zenoti / MOOV / Wound Care)", 2)
backlog = [
    ("PROD-4658", "MOOV — Coupa Integration (Epic)", "To Do", "UNASSIGNED"),
    ("PROD-4663", "Sync Purchase Orders Coupa → Zenoti", "To Do", "Kaparthi, Shirisha"),
    ("PROD-4664", "Sync Inventory Receipts Zenoti → Coupa", "To Do", "Kaparthi, Shirisha"),
    ("PROD-4702", "Spike: Feasibility Zenoti ↔ Coupa Integration", "To Do", "Kaparthi, Shirisha"),
    ("PROD-4659", "MOOV — Marketo Integration (Epic)", "To Do", "UNASSIGNED"),
    ("EDA-4651", "HOLD — COUPA data ingestion (AN/Wound Care/MOOV)", "Backlog", "Vedavyas Prabhu"),
    ("EDA-4960", "AN Locum Utilization & Cost Forecasting (Epic)", "Backlog", "Lee, Joyce"),
]
t = doc.add_table(rows=1, cols=4)
set_table_borders(t)
add_header_row(t, ["Key", "Title", "Status", "Owner"])
for key, title_text, status, owner in backlog:
    cells = t.add_row().cells
    p0 = cells[0].paragraphs[0]
    add_hyperlink(p0, jira_url(key), key)
    cells[1].text = title_text
    cells[2].text = status
    cells[3].text = owner

add_h(doc, "Zenoti ↔ Coupa Field Map (from PROD-4658)", 2)
t = doc.add_table(rows=1, cols=3)
set_table_borders(t)
add_header_row(t, ["Mapping", "Zenoti", "Coupa"])
fm = [
    ("Object", "Inventory Receipt", "Purchase Order (PO)"),
    ("Object", "Inventory Adjustment / Goods Received", "Item Receipt Confirmation"),
    ("Field", "External PO ID", "PO Number"),
    ("Field", "Inventory Item Code", "Item SKU"),
    ("Field", "Quantity Added", "Quantity Received"),
    ("Field", "Inventory Date", "Received Date"),
    ("Field", "Storage Location", "Location Code"),
]
for a, b, c in fm:
    cells = t.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

# ── SharePoint ──────────────────────────────────────────────────────────────
add_h(doc, "SharePoint", 1)

add_h(doc, "Primary Sites", 2)
sites = [
    ("Coupa SharePoint (sp280)", "https://vituity.sharepoint.com/sites/sp280", "Top-level Coupa hub"),
    ("Coupa Training (sp281)", "https://vituity.sharepoint.com/sites/sp281", "Training videos, purchase thresholds"),
    ("AN Invoicing — Coupa Files", "https://vituity.sharepoint.com/sites/sp604/AN%20Invoicing%20Project/Coupa%20Files", "Active (last modified 2026-03-25); includes Coupa Suppliers.xlsx"),
    ("Coupa Release Readiness WG", "https://vituity.sharepoint.com/sites/CoupaReleasereadiness-WG", "Release-prep workgroup site + OneNote notebook"),
    ("Application Operations — Coupa Tech Docs (sp686)", "https://vituity.sharepoint.com/sites/sp686/Shared%20Documents/Application%20Operations%20-%20Product%20Operations/Product%20Technical%20Documentation/Coupa", "Product technical documentation"),
    ("ITS Coupa Reports", "https://vituity.sharepoint.com/sites/PM-ITSAll-DEPT/Budget/ITS%20Expense%20and%20Capital%20Budget/Coupa%20reports", "ITS budget / invoice reporting"),
    ("Legal & Compliance — Coupa", "https://vituity.sharepoint.com/sites/PM-ITSAll-DEPT/ITS%20Files%20Legal%20and%20Compliance/Coupa", "Legal / compliance docs"),
    ("AP Documentation — Coupa (sp315)", "https://vituity.sharepoint.com/sites/sp315/Documentation%20and%20Process/AP/Coupa", "AP process documentation"),
]
for name, url, note in sites:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name, bold=True)
    p.add_run(f" — {note}")

add_h(doc, "Key Documents", 2)
docs_list = [
    ("Coupa Overview Guide.pdf", "https://vituity.sharepoint.com/sites/sp280/VERA%20Documents/Coupa%20Overview%20Guide.pdf"),
    ("Coupa Approver Guide.pdf", "https://vituity.sharepoint.com/sites/sp280/VERA%20Documents/Coupa%20Approver%20Guide.pdf"),
    ("Coupa Requester Guide.pdf", "https://vituity.sharepoint.com/sites/sp280/VERA%20Documents/Coupa%20Requester%20Guide.pdf"),
    ("Coupa Frequently Asked Questions.pdf", "https://vituity.sharepoint.com/sites/sp280/VERA%20Documents/Coupa%20Frequently%20Asked%20Questions.pdf"),
    ("Coupa Confirmation of Receipt.pdf", "https://vituity.sharepoint.com/sites/sp280/VERA%20Documents/Coupa%20Confirmation%20of%20Receipt.pdf"),
    ("How to Delegate Approvals in Coupa.pdf", "https://vituity.sharepoint.com/sites/sp280/VERA%20Documents/How%20to%20Delegate%20Approvals%20in%20Coupa.pdf"),
    ("Coupa-User-Purchase-Thresholds-2025.pdf", "https://vituity.sharepoint.com/sites/sp281/SiteAssets/SitePages/COUPA-USER-TRAINING-VIDEOS/Coupa-User-Purchase-Thresholds-2025.pdf"),
    ("Coupa Suppliers.xlsx (sp604)", "https://vituity.sharepoint.com/sites/sp604/AN%20Invoicing%20Project/Reference%20Files/Coupa%20Suppliers.xlsx"),
    ("Finance — NetSuite/Coupa/Navan/Floqast Product Change Matrix.xlsx", "https://vituity-my.sharepoint.com/personal/fallond_vituity_com/Documents/Documents/Change%20Management/Application%20Review/Finance%20-%20Netsuite%20Coupa%20Navan%20Floqast%20Product%20Change%20Matrix.xlsx"),
    ("COUPA-USER-TRAINING-VIDEOS landing page", "https://vituity.sharepoint.com/sites/sp281/SitePages/COUPA-USER-TRAINING-VIDEOS.aspx"),
]
for name, url in docs_list:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name)

# ── Footer ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run(
    "Sources: ServiceNow CMDB, Change/Incident/Problem/RITM/KB tables; Jira (medamerica.atlassian.net); SharePoint search."
).italic = True

# ── Save ────────────────────────────────────────────────────────────────────
out_dir = r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\Application Review"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, f"Coupa Cross-System Inventory {date.today().isoformat()}.docx")
doc.save(out_path)
print(out_path)
