"""Build Navan / TripActions cross-system inventory Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
import os

SN = "https://vituity.service-now.com"
def chg(n): return f"{SN}/change_request.do?sysparm_query=number={n}"
def inc(n): return f"{SN}/incident.do?sysparm_query=number={n}"
def prb(n): return f"{SN}/problem.do?sysparm_query=number={n}"
def kb(n):  return f"{SN}/kb_view.do?sysparm_article={n}"
def jira(k): return f"https://medamerica.atlassian.net/browse/{k}"


def add_hyperlink(p, url, text, bold=False):
    part = p.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), rid)
    r = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    if bold: rPr.append(OxmlElement("w:b"))
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = text; t.set(qn("xml:space"), "preserve"); r.append(t)
    h.append(r); p._p.append(h)


def shade(cell, hex_):
    tc = cell._tc.get_or_add_tcPr()
    s = OxmlElement("w:shd"); s.set(qn("w:val"), "clear"); s.set(qn("w:color"), "auto"); s.set(qn("w:fill"), hex_)
    tc.append(s)


def borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    b = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        x = OxmlElement(f"w:{e}")
        x.set(qn("w:val"), "single"); x.set(qn("w:sz"), "4"); x.set(qn("w:color"), "BFBFBF")
        b.append(x)
    tblPr.append(b)


def header_row(t, headers):
    cells = t.rows[0].cells
    for i, h in enumerate(headers):
        cells[i].text = ""
        run = cells[i].paragraphs[0].add_run(h)
        run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cells[i], "1F4E79")


def add_h(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


doc = Document()
style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(11)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Navan (formerly TripActions) — Cross-System Inventory")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run(f"Sources: ServiceNow • Jira • SharePoint   |   Generated {date.today().isoformat()}   |   Owner: Dan Fallon").italic = True

# Executive Summary
add_h(doc, "Executive Summary", 1)
for label, body in [
    ("App Identity", "Navan and TripActions are the same application; rebranded in 2023. CMDB still has both CIs side-by-side (Navan = Application; TripActions = Business Application, Enterprise Applications support group)."),
    ("Function", "Vituity corporate travel & expense booking (flights, hotels, rental cars). Travel rewards integrated. Sends receipts/expenses into Concur."),
    ("Owner", "Application admin / config: Edmund Trinidad (Enterprise Applications). Day-to-day support: Coupa Administrator group (Leslie Yoshida, Viresh Chibber). Auth/Okta integration: Service Delivery Optimization (Beth Vanderheiden)."),
    ("Top Integrations", "HCM (employee data feed via LEAP), Okta (SSO + provisioning), Concur (receipt/expense flow), Outlook calendar."),
    ("Cadence", "SaaS-managed (vendor handles updates) — no patch-cadence change traffic. Just 6 changes total over 5 years."),
    ("Documentation Home", "SharePoint sp480 (Travel & Expense hub) + 2 KB articles + 2020 PMO project folder for original TripActions implementation."),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{label}: ").bold = True
    p.add_run(body)

# History timeline
add_h(doc, "History (Story Arc)", 1)
for year, body in [
    ("2020 Q1", "TripActions selected as Vituity's corporate travel platform. PM Michael Broome. Project kickoff 2020-01-17 (PMO Project 2020 TripActions folder). Functional Spec v.3 published 2020-03-14."),
    ("2020 Q1–Q2", "EA-18 / EA-19 epics for HCM file extract via LEAP integration to TripActions SFTP. Vamsi Naganathanahalli built the integration; CHG0030604 migrated it into production."),
    ("2020 Q2", "VERA Category added (SNO-115, Rosalinda Rafael); user-training docs published (Book Flights/Hotels, Manage Trips, Outlook Calendar connection)."),
    ("2022", "First Vendor Risk Assessment cycle (INFOSEC-6350 CES TripActions, Bill Carter; INFOSEC-7912 batch RA covering TripActions + others, Milan Kojic). 953 active users by April 2022."),
    ("2023 Jan", "TripActions rebrands to Navan. Internal communication page published (sp480 \"TripActions Rebranding\"). Documentation begins shifting names; CMDB Navan CI added later via CHG0036552."),
    ("2024 Aug", "Functional spec re-versioned by Parvathi Arun as \"Navan (FormerlyTripsAction) Techno-Functional Specification\" with HCM mapping job-detail updates."),
    ("2025 Jan–Feb", "\"Navan Travel Relaunch\" project (PMO managed, 18-day plan 2025-01-15 → 2025-02-07). Schedule archived to PWA Retirement folder."),
    ("2025 Nov", "\"Basic economy\" fares enabled for business travel (per sp480 Improvements-to-Travel page)."),
    ("2026 Jan–Mar", "Vendor Risk Assessment renewal: PROD-9739 + PROD-9926 Phase 1 (Lokesh Dhakad + Brian Ejesieme); PROD-10848 Phase 2 (Brian Ejesieme)."),
    ("2026 Mar–Apr", "SDOP-1245 Navan/Okta Username Alignment (Beth Vanderheiden, Done Apr 2026) — proactive remediation of login friction. CHG0039562 (Open, New) requesting a dedicated \"Travel\" assignment group to resolve incident-routing gap."),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{year} — ").bold = True
    p.add_run(body)

# Hotspots
add_h(doc, "Cross-System Hotspots", 1)
for title_text, body, link, url in [
    ("CHG0039562 open (New)", "Request to add new \"Travel\" assignment group — addresses long-standing routing gap where Navan tickets land with Coupa Administrator. Edmund Trinidad assigned.", "CHG0039562", chg("CHG0039562")),
    ("Two CIs for same app", "CMDB has both \"Navan\" (Application) and \"TripActions\" (Business Application) — never consolidated post-rebrand.", "CMDB", f"{SN}/cmdb_ci.do?sysparm_query=name=Navan"),
    ("PROD-4871 stale", "Single-user login To Do unassigned for ~1 year, parent change CHG0038147 was canceled — close as superseded.", "PROD-4871", jira("PROD-4871")),
    ("EA-19 blocked since 2020", "HCM file extract → TripActions SFTP — investigate whether still relevant or kill it; LEAP integration was migrated long ago via CHG0030604.", "EA-19", jira("EA-19")),
    ("Routing gap surfaces in incidents", "Most Navan incidents route to Coupa Administrator group (Leslie Yoshida); two P2-High incidents are \"General Inquiry\" — questionable severity assignment.", "PRB0041001", prb("PRB0041001")),
    ("SDOP-1245 closed the auth root cause", "Okta username alignment fixed an entire class of Navan login incidents — should reduce 33-incident / 38-RITM volume going forward.", "SDOP-1245", jira("SDOP-1245")),
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"{title_text} — ").bold = True
    p.add_run(body + " (")
    add_hyperlink(p, url, link); p.add_run(")")

# CMDB
add_h(doc, "ServiceNow CMDB", 1)
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Name", "Class", "Status", "Support Group"])
for n, c, s, g in [
    ("Navan", "Application", "Operational", "—"),
    ("TripActions", "Business Application", "Operational", "Enterprise Applications"),
]:
    cells = t.add_row().cells
    cells[0].text = n; cells[1].text = c; cells[2].text = s; cells[3].text = g

# Changes
add_h(doc, "ServiceNow — Change Requests (6 total)", 1)
t = doc.add_table(rows=1, cols=5); borders(t); header_row(t, ["Number", "Short Description", "State", "Type", "Assignee"])
for num, desc, state, typ, who in [
    ("CHG0039562", "Request to Add new \"Travel\" Assignment Group", "New", "Normal", "Edmund Trinidad"),
    ("CHG0038147", "Patti Cordle unable to login to Navan", "Canceled", "Normal", "Zach Olsen"),
    ("CHG0036552", "Add Navan as a Configuration Item in CMDB", "Closed", "Standard", "Edmund Trinidad"),
    ("CHG0035598", "VERA KB update", "Closed", "Standard", "Edmund Trinidad"),
    ("CHG0034775", "Add Application System Administrator", "Closed", "Standard", "Edmund Trinidad"),
    ("CHG0030604", "TripAction Interface Migration into Production", "Closed", "Normal", "Vamsi Naganathanahalli"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], chg(num), num)
    cells[1].text = desc; cells[2].text = state; cells[3].text = typ; cells[4].text = who

# Incidents
add_h(doc, "ServiceNow — Incidents (33)", 1)
doc.add_paragraph("All P3–P5 except 2 P2-High; no P1. Volume dominated by login/access. Routing pattern: most route to Coupa Administrator group (gap CHG0039562 is fixing).")
doc.add_paragraph("Notable / Open incidents:", style="Intense Quote")
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Number", "Short Description", "Priority/State", "Assignee"])
for num, desc, ps, who in [
    ("INC1085765", "Reservation link for Vituity Health", "P4 / On Hold", "Diane Hoover"),
    ("INC1082226", "Delegate Access for Mike Harrington", "P3 / In Progress", "Bharat Reddy"),
    ("INC1076760", "Radiology 2026 Malpractice Rates/Cost", "P3 / On Hold", "Jitin Xavier"),
    ("INC1075902", "General Inquiry on Other / Unknown", "P2-High / Closed", "Leslie Yoshida"),
    ("INC1075078", "Navan login impossible", "P3 / Closed", "Bharat Reddy"),
    ("INC1073934", "Resolve Navan access issue with error message", "P4 / On Hold", "Leslie Yoshida"),
    ("INC1066774", "Flight Credit Expiration - Leslie Yoshida", "P3 / Closed", "Leslie Yoshida"),
    ("INC1065238", "General Inquiry on People Operations Sys", "P2-High / Closed", "Leslie Yoshida"),
    ("INC1065136", "Navan - Showing as User Disabled", "P3 / Closed", "Andrew Sanchez"),
    ("INC1063495", "Navan cancellation funds question", "P3 / Closed", "Leslie Yoshida"),
    ("INC1071621", "Navan Information Gone", "P3 / Closed", "Leslie Yoshida"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], inc(num), num)
    cells[1].text = desc; cells[2].text = ps; cells[3].text = who

# Problems
add_h(doc, "ServiceNow — Problems (2)", 1)
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Number", "Short Description", "Priority", "Assignee"])
for num, desc, pri, who in [
    ("PRB0041001", "Navan App Access", "P3", "Andrew Sanchez (Service Delivery Optimization)"),
    ("PRB0040797", "Okta: Block Login from Anonymizers (Navan-related)", "P4", "Andrew Sanchez (Service Delivery Optimization)"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], prb(num), num)
    cells[1].text = desc; cells[2].text = pri; cells[3].text = who

# Request Items
add_h(doc, "ServiceNow — Request Items (38)", 1)
doc.add_paragraph("Volume dominated by access requests, Okta-credential resets, and \"Software Request - Navan\" auto-tickets. One open: RITM0045455 (Allow Access to Navan, Work in Progress).")
doc.add_paragraph("Common patterns:", style="Intense Quote")
for s in [
    "Navan / Okta credential alignment (RITM0033230, RITM0029677, RITM0031298, RITM0031846, etc.)",
    "Generic access requests (RITM0045299, RITM0044088, RITM0036026, RITM0035363, RITM0032198)",
    "Login / sign-in failures (RITM0044890, RITM0044853, RITM0029258)",
    "Concur + Navan combo access (RITM0036758, RITM0036237)",
    "Profile setup / delegate setup (RITM0044773 Michelle Pierce, RITM0039372 Natalie Perez delegate)",
    "TripActions-era ticket: RITM0019671",
]:
    doc.add_paragraph(s, style="List Bullet")

# KB
add_h(doc, "ServiceNow — Knowledge Base (2)", 1)
t = doc.add_table(rows=1, cols=2); borders(t); header_row(t, ["Article", "Title"])
for num, title_text in [
    ("KB0011988", "Vituity's Travel Tool - Navan (formerly TripActions)"),
    ("KB0010145", "Corporate Hotel Discounts (references Navan)"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], kb(num), num)
    cells[1].text = title_text

# Jira
add_h(doc, "Jira (13 issues)", 1)

add_h(doc, "Active 2026 — Auth & Okta Integration", 2)
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Key", "Title", "Status", "Owner"])
for k, ttl, st, who in [
    ("SDOP-1245", "Navan/Okta Username Alignment — Proactive Remediation", "Done (Apr 2026)", "Beth Vanderheiden"),
    ("SDOP-1134", "Investigate Okta app name change settings", "Done (Mar 2026)", "Dan Spengler"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira(k), k)
    cells[1].text = ttl; cells[2].text = st; cells[3].text = who

add_h(doc, "Vendor Risk Assessment", 2)
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Key", "Title", "Status", "Owner"])
for k, ttl, st, who in [
    ("PROD-9739", "VRA | TripActions — Phase 1", "Done (Jan 2026)", "Lokesh Dhakad"),
    ("PROD-9926", "Vendor Eval-Scope | TripActions | Phase 1", "Done (Jan 2026)", "Brian Ejesieme"),
    ("PROD-10848", "Vendor Eval-Scope | Navan | Phase 2", "Done (Mar 2026)", "Brian Ejesieme"),
    ("INFOSEC-7912", "Vendor Risk Analysis (TripActions among vendors)", "Done (2022)", "Milan Kojic"),
    ("INFOSEC-6350", "INT: CES TripActions", "Done (2022)", "Bill Carter"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira(k), k)
    cells[1].text = ttl; cells[2].text = st; cells[3].text = who

add_h(doc, "Open / Stale Backlog", 2)
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Key", "Title", "Status", "Owner"])
for k, ttl, st, who in [
    ("PROD-4871", "Patti Cordle Navan login (CHG0038147 — canceled)", "To Do (May 2025)", "UNASSIGNED"),
    ("EA-19", "HCM file extract development → TripAction SFTP", "Blocked (since 2020)", "Vamsi Naganathanahalli"),
    ("EA-18", "Epic — TripAction (original 2020 implementation)", "Backlog", "—"),
    ("AA-160", "TripActions POC (Epic)", "Backlog", "Shreya Melkote (reporter)"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira(k), k)
    cells[1].text = ttl; cells[2].text = st; cells[3].text = who

add_h(doc, "Historical Implementation (2020)", 2)
t = doc.add_table(rows=1, cols=4); borders(t); header_row(t, ["Key", "Title", "Status", "Owner"])
for k, ttl, st, who in [
    ("SNO-115", "TripActions VERA Category Add", "Done (2020)", "Rosalinda Rafael"),
    ("ENG-1442", "Make product site provision work", "Done (2020)", "Uladzislau Loputs"),
]:
    cells = t.add_row().cells
    add_hyperlink(cells[0].paragraphs[0], jira(k), k)
    cells[1].text = ttl; cells[2].text = st; cells[3].text = who

# SharePoint
add_h(doc, "SharePoint", 1)

add_h(doc, "Primary Site & Pages", 2)
for name, url, note in [
    ("Travel & Expense Hub (sp480)", "https://vituity.sharepoint.com/sites/sp480", "Active hub — Navan training, FAQ, profile setup, expense flow. Last modified 2025-10-27."),
    ("Travel and Expense page (sp480)", "https://vituity.sharepoint.com/sites/sp480/SitePages/Travel-and-Expense.aspx", "User-facing landing page (Hotel/Rental Car/Parking)."),
    ("Improvements-to-Travel (root)", "https://vituity.sharepoint.com/SitePages/Improvements-to-Travel.aspx", "Basic-economy fare rollout 2025-11."),
    ("Traveling for Business (sp480)", "https://vituity.sharepoint.com/sites/sp480/SitePages/Traveling-for-Business-.aspx", "TripActions adoption page (\"953 Vitans\" stat from 2022)."),
    ("TripActions Rebranding announcement (sp480)", "https://vituity.sharepoint.com/sites/sp480/SitePages/TripActions-Rebranding.aspx", "2023-01 internal comms about the Navan rebrand."),
    ("PMO 2020 TripActions Project folder", "https://vituity.sharepoint.com/sites/projectmanagementoffice/Project%20%202020%20%20TripActions", "Original implementation project artifacts (kickoff, spec, minutes, logical model)."),
    ("Application Operations — Navan Tech Docs (sp686)", "https://vituity.sharepoint.com/sites/sp686/Shared%20Documents/Application%20Operations%20-%20Product%20Operations/Product%20Technical%20Documentation/Oracle%20HCM", "Navan techno-functional spec (under Oracle HCM, owner Parvathi Arun)."),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name, bold=True)
    p.add_run(f" — {note}")

add_h(doc, "Implementation & Technical Documents (TripActions-era 2020)", 2)
for name, url in [
    ("TripActions Functional Specification.docx (PMO 2020-04)", "https://vituity.sharepoint.com/sites/projectmanagementoffice/Project%20%202020%20%20TripActions/TripActions%20Functional%20Specification.docx"),
    ("TripActions Project Kickoff Deck FINAL.pptx (PMO 2020-01)", "https://vituity.sharepoint.com/sites/projectmanagementoffice/Project%20%202020%20%20TripActions/011720_TripActions%20Project%20Kickoff%20Deck_FINAL.pptx"),
    ("TripActions Project Meeting Minutes (PMO)", "https://vituity.sharepoint.com/sites/projectmanagementoffice/Project%20%202020%20%20TripActions/TripActions%20Project%20Meeting%20Minutes"),
    ("TripActions_Logical_Model.pdf (PMO 2020-03)", "https://vituity.sharepoint.com/sites/projectmanagementoffice/Project%20%202020%20%20TripActions/TripActions_Logical_Model.pdf"),
    ("TripsAction Functional Specification.docx (HCM/LEAP folder)", "https://vituity.sharepoint.com/sites/PM-ITSEngineering-DEPT/Shared%20Documents/HCM/LEAP/TripsAction%20Functional%20Specification.docx"),
    ("Navan (Formerly TripsAction) Techno-Functional Spec.docx (sp686, updated 2024-08 by Parvathi)", "https://vituity.sharepoint.com/sites/sp686/Shared%20Documents/Application%20Operations%20-%20Product%20Operations/Product%20Technical%20Documentation/Oracle%20HCM/Navan%20(FormerlyTripsAction)%20Techno-Functional%20Specification.docx"),
    ("Navan Travel Relaunch — Schedule.pdf (PWA Retirement Archive)", "https://vituity.sharepoint.com/sites/projectmanagementoffice/PWA%20Retirement/Schedule%20Archive/Navan%20Travel%20Relaunch%20-%20Schedule.pdf"),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name)

add_h(doc, "User Training & Job Aids", 2)
for name, url in [
    ("Navan FAQ Sheet.docx (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/Navan%20FAQ%20Sheet.docx"),
    ("Navan FAQ Sheet.pdf (Ambassador Orientation)", "https://vituity.sharepoint.com/sites/HMAmbassadorTeamsChannel-AH/Shared%20Documents/Ambassador%20Orientation/Providers/Navan%20FAQ%20Sheet.pdf"),
    ("Navan Guide.pdf (Ambassador Orientation)", "https://vituity.sharepoint.com/sites/HMAmbassadorTeamsChannel-AH/Shared%20Documents/Ambassador%20Orientation/Providers/Navan%20Guide.pdf"),
    ("How to Login and Register on Navan.docx (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Login%20and%20Register%20on%20Navan.docx"),
    ("How to Set Up Your Profile on Navan.docx (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Set%20Up%20Your%20Profile%20on%20Navan.docx"),
    ("How to Manage Trips on Navan.docx (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Manage%20Trips%20on%20Navan.docx"),
    ("How to Redeem Rewards on Navan.docx (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Redeem%20Rewards%20on%20Navan.docx"),
    ("How to Expense.docx — Navan → Concur receipt flow", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Expense%20.docx"),
    ("How to Book Flights.docx (TripActions-era)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Book%20Flights.docx"),
    ("How to Book Hotels.docx (TripActions-era)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20Book%20Hotels.docx"),
    ("How to connect your Outlook Calendar.pdf", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/How%20to%20connect%20your%20Outlook%20Calendar.pdf"),
    ("Navan General User Training.mp4 (sp480)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/Forms/DispForm.aspx?ID=15"),
    ("TripActions Admin Training.mp4 (sp480, 2021)", "https://vituity.sharepoint.com/sites/sp480/Shared%20Documents/Forms/DispForm.aspx?ID=14"),
    ("Discount Matrix v5.pdf (sp281)", "https://vituity.sharepoint.com/sites/sp281/SiteAssets/SitePages/Employee-Discounts/Discount-Matrix_v5.pdf"),
    ("Corporate Rental Car Program.pdf (sp281)", "https://vituity.sharepoint.com/sites/sp281/Rental%20Car/Corporate%20Rental%20Car%20Program.pdf"),
    ("Concur-Tips.pdf (LDP, mentions Navan invoice extract)", "https://vituity.sharepoint.com/sites/sp621/SiteAssets/SitePages/Leadership-Development-Program(1)(1)(1)/Concur-Tips.pdf"),
    ("Finance — NetSuite/Coupa/Navan/Floqast Product Change Matrix.xlsx", "https://vituity-my.sharepoint.com/personal/fallond_vituity_com/Documents/Documents/Change%20Management/Application%20Review/Finance%20-%20Netsuite%20Coupa%20Navan%20Floqast%20Product%20Change%20Matrix.xlsx"),
]:
    p = doc.add_paragraph(style="List Bullet")
    add_hyperlink(p, url, name)

# Footer
doc.add_paragraph()
foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot.add_run("Sources: ServiceNow CMDB, Change/Incident/Problem/RITM/KB tables; Jira (medamerica.atlassian.net); SharePoint search.").italic = True

# Save
out_dir = r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\Application Review"
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, f"Navan TripActions Cross-System Inventory {date.today().isoformat()}.docx")
doc.save(out_path)
print(out_path)
