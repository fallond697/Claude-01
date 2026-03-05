import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SNOW_URL = "https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number="

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def set_cell_text(cell, text, bold=False, font_size=10, alignment=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

def set_cell_hyperlink(cell, text, url, font_size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    add_hyperlink(p, text, url)

def add_status_badge(cell, status):
    colors = {
        "PASS": ("2E7D32", "PASS"),
        "FLAG": ("E65100", "FLAG"),
        "FAIL": ("C62828", "FAIL"),
        "N/A": ("757575", "N/A"),
        "PASS w/ minor flags": ("558B2F", "PASS*"),
        "FLAG - needs CCB discussion": ("BF360C", "FLAG"),
    }
    color_hex, label = colors.get(status, ("757575", status))
    set_cell_text(cell, label, bold=True, font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor.from_string(color_hex))

def add_bold_run(paragraph, text, font_size=10):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    return run

def add_run(paragraph, text, font_size=10, color=None, bold=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    if bold:
        run.bold = True
    return run

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_heading('Change Enablement Compliance Review', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub, 'Dual-Framework Assessment', font_size=16, color=RGBColor(0x44, 0x44, 0x44))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(sub2, 'ITIL 4 Change Enablement  |  ISMS-STA-11.01-01', font_size=13, color=RGBColor(0x66, 0x66, 0x66))

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(meta_p, 'CAB Date: Thursday, March 5, 2026 — 1:00 PM PST\n', font_size=12)
add_run(meta_p, '14 Change Requests Under Review\n', font_size=12)
add_run(meta_p, 'Prepared by: EA / Change Management\n', font_size=11, color=RGBColor(0x66, 0x66, 0x66))
add_run(meta_p, 'Date Prepared: March 4, 2026', font_size=11, color=RGBColor(0x66, 0x66, 0x66))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# FRAMEWORK REFERENCE
# ════════════════════════════════════════════════════════════════
doc.add_heading('Assessment Frameworks', level=1)

doc.add_heading('ITIL 4 Change Enablement', level=2)
p = doc.add_paragraph()
add_run(p, 'This review applies the ITIL 4 Change Enablement practice as defined in the ITIL 4 Foundation and Practice Guides (Axelos/PeopleCert). Key assessment dimensions include:', font_size=10)

itil_items = [
    ("Seven Rs of Change", "Who Raised it? What is the Reason? What Return is expected? What are the Risks? What Resources are required? Who is Responsible? What is the Relationship to other changes?"),
    ("Change Type Classification", "Normal changes require full risk assessment and CAB authorization. Emergency changes are time-sensitive with expedited authorization and mandatory post-implementation review (PIR). Standard changes are pre-approved, low-risk, repeatable."),
    ("Risk Assessment", "Every change must have risks identified, categorized, and accepted by the appropriate change authority."),
    ("Change Schedule", "Changes must be scheduled to minimize conflict with other changes and business operations."),
    ("Authorization Model", "Normal changes authorized by CAB/CCB. Emergency changes authorized by delegated authority with retroactive CAB review. Standard changes pre-authorized."),
    ("Post-Implementation Review (PIR)", "Required for all emergency changes; recommended for normal changes to confirm objectives met."),
    ("Continual Improvement", "Change models should be reviewed for optimization; patterns in failures/flags should drive process improvement."),
]

for title_text, body_text in itil_items:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_run(p, f"{title_text}: ", font_size=10)
    add_run(p, body_text, font_size=10)

doc.add_paragraph()
doc.add_heading('ISMS-STA-11.01-01 Enterprise Change Management Program', level=2)
p = doc.add_paragraph()
add_run(p, 'This review simultaneously validates compliance against Vituity\'s internal policy. Key requirements extracted from the policy:', font_size=10)

isms_items = [
    ("Minimum CR Fields (Section: Provisions)", "Change Request Type, Requested By, Category, Configuration Item, Environment, Assignment Group, Short Description, Description, Justification/Business Value, Implementation Plan, Risk and Impact Analysis, Backout Plan, Test Plan, Planned Date and Time of Change."),
    ("CCB Submission Deadline", "Change requests must be submitted and at Assess state by Wednesday 3:00 PM PT for Thursday CCB review."),
    ("Business Owner Approval", "Requester must obtain Business Owner approval prior to submitting to CCB."),
    ("Pre-Implementation Testing", "Changes to applications and operating systems are tested for security, usability, and impact prior to promoting to production. Testing in environment segregated from production where possible."),
    ("Backout Plan Requirement", "Detailed backout plan required; if implementation experiences issues, Implementor must initiate roll-back procedures to return resource to pre-change state."),
    ("Version Control", "Implementor maintains version control for all software updates."),
    ("Security Review", "InfoSec approval is required on all changes. CCB ensures changes do not compromise security."),
    ("Implementation Authorization", "All changes listed as Yes for 'CCB Approval Required' must be approved before implementation."),
    ("Post-Change Validation", "Requestor confirms and validates work is available and working as designed after production promotion."),
    ("Unauthorized Changes", "Changes implemented without full approval, outside change windows, or causing incidents are logged as unauthorized."),
    ("Communication", "Business is notified in advance of maintenance impacting service availability."),
    ("Documentation", "System documentation, operating documentation, and user procedures updated as needed."),
]

for title_text, body_text in isms_items:
    p = doc.add_paragraph(style='List Bullet')
    add_bold_run(p, f"{title_text}: ", font_size=10)
    add_run(p, body_text, font_size=10)

p = doc.add_paragraph()
add_run(p, '\nReferences: HITRUST CSF v9.2 (Change Management), ISMS Manual, Risk Management Program [ISMS-STA-05.02-01], Continuous Monitoring Program [ISMS-STA-11.09-01].', font_size=9, color=RGBColor(0x66, 0x66, 0x66))

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
doc.add_heading('Executive Summary', level=1)

summary_data = [
    ("PASS", "9", "CHG0039423, CHG0039418, CHG0039415, CHG0039403, CHG0039383, CHG0039339, CHG0039284, CHG0039414, CHG0039213"),
    ("PASS w/ minor flags", "3", "CHG0039420, CHG0039399, CHG0039283"),
    ("FLAG - needs CCB discussion", "2", "CHG0039412, CHG0039386"),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
set_cell_text(hdr[0], "Status", bold=True, font_size=10)
set_cell_text(hdr[1], "Count", bold=True, font_size=10)
set_cell_text(hdr[2], "Changes", bold=True, font_size=10)
set_cell_shading(hdr[0], "D6E4F0")
set_cell_shading(hdr[1], "D6E4F0")
set_cell_shading(hdr[2], "D6E4F0")

for status, count, changes in summary_data:
    row = tbl.add_row().cells
    add_status_badge(row[0], status)
    set_cell_text(row[1], count, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    row[2].text = ""
    p = row[2].paragraphs[0]
    chg_list = [c.strip() for c in changes.split(",")]
    for i, chg in enumerate(chg_list):
        if i > 0:
            p.add_run(", ").font.size = Pt(9)
        add_hyperlink(p, chg, SNOW_URL + chg)

doc.add_paragraph()

# Key Items
doc.add_heading('Key Items for CCB Attention', level=2)

attention_items = [
    ("CHG0039412 - ORDC Hx Firmware Upgrade",
     "ITIL 4 Risk: Backout plan explicitly states HXDP/ESXi downgrade is \"generally not supported\" and risks data loss. ITIL 4 requires identified risks to be accepted by the change authority. "
     "ISMS Gap: ESXi target version incomplete - policy requires version control for all software updates (Section: Provisions). "
     "Recommendation: CCB should formally acknowledge limited reversibility risk and confirm Cisco TAC availability. Implementor should complete the ESXi target version field."),
    ("CHG0039386 - trust-manager Emergency (Retroactive)",
     "ITIL 4: Emergency changes require expedited authorization AND mandatory Post-Implementation Review (PIR). This change was implemented 2/23 and is now at Authorize state for retroactive CCB review. "
     "ISMS: Policy states \"full approval not obtained prior to implementation\" triggers the Unauthorized Change process (Section 7: Operating Procedures). However, the Change Manager may approve emergency changes and schedule them for CCB post-mortem review. "
     "Recommendation: CCB should confirm: (1) Was the emergency justified? (2) Was PIR completed? (3) Did the Change Manager authorize the emergency per policy?"),
    ("Recurring Pattern: Weak Pre-Implementation Testing",
     "ITIL 4: Risk assessment should consider testing adequacy. "
     "ISMS: Policy requires changes \"tested for security, usability, and impact prior to promoting to production.\" "
     "Three changes (CHG0039420, CHG0039399, CHG0039283) cite \"Not applicable\" or \"None\" for pre-implementation testing. While their risk levels may justify this, the pattern should be noted for continual improvement per ITIL 4."),
]

for title_text, body_text in attention_items:
    p = doc.add_paragraph()
    add_bold_run(p, title_text + "\n", font_size=10)
    add_run(p, body_text, font_size=10)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# INDIVIDUAL CHANGE REVIEWS
# ════════════════════════════════════════════════════════════════
doc.add_heading('Individual Change Reviews', level=1)

# Each change has ITIL 4 criteria AND ISMS criteria
changes = [
    {
        "number": "CHG0039423",
        "title": "Expose Titan SFTP To Internet",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Daniel Anderson", "group": "Enterprise Networking",
        "schedule": "3/5/2026 2:00-3:00 PM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Daniel Anderson (Requester & Implementor)"),
            ("Reason", "PASS", "CornerStone SFTP replacement, continuous improvement"),
            ("Return", "PASS", "Product replacement enabling broader SFTP access"),
            ("Risks", "PASS", "R&I completed; PHI exposure acknowledged; low risk"),
            ("Resources", "PASS", "Single implementor, firewall access confirmed"),
            ("Responsible", "PASS", "Daniel Anderson - build, test, implement"),
            ("Relationship", "PASS", "No conflicting changes noted in schedule window"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All 14 required fields populated"),
            ("CCB Submission", "PASS", "Created 3/4, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "Discrete steps with specific device pdx-edgefw01 (10.100.25.103)"),
            ("Backout Plan", "PASS", "Re-add IP restrictions; rollback trigger defined"),
            ("Test Plan", "PASS", "Prior testing done; post-impl systems validation"),
            ("Risk & Impact Analysis", "PASS", "All questionnaire fields answered"),
            ("Version Control", "N/A", "Firewall rule change, no software version"),
            ("Schedule After CCB", "PASS", "3/5 2:00 PM - after 1:00 PM CCB"),
            ("Security Review", "PASS", "PHI/PII impact acknowledged; prior SAR noted"),
        ],
    },
    {
        "number": "CHG0039420",
        "title": "Deploy Jump Server for EDI",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Michael Castro", "group": "Enterprise Systems",
        "schedule": "3/5/2026 2:00-2:30 PM PST",
        "overall": "PASS w/ minor flags",
        "seven_rs": [
            ("Raised", "PASS", "Michael Castro"),
            ("Reason", "PASS", "DR testing revealed Illumio blocks preventing EDI RDP access"),
            ("Return", "PASS", "EDI team can perform required tasks without removing security rules"),
            ("Risks", "PASS", "Low - new VM, no impact to existing systems"),
            ("Resources", "PASS", "Single implementor, VM template, RDS licensing"),
            ("Responsible", "PASS", "Michael Castro - build, test, implement"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 3/4, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "5 discrete steps: deploy VM, networking, domain join, RDS, access grants"),
            ("Backout Plan", "PASS", "Delete VM"),
            ("Test Plan", "FLAG", "Pre-impl: 'Not applicable' - ISMS requires testing prior to production. Should validate template boots and domain join."),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "N/A", "New server deployment"),
            ("Schedule After CCB", "PASS", "3/5 2:00 PM - after CCB"),
            ("Security Review", "PASS", "No PHI/PII; no SAR needed"),
            ("Server/Device", "PASS", "RCM-SRVJUMPEDI identified with DHCP"),
        ],
    },
    {
        "number": "CHG0039418",
        "title": "Abnormal - Remove Domains from Safelist",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Bill Carter", "group": "Infosec",
        "schedule": "3/10-3/27/2026 (rolling)",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Bill Carter (Infosec)"),
            ("Reason", "PASS", "Compromised sender inbox detection improvement"),
            ("Return", "PASS", "Enhanced email security - detect compromised trusted senders"),
            ("Risks", "PASS", "False positives identified and mitigated with rolling approach"),
            ("Resources", "PASS", "Single implementor, Abnormal platform access"),
            ("Responsible", "PASS", "Bill Carter - implement and monitor"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 3/4, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "7 clear steps including monitoring and remediation cycle"),
            ("Backout Plan", "PASS", "Re-add domains; clear trigger defined"),
            ("Test Plan", "PASS", "Tested with 5 domains, no issues; ongoing FP monitoring"),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "N/A", "Configuration change"),
            ("Schedule After CCB", "PASS", "Starts 3/10 - after CCB"),
            ("Security Review", "PASS", "Infosec-owned change; SAR completed"),
            ("Communication", "PASS", "Rolling schedule attached to CR for visibility"),
        ],
    },
    {
        "number": "CHG0039415",
        "title": "Workato Recipe: Out-of-Sync Device Notification",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Pallav Malu", "group": "Enterprise Applications",
        "schedule": "3/5/2026 8:00-10:00 PM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "David Chu (Requester), Pallav Malu (Implementor)"),
            ("Reason", "PASS", "Intune compliance notification automation"),
            ("Return", "PASS", "Users remediate out-of-sync devices proactively"),
            ("Risks", "PASS", "Low - new recipes, no modification to existing systems"),
            ("Resources", "PASS", "Workato PROD access, Pallav Malu"),
            ("Responsible", "PASS", "Pallav Malu - all steps"),
            ("Relationship", "PASS", "Related to CHG0037461 (prior change referenced in title)"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 3/4, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "4 steps with recipe names and documentation links"),
            ("Backout Plan", "PASS", "Stop recipes in Workato"),
            ("Test Plan", "PASS", "UAT performed, data validated by requestor"),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "PASS", "New recipes identified (005, 006)"),
            ("Schedule After CCB", "PASS", "3/5 8:00 PM - after CCB"),
            ("Security Review", "PASS", "No PHI/PII; vendor support available"),
            ("Documentation", "PASS", "SharePoint recipe docs linked"),
        ],
    },
    {
        "number": "CHG0039414",
        "title": "Update Okta RADIUS Agent",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Jeffrey How", "group": "Infosec",
        "schedule": "3/6/2026 10:00-11:00 AM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Jeffrey How (Infosec)"),
            ("Reason", "PASS", "Log4j security vulnerability remediation"),
            ("Return", "PASS", "Eliminate known vulnerability in RADIUS agent"),
            ("Risks", "PASS", "Low - no production services use RADIUS currently"),
            ("Resources", "PASS", "Jeffrey How, RDP access to both servers"),
            ("Responsible", "PASS", "Jeffrey How - sequential upgrade on both servers"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 3/4, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "4 steps, sequential per-server upgrade"),
            ("Backout Plan", "PASS", "Reinstall v2.24.2 or restore snapshot"),
            ("Test Plan", "PASS", "Historical success; post-impl RADIUS service verification"),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "PASS", "v2.24.2 -> v2.26.0 documented"),
            ("Schedule After CCB", "PASS", "3/6 10:00 AM - after CCB"),
            ("Security Review", "PASS", "Infosec-owned; vulnerability remediation"),
            ("Server/Device", "PASS", "rcm-srvokta, ordc-srvokta identified"),
        ],
    },
    {
        "number": "CHG0039412",
        "title": "Upgrade Firmware on ORDC Hx",
        "type": "Normal", "risk": "Moderate", "impact": "2 - Medium",
        "assignee": "Philip Weiss", "group": "Enterprise Systems",
        "schedule": "3/11/2026 6:00-11:59 PM PST",
        "overall": "FLAG - needs CCB discussion",
        "seven_rs": [
            ("Raised", "PASS", "Philip Weiss"),
            ("Reason", "PASS", "Multiple vulnerability remediation (Cisco Hx firmware + VMware ESXi)"),
            ("Return", "PASS", "Patched infrastructure, reduced vulnerability exposure"),
            ("Risks", "FLAG", "HXDP/ESXi downgrade 'generally not supported' - data loss risk acknowledged but not fully mitigated"),
            ("Resources", "PASS", "Philip Weiss + Cisco TAC on standby"),
            ("Responsible", "PASS", "Philip Weiss - all phases"),
            ("Relationship", "PASS", "Multi-phase (vCenter already done); ESXi depends on firmware"),
        ],
        "isms": [
            ("Minimum CR Fields", "FLAG", "ESXi target version incomplete ('7.0.3 -> ???')"),
            ("CCB Submission", "PASS", "Created 3/4, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "4 phased steps; Step 1 (vCenter 8.03) already completed"),
            ("Backout Plan", "FLAG", "Explicitly states downgrade 'generally not supported' and can cause data loss. ISMS requires 'detailed backout plan' and 'roll-back procedures to return resource to pre-change state.'"),
            ("Test Plan", "FLAG", "Pre-impl: 'done in the past' only. ISMS: 'tested for security, usability, and impact prior to production.' QA: 'Not Possible: single Hx' is honest but risk."),
            ("Risk & Impact Analysis", "PASS", "All fields answered; business-critical: Yes"),
            ("Version Control", "FLAG", "Hx firmware versions complete (4.0(4k)->4.2(3o)), HxDP complete (4.5(2e)->5.5(2b)), but ESXi target missing"),
            ("Schedule After CCB", "PASS", "3/11 6:00 PM - after CCB"),
            ("Security Review", "PASS", "Vulnerability remediation; vendor support confirmed"),
            ("Server/Device", "PASS", "12 ESXi hosts + 2 fabric interconnects enumerated"),
            ("Communication", "PASS", "Impact to PM, RCM noted; after-hours window"),
        ],
    },
    {
        "number": "CHG0039403",
        "title": "Assign ASR Policies to PP Users (Warning Mode)",
        "type": "Normal", "risk": "Moderate", "impact": "2 - Medium",
        "assignee": "Chintan Myakal", "group": "Infosec",
        "schedule": "3/6/2026 7:00-10:00 AM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Chintan Myakal (Infosec)"),
            ("Reason", "PASS", "Security posture enhancement - ASR for PP department"),
            ("Return", "PASS", "Enhanced endpoint security, detect suspicious activity"),
            ("Risks", "PASS", "Warning mode only - no blocking; tested with exclusions"),
            ("Resources", "PASS", "Intune access, Chintan Myakal"),
            ("Responsible", "PASS", "Chintan Myakal - all Intune configuration"),
            ("Relationship", "PASS", "Follows previous ASR rollouts to other departments"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 3/2, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "9 detailed Intune steps with specific group names"),
            ("Backout Plan", "PASS", "Switch from Warn to Audit mode"),
            ("Test Plan", "PASS", "Tested with PP test group; exclusions validated; no user impact observed"),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "N/A", "Policy configuration, not software update"),
            ("Schedule After CCB", "PASS", "3/6 7:00 AM - after CCB"),
            ("Security Review", "PASS", "Infosec-owned change"),
        ],
    },
    {
        "number": "CHG0039399",
        "title": "Apply Security Patches to F5 Load Balancers",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Philip Weiss", "group": "Enterprise Systems",
        "schedule": "3/6/2026 6:00-8:00 PM PST",
        "overall": "PASS w/ minor flags",
        "seven_rs": [
            ("Raised", "PASS", "Philip Weiss"),
            ("Reason", "PASS", "Two specific F5 vulnerability advisories (K000156644, K000156643)"),
            ("Return", "PASS", "Remediate known vulnerabilities on production load balancers"),
            ("Risks", "PASS", "HA failover mitigates traffic interruption risk"),
            ("Resources", "PASS", "Philip Weiss, F5 vendor support available"),
            ("Responsible", "PASS", "Philip Weiss - rolling patch with failover"),
            ("Relationship", "PASS", "No conflicting changes in window"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 3/2, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "Rolling patch: standby first, failover, verify, repeat"),
            ("Backout Plan", "PASS", "Fail traffic to unpatched unit"),
            ("Test Plan", "FLAG", "Pre-impl: 'None' - ISMS requires testing prior to production. Historical experience cited but no specific validation."),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "PASS", "BIGIP-17.5.1-0.0.7 -> 17.5.1.4-0.0.20"),
            ("Schedule After CCB", "PASS", "3/6 6:00 PM - after CCB"),
            ("Security Review", "PASS", "Vulnerability remediation; vendor support during change"),
            ("Server/Device", "PASS", "ORDC-LB01, ORDC-LB02 identified"),
        ],
    },
    {
        "number": "CHG0039386",
        "title": "Upgrade trust-manager to v0.21.0",
        "type": "Emergency", "risk": "Low", "impact": "3 - Low",
        "assignee": "Stefan Nuxoll", "group": "Development - RCM",
        "schedule": "2/23/2026 10:00 AM-12:00 PM (ALREADY IMPLEMENTED)",
        "overall": "FLAG - needs CCB discussion",
        "seven_rs": [
            ("Raised", "PASS", "Stefan Nuxoll"),
            ("Reason", "PASS", "PDRS/Athena IDX CA certificate issue causing service disruption"),
            ("Return", "PASS", "Restored PDRS functionality with updated CA bundle"),
            ("Risks", "PASS", "Low - chart version bump only"),
            ("Resources", "PASS", "Stefan Nuxoll, AKS cluster access"),
            ("Responsible", "PASS", "Stefan Nuxoll"),
            ("Relationship", "PASS", "No related changes noted"),
        ],
        "isms": [
            ("Minimum CR Fields", "FLAG", "Implementation plan is single step only ('bump chart version')"),
            ("CCB Submission", "FLAG", "Emergency change - implemented 2/23 before CCB review. Policy: Change Manager may approve emergencies and schedule for CCB post-mortem (Section: Change Manager role)."),
            ("Business Owner Approval", "FLAG", "Approval state: 'Requested' - should confirm Change Manager authorized per emergency process"),
            ("Implementation Plan", "FLAG", "Single step - minimal detail for an emergency change record"),
            ("Backout Plan", "PASS", "Revert to v0.19.0"),
            ("Test Plan", "PASS", "Verify cert bundles + PDRS functionality"),
            ("Risk & Impact Analysis", "PASS", "All fields answered; QA: 'N/A - applied to all environments immediately'"),
            ("Version Control", "PASS", "v0.16.1 -> v0.21.0 (note: backout references v0.19.0, not v0.16.1)"),
            ("Schedule After CCB", "FLAG", "RETROACTIVE - already implemented 2/23. Per ISMS, this triggers Section 7 Unauthorized Change review unless Change Manager authorized the emergency."),
            ("Post-Implementation Review", "FLAG", "ITIL 4 REQUIRES PIR for all emergency changes. CCB should confirm PIR was conducted."),
            ("Security Review", "PASS", "PHI/PII: Yes - appropriately flagged"),
        ],
    },
    {
        "number": "CHG0039383",
        "title": "Vituity Stats SharePoint Page",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "DivyaRani Bhat", "group": "Enterprise Applications",
        "schedule": "3/5/2026 3:00-5:00 PM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Amy Hughes (Requester), DivyaRani Bhat (Implementor)"),
            ("Reason", "PASS", "Centralized stats page sponsored by executive leadership (Mu)"),
            ("Return", "PASS", "One-stop-shop for Vituity stats, auto-updated"),
            ("Risks", "PASS", "Read-only pipeline; no source system modification"),
            ("Resources", "PASS", "Azure SQL, Workato, SharePoint - all existing infrastructure"),
            ("Responsible", "PASS", "DivyaRani Bhat - all implementation steps"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 2/25, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "Detailed: 4 stored procedures + Workato recipe + SharePoint CSVs"),
            ("Backout Plan", "PASS", "Disable Workato recipe; clear backout triggers defined"),
            ("Test Plan", "PASS", "SP validation, Workato testing, data validation, post-impl monitoring"),
            ("Risk & Impact Analysis", "PASS", "Thoroughly answered; no PHI/PII"),
            ("Version Control", "N/A", "New pipeline, not an upgrade"),
            ("Schedule After CCB", "PASS", "3/5 3:00 PM - after CCB"),
            ("Security Review", "PASS", "No new integrations or auth models; existing approved infrastructure"),
            ("Documentation", "PASS", "Recipe documentation and Workato links provided"),
        ],
    },
    {
        "number": "CHG0039339",
        "title": "ServiceNow to Otto Integration - Agent Alerts",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Dan Spengler", "group": "Service Delivery Optimization",
        "schedule": "3/5/2026 3:00-4:00 PM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Dan Spengler"),
            ("Reason", "PASS", "Agents miss ticket assignments; email alerts lost in noise"),
            ("Return", "PASS", "Faster ticket awareness via alternate channel (Otto)"),
            ("Risks", "PASS", "Low - additive integration; opt-out field provided"),
            ("Resources", "PASS", "Dan Spengler + Beth Vanderheiden (update set promotion)"),
            ("Responsible", "PASS", "Dan Spengler (Moveworks), Beth Vanderheiden (ServiceNow)"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 2/17, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "Exceptionally detailed: update set, 2 flows, sys_user field, OAuth integration, Moveworks listener/plugin - all with dev links"),
            ("Backout Plan", "PASS", "Disable flows + disconnect listener; update set backout documented"),
            ("Test Plan", "PASS", "Tested in dev/sandbox with multiple agents"),
            ("Risk & Impact Analysis", "PASS", "SAR completed"),
            ("Version Control", "N/A", "New integration"),
            ("Schedule After CCB", "PASS", "3/5 3:00 PM - after CCB"),
            ("Security Review", "PASS", "SAR completed; OAuth 2.0 authentication documented"),
            ("Documentation", "PASS", "All components linked to dev instance"),
        ],
    },
    {
        "number": "CHG0039284",
        "title": "Enable Otto to Triage to HR and Facilities",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Beth Vanderheiden", "group": "Service Delivery Optimization",
        "schedule": "3/10/2026 9:00-10:00 AM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Mark White (Requester), Beth Vanderheiden (Implementor)"),
            ("Reason", "PASS", "Enable faster ticket routing to HR and Facilities"),
            ("Return", "PASS", "Tickets routed to correct teams faster"),
            ("Risks", "PASS", "Low - configuration change; excludes sensitive groups (LMS, L&D, HRIS Tier 3)"),
            ("Resources", "PASS", "Moveworks portal access"),
            ("Responsible", "PASS", "Beth Vanderheiden"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 2/5, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "4 clear steps in Moveworks portal"),
            ("Backout Plan", "PASS", "Re-add blocked groups"),
            ("Test Plan", "PASS", "Tested in dev; post-impl analytics monitoring"),
            ("Risk & Impact Analysis", "PASS", "All fields answered"),
            ("Version Control", "N/A", "Configuration change"),
            ("Schedule After CCB", "PASS", "3/10 9:00 AM - after CCB"),
            ("Security Review", "PASS", "Existing system already running; no SAR needed"),
        ],
    },
    {
        "number": "CHG0039283",
        "title": "Deactivate 3 ServiceNow Forms",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Beth Vanderheiden", "group": "Enterprise Applications",
        "schedule": "3/10/2026 9:00-10:00 AM PST",
        "overall": "PASS w/ minor flags",
        "seven_rs": [
            ("Raised", "PASS", "Mark White (Requester), Beth Vanderheiden (Implementor)"),
            ("Reason", "PASS", "Forms replaced by Otto and Staples; streamlining ServiceNow"),
            ("Return", "PASS", "Reduced user confusion, faster ticket processing for Service Desk"),
            ("Risks", "PASS", "Low - forms can be re-enabled if needed"),
            ("Resources", "PASS", "Update set already in dev, Beth Vanderheiden"),
            ("Responsible", "PASS", "Beth Vanderheiden"),
            ("Relationship", "PASS", "No conflicting changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 2/5, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "Update set promotion from dev to prod; forms identified with sys_ids"),
            ("Backout Plan", "PASS", "Back out update set"),
            ("Test Plan", "FLAG", "Pre-impl: 'not needed', QA: 'not needed'. ISMS: should at minimum verify update set previews cleanly before committing."),
            ("Risk & Impact Analysis", "FLAG", "'Documented prior notes': No - minor gap per ISMS Provisions"),
            ("Version Control", "N/A", "Deactivation, not upgrade"),
            ("Schedule After CCB", "PASS", "3/10 9:00 AM - after CCB"),
            ("Security Review", "PASS", "No security impact"),
        ],
    },
    {
        "number": "CHG0039213",
        "title": "New Global ServiceNow Survey",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Beth Vanderheiden", "group": "Enterprise Applications",
        "schedule": "3/5/2026 4:00-6:00 PM PST",
        "overall": "PASS",
        "seven_rs": [
            ("Raised", "PASS", "Mark White (Requester), Beth Vanderheiden (Implementor)"),
            ("Reason", "PASS", "Consolidate IT + HR surveys into single global survey"),
            ("Return", "PASS", "Unified feedback mechanism across Shared Services and Service Delivery"),
            ("Risks", "PASS", "Low - survey scope tested in dev"),
            ("Resources", "PASS", "ServiceNow dev/prod, Flow Designer, Otto integration"),
            ("Responsible", "PASS", "Beth Vanderheiden - all implementation"),
            ("Relationship", "PASS", "Disables 2 old surveys; no conflict with other changes"),
        ],
        "isms": [
            ("Minimum CR Fields", "PASS", "All required fields populated"),
            ("CCB Submission", "PASS", "Created 1/26, state: Assess"),
            ("Business Owner Approval", "PASS", "Approval: Requested"),
            ("Implementation Plan", "PASS", "Detailed: update set, survey config, triggers, email notifications, Flow Designer recipe - all with dev links"),
            ("Backout Plan", "PASS", "Back out update set"),
            ("Test Plan", "PASS", "Tested in dev with business; post-impl manual trigger + Otto reminder verification"),
            ("Risk & Impact Analysis", "PASS", "All fields answered; SAR completed"),
            ("Version Control", "N/A", "New survey configuration"),
            ("Schedule After CCB", "PASS", "3/5 4:00 PM - after CCB"),
            ("Security Review", "PASS", "No PHI/PII; SAR completed"),
            ("Documentation", "PASS", "Survey requirements linked, all components documented"),
        ],
    },
]

for chg in changes:
    # Change header
    h = doc.add_heading(level=2)
    add_hyperlink(h, chg["number"], SNOW_URL + chg["number"])
    h.add_run(f' \u2014 {chg["title"]}')

    # Meta info
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(2)
    for label, val in [("Type", chg["type"]), ("Risk", chg["risk"]), ("Impact", chg["impact"]),
                       ("Assignee", chg["assignee"]), ("Group", chg["group"])]:
        add_bold_run(meta, f"{label}: ", font_size=9)
        add_run(meta, f"{val}   ", font_size=9)
    meta2 = doc.add_paragraph()
    meta2.paragraph_format.space_after = Pt(4)
    add_bold_run(meta2, "Schedule: ", font_size=9)
    add_run(meta2, chg["schedule"], font_size=9)

    # Overall verdict
    verdict_p = doc.add_paragraph()
    add_bold_run(verdict_p, "Overall Verdict: ", font_size=11)
    vcolor = RGBColor(0x2E, 0x7D, 0x32)  # green
    if "FLAG" in chg["overall"] and "minor" not in chg["overall"]:
        vcolor = RGBColor(0xBF, 0x36, 0x0C)
    elif "minor" in chg["overall"]:
        vcolor = RGBColor(0x55, 0x8B, 0x2F)
    add_run(verdict_p, chg["overall"], font_size=11, color=vcolor, bold=True)

    # ITIL 4 Seven Rs Table
    doc.add_heading('ITIL 4 \u2014 Seven Rs of Change', level=3)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    set_cell_text(hdr[0], "Question", bold=True, font_size=9)
    set_cell_text(hdr[1], "Status", bold=True, font_size=9)
    set_cell_text(hdr[2], "Assessment", bold=True, font_size=9)
    set_cell_shading(hdr[0], "E8EAF6")
    set_cell_shading(hdr[1], "E8EAF6")
    set_cell_shading(hdr[2], "E8EAF6")

    for crit_name, crit_status, crit_notes in chg["seven_rs"]:
        row = tbl.add_row().cells
        set_cell_text(row[0], crit_name, font_size=9)
        add_status_badge(row[1], crit_status)
        set_cell_text(row[2], crit_notes, font_size=9)
        if crit_status == "FLAG":
            set_cell_shading(row[0], "FFF3E0")
            set_cell_shading(row[1], "FFF3E0")
            set_cell_shading(row[2], "FFF3E0")

    # ISMS Policy Compliance Table
    doc.add_heading('ISMS-STA-11.01-01 Policy Compliance', level=3)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Light Grid Accent 1'
    hdr2 = tbl2.rows[0].cells
    set_cell_text(hdr2[0], "Requirement", bold=True, font_size=9)
    set_cell_text(hdr2[1], "Status", bold=True, font_size=9)
    set_cell_text(hdr2[2], "Evidence / Notes", bold=True, font_size=9)
    set_cell_shading(hdr2[0], "E8F5E9")
    set_cell_shading(hdr2[1], "E8F5E9")
    set_cell_shading(hdr2[2], "E8F5E9")

    for crit_name, crit_status, crit_notes in chg["isms"]:
        row = tbl2.add_row().cells
        set_cell_text(row[0], crit_name, font_size=9)
        add_status_badge(row[1], crit_status)
        set_cell_text(row[2], crit_notes, font_size=9)
        if crit_status == "FLAG":
            set_cell_shading(row[0], "FFF3E0")
            set_cell_shading(row[1], "FFF3E0")
            set_cell_shading(row[2], "FFF3E0")

    doc.add_paragraph()
    if chg != changes[-1]:
        doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# APPENDIX
# ════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('Appendix: Assessment Methodology', level=1)

p = doc.add_paragraph()
add_run(p, 'Each change request was assessed against two frameworks simultaneously:\n\n', font_size=10)
add_bold_run(p, '1. ITIL 4 Change Enablement (Axelos/PeopleCert)\n', font_size=10)
add_run(p, '   - Seven Rs checklist applied to validate completeness of change rationale, risk identification, resource planning, and relationship mapping.\n', font_size=10)
add_run(p, '   - Change type classification validated (Normal/Emergency/Standard).\n', font_size=10)
add_run(p, '   - Emergency changes assessed for mandatory Post-Implementation Review (PIR).\n', font_size=10)
add_run(p, '   - Authorization model validated per change type.\n\n', font_size=10)
add_bold_run(p, '2. ISMS-STA-11.01-01 Enterprise Change Management Program (Vituity)\n', font_size=10)
add_run(p, '   - 14 minimum CR fields validated per Provisions section.\n', font_size=10)
add_run(p, '   - CCB submission deadline (Wednesday 3:00 PM PT) verified.\n', font_size=10)
add_run(p, '   - Schedule validated: Planned Start/End must fall after Thursday 1:00 PM PST CCB meeting.\n', font_size=10)
add_run(p, '   - Implementation, Backout, and Test Plans assessed for actionable, discrete steps.\n', font_size=10)
add_run(p, '   - Version control compliance checked for software updates.\n', font_size=10)
add_run(p, '   - Security review and InfoSec approval validated.\n', font_size=10)
add_run(p, '   - Unauthorized Change criteria checked for emergency/retroactive changes.\n', font_size=10)

p2 = doc.add_paragraph()
add_bold_run(p2, 'Sources:\n', font_size=10)
add_run(p2, '- ITIL 4 Foundation (Axelos, 2019) - Change Enablement Practice\n', font_size=9)
add_run(p2, '- ITIL 4 Change Enablement Practice Guide (PeopleCert)\n', font_size=9)
add_run(p2, '- Seven Rs of Change Management (ITIL best practice)\n', font_size=9)
add_run(p2, '- ISMS-STA-11.01-01 Enterprise Change Management Program (Vituity)\n', font_size=9)
add_run(p2, '- HITRUST CSF v9.2, Change Management\n', font_size=9)

# Save
output_dir = r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\CCB\2026"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "ITIL4-ISMS Compliance Review - CAB 2026-03-05.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
