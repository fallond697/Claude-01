#!/usr/bin/env python3
"""Generate a Word document summarizing PIR review for changes in Review state."""

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = Path.home() / "OneDrive - Vituity" / "Documents" / "Change Management"

# ---------------------------------------------------------------------------
# PIR data from 2026-02-12 review session
# ---------------------------------------------------------------------------

REVIEW_DATE = "2026-02-12"

SCORECARD = [
    {"number": "CHG0039282", "score": "11/14", "status": "GAPS",
     "short_desc": "Repair 2 domain controllers that have replication issues",
     "type": "Normal", "group": "Enterprise Systems", "assigned": "Melvin Mah"},
    {"number": "CHG0039278", "score": "14/14", "status": "PASS",
     "short_desc": "Add new table to AIDX extracts on DW server, Informatics server and VITTEST server",
     "type": "Standard", "group": "Data Ops - PM", "assigned": "Avinash Vedavyas Prabhu"},
    {"number": "CHG0039276", "score": "12/14", "status": "GAPS",
     "short_desc": "Upgrade VMware vCenter Server Appliance (7.0.3 -> 8.0.3)",
     "type": "Normal", "group": "Enterprise Systems", "assigned": "Philip Weiss"},
    {"number": "CHG0039257", "score": "13/14", "status": "GAPS",
     "short_desc": 'Create a "Pending Rescinded Providers" Salesforce Report',
     "type": "Normal", "group": "Enterprise Applications", "assigned": "Yanyan Meng"},
    {"number": "CHG0038896", "score": "14/14", "status": "PASS",
     "short_desc": "Modify ServiceNow automation from Workato to only open 1 RITM per Startup email",
     "type": "Normal", "group": "Enterprise Applications", "assigned": "Divya Meghana"},
]

NOTES = {
    "CHG0039282": {
        "summary": (
            "Implementation plan contained discrete steps (FSMO move, demote/promote cycle "
            "for CORP2 then MAW1, replication verification). Plan appears sound for the work performed."
        ),
        "gaps": [
            "Planned start/end dates not populated \u2014 unable to verify change was executed within an approved window.",
            "Configuration Item and Environment fields not populated.",
            'Backout plan stated "none \u2014 this will need to be fixed." While this may be technically accurate '
            "for a DC replication repair, policy requires a documented backout approach.",
        ],
        "questions": [
            "Was replication successfully restored on both ORDC-SRVDCMAW1 and ORDC-SRVDCCORP2?",
            "Were there any authentication disruptions during the demote/promote cycle?",
            "What dates/times was this change actually executed?",
        ],
        "recommendation": (
            "Requires implementer confirmation of successful completion and actual implementation "
            "dates before closure. Backout plan gap noted for process improvement tracking."
        ),
        "impl_status": "WEAK", "backout_status": "WEAK", "test_status": "WEAK",
    },
    "CHG0039278": {
        "summary": (
            "Standard change to add IV_BV_VISIT_NOTE_DW table to AIDX extract pipeline across DW, "
            "Informatics, and VITTEST servers. Implementation plan had 3 clear steps with named owners "
            "(Athena Support for steps 1-2, Avinash for step 3). Test plan included counts test and "
            "automated Weather testing check."
        ),
        "gaps": [
            "Planned start/end dates not populated.",
            "Configuration Item and Environment fields not populated.",
            "PHI/PII marked Yes but SAR marked as not completed. For a data extract table this should "
            "have been addressed pre-implementation.",
        ],
        "questions": [
            "Is the temp_visit_note table loading successfully on schedule?",
            "Did counts test and Weather testing check pass?",
            "Has the SAR determination been resolved given PHI/PII involvement?",
            "What dates was this implemented?",
        ],
        "recommendation": (
            "Requires confirmation that extract is running clean and SAR disposition documented before closure."
        ),
        "impl_status": "OK", "backout_status": "OK", "test_status": "OK",
    },
    "CHG0039276": {
        "summary": (
            "Major version upgrade of VMware VCSA on two servers: MBSI-SRVMGMT02 (10.10.4.176) and "
            "ORDC-SRVVCSA (10.100.27.20). Migration approach was spin-up-and-import with automatic "
            "IP/name takeover. Backout plan was clear (power off new, restore old)."
        ),
        "gaps": [
            "Planned start/end dates not populated.",
            "Configuration Item and Environment fields not populated.",
            'Pre-implementation testing was documented as "None" \u2014 for a business-critical '
            "infrastructure component with no HA, this is a notable gap.",
        ],
        "questions": [
            "Is the new vSphere 8.0.3 web interface accessible and stable on both servers?",
            "Were there any issues during the import/cutover process?",
            "Have all VM hosts reconnected and are clusters healthy?",
            "What dates/times was this executed, and was there any unplanned downtime?",
        ],
        "recommendation": (
            "Requires implementer verification of successful upgrade on both appliances and actual "
            "implementation dates before closure. Note pre-implementation testing gap for process improvement."
        ),
        "impl_status": "OK", "backout_status": "OK", "test_status": "MISSING",
    },
    "CHG0039257": {
        "summary": (
            "New automated report created via stored procedure (usp_SFDC_To_HCM_Rescinded_Discrepancy_Report) "
            "and SSIS package, scheduled via SQL Agent job to email HRPartnershipAffairs and PMCredAnalyst every "
            "Monday at 8am PST. Implementation plan had 4 discrete steps. Description was thorough with full "
            "acceptance criteria. Test plan included pre-validation with business stakeholder (Racquel Llavore)."
        ),
        "gaps": [
            "Planned start/end dates not populated.",
            "Configuration Item and Environment fields not populated.",
            'Backout decision criteria was vague ("If Business asks so") \u2014 should reference a technical failure condition.',
            "Change window did not include backout time.",
        ],
        "questions": [
            "Has the Monday 8am report fired successfully at least once?",
            "Did Racquel validate the report output and confirm acceptance criteria are met?",
            "Are providers correctly dropping off when HCM status is updated to Rescinded?",
            "What date was this deployed?",
        ],
        "recommendation": (
            "Requires confirmation that at least one successful scheduled execution has occurred "
            "and business validation is complete before closure."
        ),
        "impl_status": "WEAK", "backout_status": "OK", "test_status": "OK",
    },
    "CHG0038896": {
        "summary": (
            "Enhancement to Workato recipe (03|DB_SiteStartup_ServiceNow|Func - Create RITM for new site, "
            "v2 -> v3) to consolidate RITM creation to one per startup email instead of one per hospital. "
            "Implementation plan was well-documented with 5 steps, specific recipe names, and linked SharePoint "
            "documentation. Backout plan was clear with version restore path. Change open since 12/29/2025 (45 days)."
        ),
        "gaps": [
            "Planned start/end dates not populated.",
            "Configuration Item and Environment fields not populated.",
            "SAR question not answered in Risk & Impact section.",
        ],
        "questions": [
            "Has the updated recipe processed at least one startup email correctly, creating a single RITM?",
            "Is the update-if-exists logic working for subsequent hospitals in the same email?",
            "Were there any failed jobs since deployment?",
            "What was the actual implementation date?",
        ],
        "recommendation": (
            "Strongest submission of the batch. Requires implementer confirmation of successful "
            "operation and actual dates before closure."
        ),
        "impl_status": "OK", "backout_status": "OK", "test_status": "OK",
    },
}

CROSS_CUTTING = [
    ("Planned Start/End not populated", "5/5", "High",
     "No audit trail of approved change windows"),
    ("Configuration Item missing", "5/5", "Medium",
     "CMDB linkage broken for these changes"),
    ("Environment missing", "5/5", "Medium",
     "No environment classification recorded"),
]

# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def add_status_color(run, status):
    """Color a status run: green=OK, orange=WEAK, red=MISSING/FAIL."""
    if status == "OK":
        run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
    elif status == "WEAK":
        run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
    else:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True


def build_document() -> Document:
    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title = doc.add_heading("Post-Implementation Review Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Review Date: {REVIEW_DATE}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Policy Reference: ISMS-STA-11.01-01 Enterprise Change Management Program")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Scorecard ---
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        f"5 change requests in Review state were analyzed against the 14 policy-required fields "
        f"defined in ISMS-STA-11.01-01. 2 of 5 are fully compliant; 3 have documentation gaps "
        f"requiring attention before closure."
    )

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, label in enumerate(["Change", "Score", "Status", "Type", "Assignment Group"]):
        hdr[i].text = label
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for item in SCORECARD:
        row = table.add_row().cells
        row[0].text = item["number"]
        row[1].text = item["score"]
        row[2].text = item["status"]
        row[3].text = item["type"]
        row[4].text = item["group"]
        # Color the status cell
        for paragraph in row[2].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if item["status"] == "PASS":
                    run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
                else:
                    run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)

    doc.add_paragraph()  # spacer

    # --- Individual change reviews ---
    doc.add_heading("Individual Change Reviews", level=1)

    for item in SCORECARD:
        chg = item["number"]
        note = NOTES[chg]

        doc.add_heading(f'{chg} \u2014 {item["short_desc"]}', level=2)

        # Metadata line
        meta = doc.add_paragraph()
        meta.add_run("Assigned to: ").bold = True
        meta.add_run(f'{item["assigned"]}  |  ')
        meta.add_run("Group: ").bold = True
        meta.add_run(f'{item["group"]}  |  ')
        meta.add_run("Type: ").bold = True
        meta.add_run(item["type"])

        # Plan status line
        plan_p = doc.add_paragraph()
        plan_p.add_run("Plan Assessments:  ").bold = True
        plan_p.add_run("Implementation: ")
        r = plan_p.add_run(note["impl_status"])
        add_status_color(r, note["impl_status"])
        plan_p.add_run("  |  Backout: ")
        r = plan_p.add_run(note["backout_status"])
        add_status_color(r, note["backout_status"])
        plan_p.add_run("  |  Test: ")
        r = plan_p.add_run(note["test_status"])
        add_status_color(r, note["test_status"])

        # Summary
        doc.add_heading("Review Summary", level=3)
        doc.add_paragraph(note["summary"])

        # Gaps
        doc.add_heading("Documentation Gaps", level=3)
        for gap in note["gaps"]:
            doc.add_paragraph(gap, style="List Bullet")

        # Questions
        doc.add_heading("PIR Questions for Implementer", level=3)
        for i, q in enumerate(note["questions"], 1):
            doc.add_paragraph(f"{q}", style="List Number")

        # Recommendation
        doc.add_heading("Recommendation", level=3)
        rec_p = doc.add_paragraph()
        rec_p.add_run(note["recommendation"]).italic = True

        doc.add_page_break()

    # --- Cross-cutting observations ---
    doc.add_heading("Cross-Cutting Observations", level=1)
    doc.add_paragraph(
        "The following documentation gaps were observed across all 5 changes and indicate "
        "systemic process issues rather than individual oversights."
    )

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Light Grid Accent 1"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    for i, label in enumerate(["Gap", "Frequency", "Severity", "Impact"]):
        hdr2[i].text = label
        for paragraph in hdr2[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for gap, freq, sev, impact in CROSS_CUTTING:
        row = table2.add_row().cells
        row[0].text = gap
        row[1].text = freq
        row[2].text = sev
        row[3].text = impact

    doc.add_paragraph()

    # Process improvement recommendation
    doc.add_heading("Process Improvement Recommendation", level=2)
    doc.add_paragraph(
        "These three fields (Planned Start/End, Configuration Item, Environment) should be "
        "enforced as mandatory before a change can transition out of Assess state. This would "
        "prevent the recurring pattern of changes reaching Review without basic scheduling and "
        "CMDB linkage documented."
    )

    return doc


def main():
    doc = build_document()

    filename = f"PIR Review Summary {REVIEW_DATE}.docx"

    # Try OneDrive Change Management folder first, fall back to current dir
    if OUTPUT_DIR.exists():
        out_path = OUTPUT_DIR / filename
    else:
        out_path = Path.cwd() / filename

    doc.save(str(out_path))
    print(f"Document saved to: {out_path}")


if __name__ == "__main__":
    main()
