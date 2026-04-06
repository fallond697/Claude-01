#!/usr/bin/env python3
"""Generate PIR Review Summary for Enterprise Applications changes in Review state — 2026-04-03."""

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path.home() / "OneDrive - Vituity" / "Documents" / "Change Management"
SNOW_URL = "https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number="

REVIEW_DATE = "2026-04-03"

# ---------------------------------------------------------------------------
# Hyperlink helper
# ---------------------------------------------------------------------------

def add_hyperlink(paragraph, text, url):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# ---------------------------------------------------------------------------
# PIR data — 2026-04-03 review of Enterprise Applications changes
# ---------------------------------------------------------------------------

SCORECARD = [
    {
        "number": "CHG0039589",
        "score": "14/14",
        "status": "PASS",
        "short_desc": "Retirement Statement Data Update: 04/01/2026",
        "type": "Standard",
        "group": "Enterprise Applications",
        "assigned": "Xiu Lu",
        "requested_by": "Gabriela Ruiz",
        "ci": "Portal \u2013 Content Update",
        "planned_start": "04/02/2026 09:00",
        "planned_end": "04/02/2026 11:00",
        "release": "RLSE0012936",
    },
    {
        "number": "CHG0039513",
        "score": "14/14",
        "status": "PASS",
        "short_desc": "PE-16: Deploy Automated IEC/GEC Historical Status Update on Contract Termination",
        "type": "Normal",
        "group": "Enterprise Applications",
        "assigned": "Paramasivan Arunachalam",
        "requested_by": "Nancy Nair",
        "ci": "Salesforce",
        "planned_start": "04/02/2026 08:00",
        "planned_end": "04/02/2026 11:00",
        "release": "RLSE0012936",
    },
    {
        "number": "CHG0039333",
        "score": "13/14",
        "status": "GAPS",
        "short_desc": "CPAC Report Update \u2013 Special Agreement Visibility",
        "type": "Normal",
        "group": "Enterprise Applications",
        "assigned": "Yanyan Meng",
        "requested_by": "Tara Bernstein",
        "ci": "Office 365 - SharePoint",
        "planned_start": "03/31/2026 19:30",
        "planned_end": "03/31/2026 20:30",
        "release": "RLSE0012938",
    },
    {
        "number": "CHG0039131",
        "score": "13/14",
        "status": "GAPS",
        "short_desc": "Memo Field Population on Depreciation Journals",
        "type": "Normal",
        "group": "Enterprise Applications",
        "assigned": "Jitin Xavier (Contractor)",
        "requested_by": "Johnny Chan",
        "ci": "Netsuite",
        "planned_start": "03/26/2026 03:00",
        "planned_end": "03/26/2026 05:00",
        "release": "RLSE0012902",
    },
]

NOTES = {
    "CHG0039589": {
        "summary": (
            "Recurring Standard change for loading retirement contribution statement data to the Portal. "
            "Xiu Lu executed SQL scripts to backup, truncate, and reload data from Excel into the "
            "WebBenefitsContributions table on sacdc-svrsql2. Pre-validated by Gabriela Ruiz on qa.cep.com. "
            "Brian Ouderkirk approved and assigned to release RLSE0012936 for the 04/02 8\u201311am EA window."
        ),
        "activity": [
            ("04/01 10:28", "Gabriela Ruiz", "Created CR with attached data file (Excel, 470KB)"),
            ("04/01 14:55", "Gabriela Ruiz", "QA validation: \u201cUpdates look good!\u201d"),
            ("04/02 06:21", "Brian Ouderkirk", "Approved, assigned to RLSE0012936 for 04/02 8\u201311am"),
            ("04/02 10:17", "Xiu Lu", "Deployment to prod complete, requested requester verification"),
            ("04/02 12:38", "Gabriela Ruiz", "Production validation: \u201cData looks good.\u201d"),
        ],
        "gaps": [],
        "observations": [
            "Standard Change Matrix: Retirement contribution data loading is not explicitly listed in the "
            "Pre-Approved Matrix. The closest entries are Portal > 401K data loading / Profit Distribution / "
            "K1 data loading. Recommend adding \u201cRetirement Contribution Statement data loading\u201d to "
            "the Portal section of the matrix.",
        ],
        "recommendation": (
            "All required fields present. Implementation successful and validated by requester. "
            "Ready for closure."
        ),
        "closure_ready": True,
        "impl_status": "OK", "backout_status": "OK", "test_status": "OK",
    },
    "CHG0039513": {
        "summary": (
            "Deploy PE-16 (PROD-10504) Contract Auto-Cancel automation to Salesforce Production. "
            "Record-triggered flow that updates IEC/GEC records to \u201cHistorical ECHO/Incomplete Record\u201d "
            "when a Contract End Date is set. Additional logic within existing Apex class \u2014 not a new automation. "
            "Tested in SB2 by Hunter Dix, Siva, and UAT by Nancy Nair. Brian Ouderkirk approved and assigned "
            "to RLSE0012936."
        ),
        "activity": [
            ("03/19 14:56", "Hunter Dix", "Uploaded Nancy Nair approval screenshot"),
            ("03/25 12:24", "Hunter Dix", "Uploaded Yanyan Meng approval screenshot"),
            ("03/25 12:27", "Hunter Dix", "Work note: Yanyan confirmed no mainload impact from Apex class changes"),
            ("03/25 12:43", "Hunter Dix", "Uploaded Brian Ouderkirk approval screenshot"),
            ("03/31 09:29", "Brian Ouderkirk", "Approved, assigned to RLSE0012936 for 04/02 8\u201311am"),
            ("04/01 14:51", "Hunter Dix", "Uploaded Bala/Workato approval screenshot"),
            ("04/02 08:04", "Paramasivan Arunachalam", "Confirmed deployment validated, proceeding with deploy"),
            ("04/02 08:34", "Paramasivan Arunachalam", "Deployed to production, requested Nancy Nair validation"),
        ],
        "gaps": [],
        "observations": [
            "Post-implementation test is limited to checking field visibility on layouts. Does not verify "
            "the core automation logic (contract termination triggering IEC/GEC status flips).",
            "Awaiting Nancy Nair\u2019s production validation as of review date.",
        ],
        "recommendation": (
            "All required fields present. Deployed successfully. Pending Nancy Nair\u2019s confirmation "
            "that the automation is firing correctly on contract terminations. Ready for closure once "
            "validation is received."
        ),
        "closure_ready": False,
        "impl_status": "OK", "backout_status": "OK", "test_status": "OK",
    },
    "CHG0039333": {
        "summary": (
            "Updates to CPAC reports (Raw and Calc) to remove legacy 2018\u20132019 columns, add Pre2019 "
            "summary column, and add Special Agreement visibility (AA and Incentive columns) pulling from "
            "Salesforce. Implementation involved 8 steps: table modifications, 4 stored procedure updates, "
            "and SSIS package changes on PSQL1. Deployed during 03/31 7\u201310pm EA window per RLSE0012938."
        ),
        "activity": [
            ("02/13 15:26", "Tara Bernstein", "Uploaded requirement screenshots (CPAC Calc, PAD views)"),
            ("03/31 09:23", "Yanyan Meng", "Work note requesting Brian\u2019s approval for tonight\u2019s deploy"),
            ("03/31 09:39", "Brian Ouderkirk", "Approved, assigned to RLSE0012938 for 03/31 7\u201310pm"),
            ("03/31 09:21", "Yanyan Meng", "Uploaded implementation screenshot"),
            ("03/31 21:12", "Yanyan Meng", "Deployed and sent report to Tara for verification"),
            ("04/01 14:05", "Tara Bernstein", "Validated: \u201cReports look good. Please feel free to close out this request.\u201d"),
            ("04/03 07:14", "Dan Fallon", "Work note: Change Management Review Completed"),
        ],
        "gaps": [
            "Backout plan is inadequate \u2014 \u201cRollback to previous version\u201d with no specifics on "
            "which version or SQL rollback steps for 4+ table/SP changes. Backout window explicitly says \u201cno.\u201d",
        ],
        "observations": [
            "Configuration Item listed as \u201cOffice 365 - SharePoint\u201d but actual implementation was on "
            "PSQL1 database (stored procedures, tables, SSIS). The Partner Advancement Dashboard may live on "
            "SharePoint, but the changes were to SQL Server. Consider updating CI.",
            "Vendor support answered \u201cNone\u201d \u2014 should be \u201cNo\u201d or \u201cN/A\u201d for clarity.",
        ],
        "recommendation": (
            "Implementation successful and validated by requester. Backout plan gap noted for process "
            "improvement tracking. Ready for closure."
        ),
        "closure_ready": True,
        "impl_status": "OK", "backout_status": "WEAK", "test_status": "OK",
    },
    "CHG0039131": {
        "summary": (
            "New SuiteScript (UES_updateMainMemo) deployed to NetSuite Production to automate population "
            "of the body-level memo field on depreciation journals. Tested in NS Sandbox by Jitin, Leslie, "
            "and Johnny Chan. Deployed during 03/26 3\u20135am offshore window per RLSE0012902."
        ),
        "activity": [
            ("03/19 05:20", "Brian Ouderkirk", "Work note asking Jitin about script location and support documentation"),
            ("03/23 23:23", "Jitin Xavier", "Script is in NS Sandbox Files and Documents"),
            ("03/24 05:23", "Brian Ouderkirk", "Approved, assigned to RLSE0012902 for 03/26 3\u20135am"),
            ("03/26 07:25", "Jitin Xavier", "Deployed: script UES_updateMainMemo deployed to NS Production for all roles"),
            ("03/27 12:49", "Johnny Chan", "Memo field works, but received fail status on lease interest run for first time"),
            ("03/27 14:29", "Dan Fallon", "Work note forwarding Johnny\u2019s concern to Jitin and Hunter Dix"),
            ("03/30 07:57", "Jitin Xavier", "Requested error details from Johnny"),
            ("03/30 10:01", "Johnny Chan", "Uploaded Excel with error details (CHG0039131.xlsx)"),
            ("03/31 05:13", "Brian Ouderkirk", "Forwarded Johnny\u2019s update to Jitin"),
            ("04/01 12:37", "Jitin Xavier", "Confirmed fail is not caused by the script \u2014 unrelated to lease/depreciation run"),
        ],
        "gaps": [
            "Post-implementation test plan was \u201cNone, we will be monitoring\u201d \u2014 no defined validation "
            "criteria or timeframe for a script in a business-critical system.",
            "SAR question left blank \u2014 should be answered (N/A for a SuiteScript deployment).",
        ],
        "observations": [
            "Johnny Chan reported a lease interest fail status post-deployment. Jitin confirmed it\u2019s "
            "unrelated to the script. If the lease interest issue persists, a separate INC should be opened.",
            "Change was in pipeline for 90 days (created 01/02/2026).",
        ],
        "recommendation": (
            "Implementation successful. Memo field is populating correctly. Lease interest fail status "
            "confirmed unrelated by implementer. Ready for closure."
        ),
        "closure_ready": True,
        "impl_status": "OK", "backout_status": "OK", "test_status": "WEAK",
    },
}

CROSS_CUTTING = [
    ("Backout plan quality varies",
     "1/4 WEAK",
     "Medium",
     "CHG0039333 has 8 implementation steps but generic rollback with no backout window"),
    ("Post-implementation test gaps",
     "2/4 WEAK",
     "Medium",
     "CHG0039131 had no defined test criteria; CHG0039513 tests layout only, not automation logic"),
    ("Standard Change Matrix gap",
     "1/4",
     "Low",
     "Retirement contribution data loading not in Pre-Approved Matrix \u2014 recommend adding to Portal section"),
    ("CI accuracy",
     "1/4",
     "Low",
     "CHG0039333 lists SharePoint but implementation was on PSQL1 SQL Server"),
]

PROCESS_IMPROVEMENTS = [
    "Add \u201cRetirement Contribution Statement data loading\u201d to the Portal section of the Standard "
    "Change Pre-Approved Matrix.",
    "Encourage implementers to use Work Notes (not just Comments) for implementation status updates \u2014 "
    "Work Notes are internal and more appropriate for change audit trail.",
    "For changes involving multiple database objects (tables, stored procedures, SSIS packages), require "
    "specific rollback steps for each object, not a generic \u201crollback to previous version.\u201d",
    "For automation deployments (scripts, flows, triggers), post-implementation testing should verify "
    "the automation logic fires correctly \u2014 not just that UI elements are visible.",
]


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def add_status_color(run, status):
    if status == "OK":
        run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
    elif status == "WEAK":
        run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
    else:
        run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run.bold = True


def build_document() -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)

    # Title
    title = doc.add_heading("Post-Implementation Review Summary", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Review Date: {REVIEW_DATE}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Assignment Group: Enterprise Applications (Brian Ouderkirk)")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Policy Reference: ISMS-STA-11.01-01 Enterprise Change Management Program")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)

    passing = sum(1 for s in SCORECARD if s["status"] == "PASS")
    total = len(SCORECARD)
    ready = sum(1 for n in NOTES.values() if n["closure_ready"])
    doc.add_paragraph(
        f"{total} Enterprise Applications change requests in Review state were analyzed against "
        f"the 14 policy-required fields defined in ISMS-STA-11.01-01. {passing} of {total} are fully "
        f"compliant; {total - passing} have minor documentation gaps. {ready} of {total} are ready for closure."
    )

    # Scorecard table
    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Change", "Score", "Status", "Type", "Assigned To", "Closure Ready", "Release"]
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr[i].text = label
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for item in SCORECARD:
        note = NOTES[item["number"]]
        row = table.add_row().cells
        # Hyperlinked change number
        row[0].paragraphs[0].clear()
        add_hyperlink(row[0].paragraphs[0], item["number"], SNOW_URL + item["number"])
        row[1].text = item["score"]
        row[2].text = item["status"]
        row[3].text = item["type"]
        row[4].text = item["assigned"]
        row[5].text = "Yes" if note["closure_ready"] else "Pending"
        row[6].text = item.get("release", "")
        # Color status
        for paragraph in row[2].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if item["status"] == "PASS":
                    run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
                else:
                    run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)
        # Color closure ready
        for paragraph in row[5].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if note["closure_ready"]:
                    run.font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
                else:
                    run.font.color.rgb = RGBColor(0xCC, 0x7A, 0x00)

    doc.add_paragraph()

    # --- Individual Change Reviews ---
    doc.add_heading("Individual Change Reviews", level=1)

    for item in SCORECARD:
        chg = item["number"]
        note = NOTES[chg]

        # Change heading with hyperlink
        h = doc.add_heading(level=2)
        add_hyperlink(h, chg, SNOW_URL + chg)
        h.add_run(f" \u2014 {item['short_desc']}")

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run("Assigned to: ").bold = True
        meta.add_run(f"{item['assigned']}  |  ")
        meta.add_run("Requested by: ").bold = True
        meta.add_run(f"{item.get('requested_by', '')}  |  ")
        meta.add_run("Type: ").bold = True
        meta.add_run(f"{item['type']}  |  ")
        meta.add_run("CI: ").bold = True
        meta.add_run(item.get("ci", ""))

        meta2 = doc.add_paragraph()
        meta2.add_run("Planned Window: ").bold = True
        meta2.add_run(f"{item.get('planned_start', '')} \u2013 {item.get('planned_end', '')}  |  ")
        meta2.add_run("Release: ").bold = True
        meta2.add_run(item.get("release", ""))

        # Plan assessments
        plan_p = doc.add_paragraph()
        plan_p.add_run("Plan Assessments:  ").bold = True
        plan_p.add_run("Implementation: ")
        r = plan_p.add_run(note["impl_status"])
        add_status_color(r, note["impl_status"])
        plan_p.add_run("  |  Backout: ")
        r = plan_p.add_run(note["backout_status"])
        add_status_color(r, note["backout_status"])
        plan_p.add_run("  |  Test: ")
        r = plan_p.add_run(note["test_status"])
        add_status_color(r, note["test_status"])

        # Summary
        doc.add_heading("Review Summary", level=3)
        doc.add_paragraph(note["summary"])

        # Activity Log
        doc.add_heading("Activity Log", level=3)
        if note["activity"]:
            act_table = doc.add_table(rows=1, cols=3)
            act_table.style = "Light Grid Accent 1"
            act_hdr = act_table.rows[0].cells
            for i, label in enumerate(["Date/Time", "Author", "Activity"]):
                act_hdr[i].text = label
                for paragraph in act_hdr[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for date_str, author, desc in note["activity"]:
                arow = act_table.add_row().cells
                arow[0].text = date_str
                arow[1].text = author
                arow[2].text = desc
        else:
            doc.add_paragraph("No activity entries.")

        # Gaps
        if note["gaps"]:
            doc.add_heading("Documentation Gaps", level=3)
            for gap in note["gaps"]:
                doc.add_paragraph(gap, style="List Bullet")

        # Observations
        if note["observations"]:
            doc.add_heading("Observations", level=3)
            for obs in note["observations"]:
                doc.add_paragraph(obs, style="List Bullet")

        # Recommendation
        doc.add_heading("Recommendation", level=3)
        rec_p = doc.add_paragraph()
        closure_icon = "\u2705" if note["closure_ready"] else "\u23f3"
        rec_p.add_run(f"{closure_icon} ").font.size = Pt(12)
        rec_p.add_run(note["recommendation"]).italic = True

        doc.add_page_break()

    # --- Cross-Cutting Observations ---
    doc.add_heading("Cross-Cutting Observations", level=1)
    doc.add_paragraph(
        "The following patterns were observed across the reviewed changes."
    )

    table2 = doc.add_table(rows=1, cols=4)
    table2.style = "Light Grid Accent 1"
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = table2.rows[0].cells
    for i, label in enumerate(["Finding", "Frequency", "Severity", "Detail"]):
        hdr2[i].text = label
        for paragraph in hdr2[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for finding, freq, sev, detail in CROSS_CUTTING:
        row = table2.add_row().cells
        row[0].text = finding
        row[1].text = freq
        row[2].text = sev
        row[3].text = detail

    doc.add_paragraph()

    # --- Process Improvements ---
    doc.add_heading("Process Improvement Recommendations", level=1)
    for i, rec in enumerate(PROCESS_IMPROVEMENTS, 1):
        doc.add_paragraph(rec, style="List Number")

    return doc


def main():
    doc = build_document()
    filename = f"PIR Review Summary {REVIEW_DATE} - Enterprise Applications.docx"
    if OUTPUT_DIR.exists():
        out_path = OUTPUT_DIR / filename
    else:
        out_path = Path.cwd() / filename
    doc.save(str(out_path))
    print(f"Document saved to: {out_path}")


if __name__ == "__main__":
    main()
