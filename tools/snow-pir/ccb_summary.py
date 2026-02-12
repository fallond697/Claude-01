"""Generate CCB Review Summary Word document."""
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.expanduser(
    r"~\OneDrive - Vituity\Documents\Change Management\CCB\2026"
)
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "CCB Review Summary 2026-02-12.docx")

doc = Document()

# -- Styles --
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

# -- Title --
title = doc.add_heading("Enterprise CCB Review Summary", level=1)
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph(f"Date: Thursday, February 12, 2026")
doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
doc.add_paragraph("Changes at Assess: 10  |  Changes at Authorize: 0")
doc.add_paragraph("")

# -- Readiness Summary Table --
doc.add_heading("Plan Readiness Summary", level=2)

changes = [
    ("CHG0039318", "QuantaStor Core Upgrade 6.6.5", "Normal", "Moderate", "OK", "WEAK", "OK", "Review"),
    ("CHG0039314", "Add fields to AIDX extracts (DW/Informatics/VITTEST)", "Standard", "Low", "OK", "OK", "OK", "Ready"),
    ("CHG0039310", "JAMF: Deploy Tanium client to Mac devices", "Normal", "Low", "OK", "OK", "OK", "Ready"),
    ("CHG0039298", "athenaIDX Monthly PM (Feb 15)", "Normal", "Low", "OK", "WEAK", "OK", "Ready"),
    ("CHG0039288", "Server Decommission ORDC-SQLTMSRV", "Normal", "Low", "OK", "OK", "OK", "Ready"),
    ("CHG0039212", "IsLatest logic for dataset scan reporting", "Standard", "Low", "MISSING", "MISSING", "MISSING", "Not Ready"),
    ("CHG0039211", "409 Provider Errors for Cred Events", "Standard", "Low", "OK", "OK", "OK", "Ready"),
    ("CHG0039154", "MOOV Email Automation in HCM", "Normal", "Low", "OK", "OK", "OK", "Ready"),
    ("CHG0039148", "Oregon license expiration date issue", "Standard", "Low", "WEAK", "MISSING", "MISSING", "Not Ready"),
    ("CHG0039146", "Salesforce Spring '26 Release", "Normal", "Low", "OK", "OK", "OK", "Ready"),
]

headers = ["CHG #", "Description", "Type", "Risk", "Impl", "Backout", "Test", "Verdict"]
table = doc.add_table(rows=1, cols=len(headers))
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)

# Color map for verdicts and plan status
verdict_colors = {
    "Ready": RGBColor(0x22, 0x7A, 0x22),
    "Review": RGBColor(0xCC, 0x88, 0x00),
    "Not Ready": RGBColor(0xCC, 0x22, 0x22),
}
plan_colors = {
    "OK": RGBColor(0x22, 0x7A, 0x22),
    "WEAK": RGBColor(0xCC, 0x88, 0x00),
    "MISSING": RGBColor(0xCC, 0x22, 0x22),
}

for row_data in changes:
    row = table.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
            # Color the verdict column
            if i == 7 and val in verdict_colors:
                p.runs[0].font.color.rgb = verdict_colors[val]
                p.runs[0].bold = True
            # Color plan status columns
            if i in (4, 5, 6) and val in plan_colors:
                p.runs[0].font.color.rgb = plan_colors[val]
                if val != "OK":
                    p.runs[0].bold = True

doc.add_paragraph("")

# -- Not Ready Section --
doc.add_heading("Not Ready (2)", level=2)

p = doc.add_paragraph()
r = p.add_run("CHG0039212 - IsLatest logic for dataset scan reporting")
r.bold = True
r.font.size = Pt(10)
doc.add_paragraph(
    "Assigned To: Parvathi Arun (Enterprise Applications)\n"
    "Issue: Implementation, Backout, and Test plans are all empty. "
    "Justification is listed as \"tbd\". Risk & Impact section is entirely unanswered.\n"
    "Action: Send back to assignee. All plan fields must be completed before CCB can review."
)
doc.add_paragraph("")

p = doc.add_paragraph()
r = p.add_run("CHG0039148 - Oregon license expiration date issue")
r.bold = True
r.font.size = Pt(10)
doc.add_paragraph(
    "Assigned To: Parvathi Arun (Enterprise Applications)\n"
    "Issue: Implementation plan describes intent but lacks discrete numbered steps. "
    "Backout and Test plans are both empty. Risk & Impact section is mostly unanswered.\n"
    "Action: Send back to assignee. Needs concrete implementation steps, backout procedure, "
    "and test plan before CCB can review."
)
doc.add_paragraph("")

# -- Needs Discussion Section --
doc.add_heading("Needs Discussion (1)", level=2)

p = doc.add_paragraph()
r = p.add_run("CHG0039318 - QuantaStor Core Upgrade 6.6.5 + TRIM Remediation")
r.bold = True
r.font.size = Pt(10)
doc.add_paragraph(
    "Assigned To: Michael Castro (Enterprise Systems)\n"
    "Risk: Moderate | PHI/PII: Yes | Business-Critical: Yes\n"
    "Issue: Backout plan explicitly states \"No steps to revert\" and change window does not "
    "include backout time. This is a moderate-risk change affecting business-critical storage "
    "infrastructure (Mirth, Laserfiche, EDI, all Modesto VMs).\n"
    "Mitigating factors: OSNEXUS vendor support will be live during the change. "
    "Non-reboot upgrade per vendor recommendation.\n"
    "Discussion point: What is the recovery path if the upgrade destabilizes the cluster? "
    "Is there a snapshot or backup strategy before the upgrade begins?"
)
doc.add_paragraph("")

# -- Recurring/Routine Section --
doc.add_heading("Recurring / Routine (1)", level=2)

p = doc.add_paragraph()
r = p.add_run("CHG0039298 - athenaIDX Monthly PM (Feb 15)")
r.bold = True
r.font.size = Pt(10)
doc.add_paragraph(
    "Assigned To: Jose Ortiz (App Support - RCM MedFM)\n"
    "Monthly vendor-managed maintenance. Runs every month with established procedures. "
    "Backout plan is thin (\"vendor will complete any roll back\") but consistent with "
    "prior months. Window: Sun Feb 15, 1AM-8AM ET. Sites: RCM.\n"
    "Recommendation: Approve as routine."
)
doc.add_paragraph("")

# -- Ready Section --
doc.add_heading("Ready (6)", level=2)

ready_changes = [
    ("CHG0039314", "Add fields to AIDX extracts", "Avinash Vedavyas Prabhu", "Data Ops - PM",
     "Standard change. 3 discrete implementation steps with owners. Backout: revert DDL. "
     "Test: counts test + automated Weather check. 0 downtime."),
    ("CHG0039310", "JAMF: Deploy Tanium client to Macs", "Jason Nguyen", "Service Delivery Optimization",
     "Tanium client v7.8.2.1111 deployed silently via Jamf Pro policy. "
     "Pre-tested on IT dept devices. Backout: revoke config profile via APNS. No user impact."),
    ("CHG0039288", "Server Decommission ORDC-SQLTMSRV", "Khalil Douglas", "Enterprise Systems",
     "Decommission of unused server. Users already migrated to ORDC-SRVJMPTM1/TM2 and validated. "
     "Power off for 30 days before removal. Backout: power back on."),
    ("CHG0039211", "409 Provider Errors for Cred Events", "Parvathi Arun", "Enterprise Applications",
     "Data fix for 409 providers with mismatched cred event statuses between Salesforce and Verifiable. "
     "Detailed pre/impl/post steps. Pre-tested manually on one cred event. Vy Vu will validate."),
    ("CHG0039154", "MOOV Email Automation in HCM", "Bala Subramanyam Tadisetty", "Enterprise Applications",
     "Workato recipe update for MOOV employee email automation. 4 implementation steps. "
     "Tested in lower env. Backout: revert to previous recipe version. 0 downtime."),
    ("CHG0039146", "Salesforce Spring '26 Release", "Brian Ouderkirk", "Enterprise Applications",
     "Vendor-executed quarterly upgrade (Winter '26 -> Spring '26). "
     "Window: Feb 20, 10:30-11:00 PM PST. Tested on SB2 by 8 business teams. "
     "Post-impl validation by all teams. <30min outage."),
]

for chg_num, desc, assignee, group, notes in ready_changes:
    p = doc.add_paragraph()
    r = p.add_run(f"{chg_num} - {desc}")
    r.bold = True
    r.font.size = Pt(10)
    doc.add_paragraph(f"Assigned To: {assignee} ({group})\n{notes}")
    doc.add_paragraph("")

# -- Footer --
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("--- End of CCB Review Summary ---")
r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# Save
os.makedirs(OUTPUT_DIR, exist_ok=True)
doc.save(OUTPUT_FILE)
print(f"Saved: {OUTPUT_FILE}")
print(f"Size: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")
