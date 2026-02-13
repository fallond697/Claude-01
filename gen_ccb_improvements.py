from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('CCB Process Improvement Recommendations', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Weekly Enterprise Change Control Board\nFebruary 13, 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()

# Meta table
meta = doc.add_table(rows=4, cols=2)
meta.style = 'Light Shading Accent 1'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Prepared by', 'Dan Fallon, Enterprise Architecture'),
    ('Based on', 'CCB Meeting of February 12, 2026'),
    ('Policy Reference', 'ISMS-STA-11.01-01 Enterprise Change Management Program'),
    ('Classification', 'Internal Use Only'),
]
for i, (k, v) in enumerate(meta_data):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'Following the February 12, 2026 CCB meeting, a post-session review of prep documents '
    'and the meeting transcript identified six process improvement opportunities. These '
    'recommendations address documentation gaps, audit trail weaknesses, and workflow '
    'inefficiencies observed across multiple CCB cycles. Implementation of these improvements '
    'will reduce board prep time, improve change quality, and strengthen compliance with '
    'ISMS-STA-11.01-01.'
)

# Context
doc.add_heading('Review Context', level=1)
doc.add_paragraph(
    'The February 12 CCB reviewed 6 Normal change requests plus 1 reschedule (CHG0039273). '
    'Of these, 5 were approved, 1 was approved with full quorum vote (QuantaStor upgrade), '
    'and 1 was rescheduled to the following week. The meeting ran approximately 25 minutes '
    'with transcription enabled. Three prep documents were generated for this session, two '
    'of which were identical. The meeting transcript was downloaded from the Teams meeting '
    'recap and used as a primary source for this analysis.'
)

# Findings table
doc.add_heading('Findings Summary', level=1)

summary_table = doc.add_table(rows=7, cols=4)
summary_table.style = 'Medium Shading 1 Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['#', 'Recommendation', 'Effort', 'Impact']
for i, h in enumerate(headers):
    summary_table.rows[0].cells[i].text = h

rows_data = [
    ('1', 'Auto-sync transcript to OneDrive', 'Low', 'High'),
    ('2', 'Pre-screen changes with empty plans', 'Low', 'High'),
    ('3', 'Consolidate to single prep document', 'Low', 'Medium'),
    ('4', 'Standardize plan rating rubric', 'Medium', 'Medium'),
    ('5', 'Add question disposition tracking', 'Medium', 'Medium'),
    ('6', 'Implement post-CCB action log', 'Medium', 'High'),
]
for i, (num, rec, effort, impact) in enumerate(rows_data):
    row = summary_table.rows[i + 1]
    row.cells[0].text = num
    row.cells[1].text = rec
    row.cells[2].text = effort
    row.cells[3].text = impact

doc.add_paragraph()

# Detailed recommendations
doc.add_heading('Detailed Recommendations', level=1)

# Rec 1
doc.add_heading('1. Auto-Sync Meeting Transcript to OneDrive', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'The February 12 CCB meeting transcript was only available via manual download from the '
    'Teams meeting recap page. Other recurring meetings (HRIS Change Management, CMS Quality, '
    'Enterprise CM doc review) have transcripts that automatically sync to OneDrive. Without '
    'automatic sync, the transcript is at risk of being lost or inaccessible for audit purposes.'
)
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'Ensure the CCB meeting transcript automatically syncs to OneDrive alongside the '
    'recording, consistent with other recurring meetings. The transcript should be archived '
    'with the CCB prep documents for each session to create a readily accessible, auditable '
    'record of board decisions for change management compliance and post-incident review.'
)
p = doc.add_paragraph()
p.add_run('Effort: ').bold = True
p.add_run('Low \u2014 Teams meeting or OneDrive sync setting adjustment.')

# Rec 2
doc.add_heading('2. Pre-Screen Changes with Empty Plans', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'Changes with completely empty Implementation, Backout, and Test plans have reached the '
    'CCB agenda. In some cases, justification fields contain only "tbd" and Risk & Impact '
    'sections are entirely unanswered. These consume board discussion time despite being '
    'obviously incomplete.'
)
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'Implement a pre-CCB gate on Wednesday (before the 3:00 PM PT cutoff) that checks for '
    'empty plan fields on all changes at Assess state. Changes with empty Implementation, '
    'Backout, or Test plans should be automatically flagged and returned to the assignee '
    'with a standard notification. This could be a ServiceNow business rule or a scheduled '
    'report sent to change owners.'
)
p = doc.add_paragraph()
p.add_run('Effort: ').bold = True
p.add_run('Low \u2014 ServiceNow scheduled job or report rule.')

# Rec 3
doc.add_heading('3. Consolidate to Single Prep Document', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'Three prep documents were generated for the February 12 CCB: "CCB Prep - CAB Review," '
    '"CCB Prep - CCB Review," and "CCB-Review-2026-02-12." The first two are identical. '
    'The third uses a different format with implementation schedule tables, structured '
    'metadata cards per change, and a prioritized discussion section.'
)
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'Consolidate into a single document using the CCB-Review format, which is the most '
    'comprehensive. Retire the duplicate "CCB Prep" naming convention. The single document '
    'should serve as both preparation material and the meeting agenda.'
)
p = doc.add_paragraph()
p.add_run('Effort: ').bold = True
p.add_run('Low \u2014 Template standardization.')

# Rec 4
doc.add_heading('4. Standardize Plan Rating Rubric', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'The two prep document formats applied different ratings to the same change plans:'
)

rating_table = doc.add_table(rows=4, cols=4)
rating_table.style = 'Light List Accent 1'
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rating_headers = ['Change', 'Field', 'CCB Prep Rating', 'CCB-Review Rating']
for i, h in enumerate(rating_headers):
    rating_table.rows[0].cells[i].text = h
rating_data = [
    ('CHG0039298', 'Backout Plan', 'INCOMPLETE', 'OK'),
    ('CHG0039288', 'Test Plan', 'INCOMPLETE', 'OK'),
    ('CHG0039318', 'Test Plan', 'INCOMPLETE', 'OK'),
]
for i, (chg, field, prep, review) in enumerate(rating_data):
    row = rating_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = field
    row.cells[2].text = prep
    row.cells[3].text = review

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'Define a single plan rating rubric with clear criteria for each level:'
)

rubric = doc.add_table(rows=5, cols=2)
rubric.style = 'Light List Accent 1'
rubric.alignment = WD_TABLE_ALIGNMENT.CENTER
rubric.rows[0].cells[0].text = 'Rating'
rubric.rows[0].cells[1].text = 'Criteria'
rubric_data = [
    ('OK', 'Contains discrete, actionable steps appropriate to the change type.'),
    ('WEAK', 'Steps exist but lack specificity, rely entirely on vendor, or omit key details.'),
    ('INCOMPLETE', 'Field is empty, contains only "TBD" or "N/A" where a plan is required.'),
    ('N/A', 'Plan type genuinely does not apply (e.g., test plan for a decommission with 30-day hold).'),
]
for i, (rating, criteria) in enumerate(rubric_data):
    rubric.rows[i + 1].cells[0].text = rating
    rubric.rows[i + 1].cells[1].text = criteria

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Effort: ').bold = True
p.add_run('Medium \u2014 Requires agreement on definitions and consistent application.')

# Rec 5
doc.add_heading('5. Add Question Disposition Tracking', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'The prep documents raised 25+ questions across all changes. After the meeting, there '
    'is no record of which questions were answered, what the answers were, or which were '
    'deferred. This breaks the audit trail between prep and outcome.'
)
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'Add a disposition field to each question in the post-meeting document. Example statuses: '
    'Answered (with summary), Deferred (with owner/date), Waived (with rationale). This '
    'closes the loop between preparation and decision and provides evidence for audit.'
)
p = doc.add_paragraph()
p.add_run('Effort: ').bold = True
p.add_run('Medium \u2014 Requires real-time or immediate post-meeting annotation.')

# Rec 6
doc.add_heading('6. Implement Post-CCB Action Log', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'The Minutes of Meeting email instructed change owners to "advance the CR to the Review '
    'state and add post-deployment comments to the Notes." However, there is no structured '
    'mechanism to track whether these actions were completed. Conditional approvals (e.g., '
    '"fill in the blank risk fields before implementing") have no follow-up trigger.'
)
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'Create a post-CCB action log that is reviewed at the start of the next CCB meeting. '
    'Each action should include the change number, action description, owner, due date, '
    'and completion status:'
)

action_table = doc.add_table(rows=5, cols=5)
action_table.style = 'Medium Shading 1 Accent 1'
action_table.alignment = WD_TABLE_ALIGNMENT.CENTER
action_headers = ['Change', 'Action', 'Owner', 'Due', 'Status']
for i, h in enumerate(action_headers):
    action_table.rows[0].cells[i].text = h
action_data = [
    ('CHG0039212', 'Complete all plan fields and resubmit', 'Parvathi Arun', 'Next CCB', 'Open'),
    ('CHG0039148', 'Complete impl steps, backout, and test plan', 'Parvathi Arun', 'Next CCB', 'Open'),
    ('CHG0039318', 'Obtain documented recovery procedure from OSNEXUS', 'Michael Castro', 'Before 2/18', 'Open'),
    ('CHG0039310', 'Complete blank risk assessment fields', 'Jason Nguyen', 'Before impl', 'Open'),
]
for i, (chg, action, owner, due, status) in enumerate(action_data):
    row = action_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = action
    row.cells[2].text = owner
    row.cells[3].text = due
    row.cells[4].text = status

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Effort: ').bold = True
p.add_run('Medium \u2014 Requires discipline to maintain weekly; could be automated from ServiceNow.')

# Transcript Analysis
doc.add_heading('Meeting Transcript Analysis', level=1)
doc.add_paragraph(
    'The following analysis is based on the meeting transcript downloaded from the Teams '
    'meeting recap (February 12, 2026, 25 minutes, 12 department representatives present). '
    'It documents what was discussed, verbal commitments made, and gaps between prep '
    'questions and actual meeting coverage.'
)

# Changes reviewed table
doc.add_heading('Changes Reviewed', level=2)
changes_table = doc.add_table(rows=8, cols=5)
changes_table.style = 'Medium Shading 1 Accent 1'
changes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
ch_headers = ['Change', 'Description', 'Presenter', 'Time', 'Outcome']
for i, h in enumerate(ch_headers):
    changes_table.rows[0].cells[i].text = h
ch_data = [
    ('CHG0039146', 'SF Spring \'26 Release', 'Brian Ouderkirk', '2:00-3:48', 'Approved'),
    ('CHG0039310', 'JAMF/Tanium Mac Deploy', 'Jason Nguyen', '4:00-6:24', 'Approved'),
    ('CHG0039298', 'athenaIDX Monthly PM', 'Melissa Ramirez', '6:24-9:33', 'Approved'),
    ('CHG0039288', 'Server Decommission', 'Khalil Douglas', '9:33-11:31', 'Approved'),
    ('CHG0039154', 'MOOV Email Automation', 'Bala Tadisetty', '11:31-14:35', 'Approved'),
    ('CHG0039273', 'BGP Phase 1 (reschedule)', 'Daniele Bartoli', '14:35-16:47', 'Rescheduled to 2/17'),
    ('CHG0039318', 'QuantaStor Upgrade', 'Michael Castro', '16:47-25:28', 'Approved (quorum vote)'),
]
for i, (chg, desc, presenter, time, outcome) in enumerate(ch_data):
    row = changes_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = desc
    row.cells[2].text = presenter
    row.cells[3].text = time
    row.cells[4].text = outcome

doc.add_paragraph()

# Backout discussion finding
doc.add_heading('Finding: Backout Details Emerged Only During Discussion', level=2)
p = doc.add_paragraph()
p.add_run(
    'For CHG0039318 (QuantaStor), the backout discussion revealed critical information '
    'that was not documented in the change request:'
)
doc.add_paragraph(
    'Shawn Wood (TechOps) clarified that rollback means "they rebuild the machine" \u2014 '
    'reflashing firmware from scratch. This is a significantly more involved recovery than '
    'implied by the CR\'s "no steps to revert."',
    style='List Bullet'
)
doc.add_paragraph(
    'Michael Castro acknowledged "I\'m not 100% certain" on whether rollback was possible, '
    'confirming the backout plan gap identified in the prep documents.',
    style='List Bullet'
)
doc.add_paragraph(
    'Shawn Wood added context that this rebuild has only been needed once historically and '
    'that upgrades have been "pretty clean."',
    style='List Bullet'
)
p = doc.add_paragraph()
p.add_run('Implication: ').bold = True
p.add_run(
    'Change owners should document the specific rollback mechanism in the CR before CCB, '
    'especially for Moderate+ risk changes. Verbal clarifications like these should be '
    'captured in the CR notes during or immediately after the meeting.'
)

doc.add_paragraph()

# Verbal commitments
doc.add_heading('Finding: Untracked Verbal Commitments', level=2)
p = doc.add_paragraph()
p.add_run(
    'The following commitments were made verbally during the meeting. None are tracked '
    'in a post-meeting action log:'
)

commit_table = doc.add_table(rows=8, cols=3)
commit_table.style = 'Light List Accent 1'
commit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
commit_headers = ['Change', 'Commitment', 'Who']
for i, h in enumerate(commit_headers):
    commit_table.rows[0].cells[i].text = h
commit_data = [
    ('CHG0039298', 'Jose will log in after upgrade to validate', 'Melissa Ramirez'),
    ('CHG0039298', 'Follow up with Jose on test server patching', 'Dan Fallon'),
    ('CHG0039318', 'Send pre/post email notification to teams', 'Michael Castro'),
    ('CHG0039318', 'Contacted Brendan, Nina, George for awareness', 'Michael Castro'),
    ('CHG0039318', 'Will have someone verify on RCM DevOps end', 'Eric Sater'),
    ('CHG0039318', 'Will review Laserfiche impact and test', 'Avinash Vedavyas Prabhu'),
    ('CHG0039318', 'On-call person available for EDI validation', 'Mike Winget'),
]
for i, (chg, commit, who) in enumerate(commit_data):
    row = commit_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = commit
    row.cells[2].text = who

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Implication: ').bold = True
p.add_run(
    'These commitments currently exist only in the transcript. The post-CCB action log '
    '(Recommendation #6) would provide a structured mechanism to track these to completion.'
)

doc.add_paragraph()

# Prep question coverage
doc.add_heading('Finding: Prep Question Coverage Gaps', level=2)
p = doc.add_paragraph()
p.add_run(
    'The prep documents raised 25+ questions. Cross-referencing against the transcript '
    'reveals which were addressed in the meeting and which were not:'
)

q_table = doc.add_table(rows=9, cols=4)
q_table.style = 'Light List Accent 1'
q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
q_headers = ['Change', 'Prep Question', 'Asked?', 'Answer']
for i, h in enumerate(q_headers):
    q_table.rows[0].cells[i].text = h
q_data = [
    ('CHG0039310', 'Vendor support available?', 'Yes', 'Jason confirmed yes'),
    ('CHG0039310', 'Sites/BU impacted (blank field)?', 'No', 'Not discussed'),
    ('CHG0039298', 'Test servers patched before prod?', 'Yes', 'Confirmed done same morning'),
    ('CHG0039298', 'OpenID authcfg scope?', 'No', 'Not discussed'),
    ('CHG0039154', 'Excluded JobIDs \u2014 who maintains?', 'Yes', 'HCM team maintains'),
    ('CHG0039154', 'Retroactive email updates needed?', 'Yes', 'Manual; only new users'),
    ('CHG0039318', 'Backout/recovery path?', 'Yes', 'Fix-forward; rebuild if needed'),
    ('CHG0039318', 'Snapshots/backups before upgrade?', 'No', 'Not discussed'),
]
for i, (chg, question, asked, answer) in enumerate(q_data):
    row = q_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = question
    row.cells[2].text = asked
    row.cells[3].text = answer

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Implication: ').bold = True
p.add_run(
    'Several important prep questions went unasked during the meeting. A printed or shared '
    'checklist during the meeting would help ensure complete coverage. Question disposition '
    'tracking (Recommendation #5) would close this gap.'
)

doc.add_paragraph()

# Positive observation
doc.add_heading('Positive Observation: Quorum Voting on CHG0039318', level=2)
p = doc.add_paragraph()
p.add_run(
    'The QuantaStor upgrade (CHG0039318) correctly went through a per-department approval '
    'vote. Each department confirmed both approval and post-change validation availability:'
)

vote_table = doc.add_table(rows=10, cols=3)
vote_table.style = 'Light List Accent 1'
vote_table.alignment = WD_TABLE_ALIGNMENT.CENTER
vote_headers = ['Department', 'Representative', 'Validation Committed']
for i, h in enumerate(vote_headers):
    vote_table.rows[0].cells[i].text = h
vote_data = [
    ('InfoSec', 'Vinod Kashyap', 'Approved'),
    ('TechOps', 'Shawn Wood', 'Approved'),
    ('App Support / Business Analytics', 'Melissa Ramirez', 'Jose to validate'),
    ('Enterprise Applications', 'Brian Ouderkirk', 'George\'s team contacted'),
    ('Platform Integration', 'Divyarani Bhat', 'Will validate'),
    ('RCM DevOps', 'Eric Sater', 'Someone to verify'),
    ('Service Delivery', 'Mark White', 'Service Desk available'),
    ('Enterprise Data', 'Avinash Vedavyas Prabhu', 'Will review Laserfiche + test'),
    ('RCM EDI', 'Mike Winget', 'On-call available'),
]
for i, (dept, rep, validation) in enumerate(vote_data):
    row = vote_table.rows[i + 1]
    row.cells[0].text = dept
    row.cells[1].text = rep
    row.cells[2].text = validation

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Recommendation: ').bold = True
p.add_run(
    'This per-department vote with validation commitment is good practice. Consider '
    'formalizing it as the standard procedure for all Moderate and higher risk changes.'
)

doc.add_paragraph()

# Meeting efficiency
doc.add_heading('Meeting Efficiency', level=2)
p = doc.add_paragraph()
p.add_run(
    'The meeting was well-paced at 25 minutes for 7 agenda items. Straightforward changes '
    '(SF release, JAMF, decommission) took 2\u20133 minutes each. The QuantaStor upgrade '
    'appropriately consumed the most time (~9 minutes) given its Moderate risk rating and '
    'large blast radius. No improvement needed on pacing.'
)

doc.add_paragraph()

# Implementation priority
doc.add_heading('Implementation Priority', level=1)
doc.add_paragraph(
    'The following implementation order maximizes value while building on each improvement:'
)

priority = doc.add_table(rows=4, cols=3)
priority.style = 'Medium Shading 1 Accent 1'
priority.alignment = WD_TABLE_ALIGNMENT.CENTER
priority.rows[0].cells[0].text = 'Phase'
priority.rows[0].cells[1].text = 'Recommendations'
priority.rows[0].cells[2].text = 'Target'
priority_data = [
    ('Immediate', '#1 (Transcript sync), #3 (Consolidate docs)', 'Next CCB (Feb 19)'),
    ('Short-term', '#2 (Pre-screen), #4 (Rating rubric)', 'Within 2 weeks'),
    ('Medium-term', '#5 (Question tracking), #6 (Action log)', 'Within 30 days'),
]
for i, (phase, recs, target) in enumerate(priority_data):
    row = priority.rows[i + 1]
    row.cells[0].text = phase
    row.cells[1].text = recs
    row.cells[2].text = target

doc.add_paragraph()

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Enterprise Architecture | Vituity Information Technology Services')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

output_path = r'C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\CCB\2026\CCB Process Improvement Recommendations 2026-02-13.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
