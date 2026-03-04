import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def chg_url(number):
    return f"https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number={number}"

def add_table_row(table, cells, bold_first=True):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if i == 0 and bold_first:
            run.bold = True
    return row

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

title = doc.add_heading('CCB Meeting Preparation \u2014 March 5, 2026', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Meeting: ').bold = True
p.add_run('Change Control Board (CCB) \u2014 Thursday, March 5, 2026, 1:00 PM PST')
p = doc.add_paragraph()
p.add_run('Prepared by: ').bold = True
p.add_run('Dan Fallon (EA)')
p = doc.add_paragraph()
p.add_run('Changes for Review: ').bold = True
p.add_run('12 Normal Changes')
p = doc.add_paragraph()
p.add_run('Prepared: ').bold = True
p.add_run('March 4, 2026')

# SUMMARY TABLE
doc.add_heading('Executive Summary', level=1)
summary_table = doc.add_table(rows=1, cols=6)
summary_table.style = 'Light Grid Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = summary_table.rows[0].cells
for i, h in enumerate(['CHG #', 'Title', 'Group', 'Risk', 'Schedule', 'Ready?']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

changes_summary = [
    ('CHG0039420', 'Deploy Jump Server for EDI', 'Ent Sys', 'Low', '3/5 2\u20132:30PM', 'YES'),
    ('CHG0039418', 'Abnormal - Remove Domains from Safelist', 'Infosec', 'Low', '3/10\u20133/27', 'YES'),
    ('CHG0039415', 'Workato - Intune Out of Sync Devices', 'Ent Apps', 'Low', '3/5 8\u201310PM', 'YES'),
    ('CHG0039414', 'Update Okta RADIUS Agent', 'Infosec', 'Low', '3/6 10\u201311AM', 'YES'),
    ('CHG0039412', 'Upgrade Firmware on ORDC Hx', 'Ent Sys', 'Moderate', '3/11 6PM\u201312AM', 'YES*'),
    ('CHG0039403', 'ASR Policies for PP Users (Warning)', 'Infosec', 'Moderate', '3/6 7\u201310AM', 'YES'),
    ('CHG0039399', 'Apply Security Patches to F5 LBs', 'Ent Sys', 'Low', '3/6 6\u20138PM', 'YES'),
    ('CHG0039383', 'Vituity Stats SharePoint Page', 'Ent Apps', 'Low', '3/5 3\u20135PM', 'YES'),
    ('CHG0039339', 'SN to Otto Integration - Agent Alerts', 'SDO', 'Low', '3/5 3\u20134PM', 'YES'),
    ('CHG0039284', 'Enable Otto Triage to HR & Facilities', 'SDO', 'Low', '3/10 9\u201310AM', 'YES'),
    ('CHG0039283', 'Deactivate Forms - IT/Equip/Software', 'Ent Apps', 'Low', '3/10 9\u201310AM', 'YES'),
    ('CHG0039213', 'New Global ServiceNow Survey', 'Ent Apps', 'Low', '3/5 4\u20136PM', 'YES'),
]

for row_data in changes_summary:
    row = summary_table.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        if i == 0:
            add_hyperlink(p, val, chg_url(val))
        elif i == 5:
            run = p.add_run(val)
            run.font.size = Pt(9)
            if val == 'NO':
                run.font.color.rgb = RGBColor(0xCC, 0, 0)
                run.bold = True
            elif val.startswith('YES*'):
                run.font.color.rgb = RGBColor(0xCC, 0x88, 0)
                run.bold = True
            else:
                run.font.color.rgb = RGBColor(0, 0x80, 0)
                run.bold = True
        else:
            run = p.add_run(val)
            run.font.size = Pt(9)

doc.add_paragraph('')
p = doc.add_paragraph()
run = p.add_run('YES* = Ready with questions/notes for CCB discussion')
run.font.size = Pt(9)
run.italic = True

# DETAILED REVIEWS
doc.add_heading('Detailed Change Reviews', level=1)

# --- 1. CHG0039420 ---
doc.add_heading('1. CHG0039420 \u2014 Deploy Jump Server for EDI', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Michael Castro (Enterprise Systems)'])
add_table_row(t, ['Requested By', 'Michael Castro'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/05 2:00 PM \u2013 2:30 PM PST (30 min)'])
add_table_row(t, ['Server', 'RCM-SRVJUMPEDI (Modesto, DHCP)'])
add_table_row(t, ['Description', 'Deploy a new jump server VM in Modesto for EDI team to RDP into systems blocked by Illumio rules. Discovered during DR testing.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 Deploy from VM template, set DHCP, domain join, import RDS licensing, grant access to 5 named EDI users'])
add_table_row(t2, ['Backout', 'OK \u2014 Delete the VM if not needed'])
add_table_row(t2, ['Test', 'OK \u2014 EDI team will confirm tasks/functions work through jump server'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 All answered. Not business-critical, no PHI/PII, 0 outage, RCM only'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Will the jump server be hardened (e.g., restricted RDP access, MFA, session logging)?', style='List Bullet')
doc.add_paragraph('Is DHCP appropriate for a server, or should it have a static/reserved IP for security monitoring?', style='List Bullet')
doc.add_paragraph('Are Illumio rules being updated to allow RDP from this jump server to the target EDI systems?', style='List Bullet')
doc.add_paragraph('Will the 5 named users have local admin or standard user access on the jump server?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 2. CHG0039418 ---
doc.add_heading('2. CHG0039418 \u2014 Abnormal - Remove Domains from Safelist', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Bill Carter (Infosec)'])
add_table_row(t, ['Requested By', 'Bill Carter'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/10 10:00 AM \u2013 03/27 10:30 AM (rolling schedule)'])
add_table_row(t, ['Description', 'Remove domains from Abnormal safelist so security scans detect compromised trusted sender inboxes.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 7 steps incl. monitoring and false positive remediation'])
add_table_row(t2, ['Backout', 'OK \u2014 Re-add domains'])
add_table_row(t2, ['Test', 'OK \u2014 Tested with 5 domains, no issues'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 All answered. Business-critical: Yes, PHI/PII: Yes, 0 outage'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('How many total domains are being removed from the safelist?', style='List Bullet')
doc.add_paragraph('Can you walk CCB through the phased rolling schedule (attachment)?', style='List Bullet')
doc.add_paragraph('PHI/PII marked Yes \u2014 what is the data exposure risk if false positives quarantine legitimate email?', style='List Bullet')
doc.add_paragraph('What was the false positive rate in the 5-domain test?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 3. CHG0039415 ---
doc.add_heading('3. CHG0039415 \u2014 Workato: Intune Out of Sync Device Notifications', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Pallav Malu (Enterprise Applications)'])
add_table_row(t, ['Requested By', 'David Chu'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/05 8:00 PM \u2013 10:00 PM PST (2 hours)'])
add_table_row(t, ['Description', 'Two new Workato recipes to auto-notify PM/RCM users with Intune devices not checked in 30\u201390 days.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 4 steps, two named recipes, documentation linked'])
add_table_row(t2, ['Backout', 'OK \u2014 Stop both recipes in Workato PROD'])
add_table_row(t2, ['Test', 'OK \u2014 UAT performed, data validated by requestor'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 All answered. No PHI/PII, 0 outage, vendor support available'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('How does the recipe prevent duplicate notifications if a device stays out of sync across daily runs?', style='List Bullet')
doc.add_paragraph('What does the remediation email instruct users to do? Is there a self-service link?', style='List Bullet')
doc.add_paragraph('Are shared/kiosk devices excluded from the notification criteria?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 4. CHG0039414 ---
doc.add_heading('4. CHG0039414 \u2014 Update Okta RADIUS Agent', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Jeffrey How (Infosec)'])
add_table_row(t, ['Requested By', 'Jeffrey How'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/06 10:00 AM \u2013 11:00 AM PST (1 hour)'])
add_table_row(t, ['Servers', 'rcm-srvokta, ordc-srvokta'])
add_table_row(t, ['Version', '2.24.2 \u2192 2.26.0'])
add_table_row(t, ['Description', 'Update Okta RADIUS Agent on two servers to address security vulnerability. No production services currently use Okta RADIUS (VPN/Horizon not in use).'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 4 steps: RDP, install v2.26.0, confirm, repeat for second server'])
add_table_row(t2, ['Backout', 'OK \u2014 Reinstall v2.24.2 or restore snapshot/backup. Trigger: RADIUS services don\'t come back up.'])
add_table_row(t2, ['Test', 'OK \u2014 Prior smooth updates; post-check: RADIUS service started, agent registered in Okta'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 All answered. Not business-critical, no PHI/PII, 0 outage, HA: Yes, vendor support: Yes'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('What specific CVE or vulnerability is being remediated?', style='List Bullet')
doc.add_paragraph('If no production services currently use RADIUS, what is the planned future use case?', style='List Bullet')
doc.add_paragraph('Are the servers being patched for OS vulnerabilities at the same time?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 5. CHG0039412 ---
doc.add_heading('5. CHG0039412 \u2014 Upgrade Firmware on ORDC Hx', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Philip Weiss (Enterprise Systems)'])
add_table_row(t, ['Requested By', 'Philip Weiss'])
add_table_row(t, ['Risk / Impact', 'Moderate / 2 - Medium'])
add_table_row(t, ['Planned Window', '03/11 6:00 PM \u2013 11:59 PM PST (~6 hours)'])
add_table_row(t, ['Description', 'Multi-step firmware upgrade on ORDC HyperFlex cluster.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('In-Scope Devices (14):').bold = True
doc.add_paragraph('Fabric Interconnects: Hx_Fab_A, Hx_Fab_B', style='List Bullet')
doc.add_paragraph('ESXi Hosts: ordc-hxsrvesxi01 through ordc-hxsrvesxi12 (.mbsi.medamerica.local)', style='List Bullet')
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Version Upgrades:').bold = True
t3 = doc.add_table(rows=0, cols=2); t3.style = 'Light Grid Accent 1'
add_table_row(t3, ['Hx Firmware', '4.0(4k) \u2192 4.0(4n) \u2192 4.2(3o)'])
add_table_row(t3, ['HyperFlex DP', '4.5(2e) \u2192 5.5(2b)'])
add_table_row(t3, ['ESXi', '7.0.3 \u2192 8.03'])
add_table_row(t3, ['vCenter', 'Already upgraded to 8.03 (completed)'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 4 phased steps, vCenter already done'])
add_table_row(t2, ['Backout', 'LIMITED \u2014 States reverting HxDP/ESXi is "generally not supported"'])
add_table_row(t2, ['Test', 'OK \u2014 Prior UCS upgrade experience, Cisco TAC on standby'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 Business-critical: Yes, 0 outage, HA: Yes, vendor support: Yes'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Backout states reverting is "generally not supported" \u2014 what is the actual recovery path if the combined upgrade fails?', style='List Bullet')
doc.add_paragraph('Are VMs live-migrated during the host-by-host upgrade, or is there expected downtime?', style='List Bullet')
doc.add_paragraph('ESXi target "8.03" \u2014 confirmed compatible with all running workloads?', style='List Bullet')
doc.add_paragraph('Is Cisco TAC proactively engaged or on standby only? This is the highest-risk change on the agenda.', style='List Bullet')
doc.add_paragraph('ESXi prior version field is incomplete ("7.0.3, 24784741 ->") \u2014 what is the exact target build?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE WITH DISCUSSION \u2014 Backout limitations and scope warrant CCB discussion.'); run.bold = True; run.font.color.rgb = RGBColor(0xCC, 0x88, 0)

# --- 6. CHG0039403 ---
doc.add_heading('6. CHG0039403 \u2014 ASR Policies for PP Department Users (Warning Mode)', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Chintan Myakal (Infosec)'])
add_table_row(t, ['Requested By', 'Chintan Myakal'])
add_table_row(t, ['Risk / Impact', 'Moderate / 2 - Medium'])
add_table_row(t, ['Planned Window', '03/06 7:00 AM \u2013 10:00 AM PST (3 hours)'])
add_table_row(t, ['Target Groups', 'PP - Medical Directors-HM-AL, PP - Scribes-AL'])
add_table_row(t, ['Description', 'Deploy MS Defender Attack Surface Reduction policies in Warning mode to PP users via Intune.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 9 detailed Intune steps: create policy, Warn mode, exclusions, assign to PP groups'])
add_table_row(t2, ['Backout', 'OK \u2014 Change ASR rules from Warn to Audit mode'])
add_table_row(t2, ['Test', 'OK \u2014 Tested with PP test group; exclusions tuned; no false warnings after tuning'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 Not business-critical, no PHI/PII, 0 outage, PP only'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('How many users are in the two target groups (Medical Directors-HM-AL and Scribes-AL)?', style='List Bullet')
doc.add_paragraph('Warning mode allows the action with a prompt \u2014 is there a planned timeline to move to Block?', style='List Bullet')
doc.add_paragraph('How many exclusions were needed during testing? Are they documented?', style='List Bullet')
doc.add_paragraph('Controlled Folder Access set to Audit (not Warn) \u2014 rationale for different mode?', style='List Bullet')
doc.add_paragraph('Is there a communication plan to PP users about the warning prompts?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 7. CHG0039399 ---
doc.add_heading('7. CHG0039399 \u2014 Apply Security Patches to F5 Load Balancers', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Philip Weiss (Enterprise Systems)'])
add_table_row(t, ['Requested By', 'Philip Weiss'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/06 6:00 PM \u2013 8:00 PM PST (2 hours)'])
add_table_row(t, ['Servers', 'ORDC-LB01, ORDC-LB02'])
add_table_row(t, ['Version', 'BIGIP-17.5.1-0.0.7 \u2192 BIGIP-17.5.1.4-0.0.20'])
add_table_row(t, ['Description', 'Security patching on F5 LBs for vulnerabilities K000156644, K000156643.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 Patch standby, failover, verify, repeat'])
add_table_row(t2, ['Backout', 'OK \u2014 Fail traffic back to unpatched unit'])
add_table_row(t2, ['Test', 'OK \u2014 Verify traffic passing; prior patching experience'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 Business-critical: Yes, PHI/PII: Yes, 0 outage (rolling), HA: Yes'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('F5 advisories K000156644, K000156643 \u2014 severity rating and actively exploited?', style='List Bullet')
doc.add_paragraph('During failover, any brief traffic interruption or truly seamless?', style='List Bullet')
doc.add_paragraph('PHI/PII marked Yes \u2014 because F5 handles SSL termination for PHI-bearing apps?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 8. CHG0039383 ---
doc.add_heading('8. CHG0039383 \u2014 Vituity Stats SharePoint Page', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'DivyaRani Bhat (Enterprise Applications)'])
add_table_row(t, ['Requested By', 'Amy Hughes'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/05 3:00 PM \u2013 5:00 PM PST (2 hours)'])
add_table_row(t, ['Description', 'New pipeline: Azure SQL SPs \u2192 Workato recipe \u2192 CSV files \u2192 SharePoint page.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 4 stored procedures, 1 Workato recipe, 4 CSVs, docs linked'])
add_table_row(t2, ['Backout', 'OK \u2014 Disable Workato recipe'])
add_table_row(t2, ['Test', 'OK \u2014 SP validation, Workato testing, data validation, monitoring'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 All answered. No PHI/PII, 0 outage, QA tested, HA: Yes'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Do the stored procedures impact Azure SQL SFDC performance during execution?', style='List Bullet')
doc.add_paragraph('What is the scheduled frequency for the Workato recipe?', style='List Bullet')
doc.add_paragraph('Who will have access to the SharePoint page \u2014 all Vitans immediately?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 9. CHG0039339 ---
doc.add_heading('9. CHG0039339 \u2014 ServiceNow to Otto Integration - Agent Alerts', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Dan Spengler (Service Delivery Optimization)'])
add_table_row(t, ['Requested By', 'Dan Spengler'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/05 3:00 PM \u2013 4:00 PM PST (1 hour)'])
add_table_row(t, ['Description', 'New SN-to-Otto integration: alerts agents via Otto when tickets assigned. Update set with 2 flows, Moveworks listener/plugin, opt-out field.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 Extremely detailed. All dev links provided.'])
add_table_row(t2, ['Backout', 'OK \u2014 Disable flows/plugins, back out update set'])
add_table_row(t2, ['Test', 'OK \u2014 Tested in dev/sandbox with multiple agents'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 No PHI/PII, 0 outage, SAR completed'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Opt-out checkbox on sys_user \u2014 reviewed with SN platform team for schema concerns?', style='List Bullet')
doc.add_paragraph('OAuth 2.0 credentials for Moveworks \u2014 stored securely in SN (connection alias)?', style='List Bullet')
doc.add_paragraph('If Otto/Moveworks is down, do flows fail silently or generate SN errors?', style='List Bullet')
doc.add_paragraph('Vendor support not during change \u2014 risk if prod behaves differently than sandbox?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 10. CHG0039284 ---
doc.add_heading('10. CHG0039284 \u2014 Enable Otto Triage to HR & Facilities', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Beth Vanderheiden (Service Delivery Optimization)'])
add_table_row(t, ['Requested By', 'Mark White'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/10 9:00 AM \u2013 10:00 AM PST (1 hour)'])
add_table_row(t, ['Description', 'Remove Moveworks/Otto triage restrictions for HR and Facilities groups.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 4 steps in Moveworks portal; 3 groups remain blocked'])
add_table_row(t2, ['Backout', 'OK \u2014 Re-add groups to blocked list'])
add_table_row(t2, ['Test', 'OK \u2014 Tested in dev; analytics monitoring planned'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 Not business-critical, no PHI/PII, 0 outage'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Why are LMS Support, L&D, and HRIS Tier 3 kept blocked?', style='List Bullet')
doc.add_paragraph('Have HR and Facilities teams been informed Otto will route tickets to them?', style='List Bullet')
doc.add_paragraph('What is the expected ticket volume to these new groups?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 11. CHG0039283 ---
doc.add_heading('11. CHG0039283 \u2014 Deactivate Forms: IT Request, Equipment, Software', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Beth Vanderheiden (Enterprise Applications)'])
add_table_row(t, ['Requested By', 'Mark White'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/10 9:00 AM \u2013 10:00 AM PST (1 hour)'])
add_table_row(t, ['Description', 'Deactivate 3 SN catalog forms replaced by Staples ordering, incidents, and Otto.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 Move update set from dev to prod'])
add_table_row(t2, ['Backout', 'OK \u2014 Back out update set'])
add_table_row(t2, ['Test', 'OK \u2014 Tested in dev; verify forms inaccessible post-deploy'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 Not business-critical, no PHI/PII, no outage'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Have end users been notified these forms are going away? KB article or redirect?', style='List Bullet')
doc.add_paragraph('Any active/open requests through these forms that need processing first?', style='List Bullet')
doc.add_paragraph('IT Request replaced by incidents \u2014 does Service Desk have updated procedures?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# --- 12. CHG0039213 ---
doc.add_heading('12. CHG0039213 \u2014 New Global ServiceNow Survey', level=2)
t = doc.add_table(rows=0, cols=2); t.style = 'Light Grid Accent 1'
add_table_row(t, ['Assignee', 'Beth Vanderheiden (Enterprise Applications)'])
add_table_row(t, ['Requested By', 'Mark White'])
add_table_row(t, ['Risk / Impact', 'Low / 3 - Low'])
add_table_row(t, ['Planned Window', '03/05 4:00 PM \u2013 6:00 PM PST (2 hours)'])
add_table_row(t, ['Description', 'New consolidated survey for Shared Services & Service Delivery, replacing IT/HR surveys. Otto 7-day reminder flow.'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Plan Assessment:').bold = True
t2 = doc.add_table(rows=0, cols=2); t2.style = 'Light Grid Accent 1'
add_table_row(t2, ['Implementation', 'OK \u2014 Update set: survey, triggers, email notifications, Otto reminder. Old surveys disabled.'])
add_table_row(t2, ['Backout', 'OK \u2014 Back out update set'])
add_table_row(t2, ['Test', 'OK \u2014 Tested in dev with business; prod validation plan documented'])
add_table_row(t2, ['Risk Analysis', 'OK \u2014 Not business-critical, no PHI/PII, SAR completed'])
doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('CCB Questions:').bold = True
doc.add_paragraph('Low-score alert emails \u2014 what threshold triggers them? Who receives (IT vs HR)?', style='List Bullet')
doc.add_paragraph('Otto 7-day reminder \u2014 is there a cap on reminders per user?', style='List Bullet')
doc.add_paragraph('Will historical survey data from old surveys remain accessible?', style='List Bullet')
p = doc.add_paragraph()
run = p.add_run('Recommendation: AUTHORIZE'); run.bold = True; run.font.color.rgb = RGBColor(0, 0x80, 0)

# SUMMARY
doc.add_page_break()
doc.add_heading('Summary of Recommendations', level=1)
rec_table = doc.add_table(rows=1, cols=3)
rec_table.style = 'Light Grid Accent 1'
hdr = rec_table.rows[0].cells
for i, h in enumerate(['CHG #', 'Title', 'Recommendation']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True; r.font.size = Pt(9)

recs = [
    ('CHG0039420', 'Deploy Jump Server for EDI', 'AUTHORIZE', '00802B'),
    ('CHG0039418', 'Abnormal - Remove Domains from Safelist', 'AUTHORIZE', '00802B'),
    ('CHG0039415', 'Workato - Intune Out of Sync Devices', 'AUTHORIZE', '00802B'),
    ('CHG0039414', 'Update Okta RADIUS Agent', 'AUTHORIZE', '00802B'),
    ('CHG0039412', 'Upgrade Firmware on ORDC Hx', 'AUTHORIZE WITH DISCUSSION', 'CC8800'),
    ('CHG0039403', 'ASR Policies for PP Users (Warning)', 'AUTHORIZE', '00802B'),
    ('CHG0039399', 'Apply Security Patches to F5 LBs', 'AUTHORIZE', '00802B'),
    ('CHG0039383', 'Vituity Stats SharePoint Page', 'AUTHORIZE', '00802B'),
    ('CHG0039339', 'SN to Otto Integration - Agent Alerts', 'AUTHORIZE', '00802B'),
    ('CHG0039284', 'Enable Otto Triage to HR & Facilities', 'AUTHORIZE', '00802B'),
    ('CHG0039283', 'Deactivate Forms', 'AUTHORIZE', '00802B'),
    ('CHG0039213', 'New Global ServiceNow Survey', 'AUTHORIZE', '00802B'),
]

for chg, title, rec, color in recs:
    row = rec_table.add_row()
    row.cells[0].text = ''
    add_hyperlink(row.cells[0].paragraphs[0], chg, chg_url(chg))
    run = row.cells[1].paragraphs[0].add_run(title)
    run.font.size = Pt(9)
    row.cells[2].text = ''
    run = row.cells[2].paragraphs[0].add_run(rec)
    run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))

doc.add_paragraph('')
p = doc.add_paragraph(); p.add_run('Key Takeaways:').bold = True
doc.add_paragraph('11 of 12 changes ready for authorization', style='List Bullet')
doc.add_paragraph('1 change (CHG0039412) is highest-risk, warrants CCB discussion on backout limitations', style='List Bullet')
doc.add_paragraph('5 changes scheduled same day as CCB (3/5) \u2014 all after 1 PM PST', style='List Bullet')
doc.add_paragraph('2 changes have Moderate risk (CHG0039412, CHG0039403)', style='List Bullet')

outpath = os.path.expanduser('~/OneDrive - Vituity/Documents/Change Management/CCB/2026/CCB_Prep_2026-03-05.docx')
os.makedirs(os.path.dirname(outpath), exist_ok=True)
doc.save(outpath)
print(f"Saved: {outpath}")
