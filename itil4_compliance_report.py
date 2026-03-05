import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import subprocess

SNOW_URL = "https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number="

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '20')
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def set_cell_text(cell, text, bold=False, font_size=10, alignment=None, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

def set_cell_hyperlink(cell, text, url, font_size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    add_hyperlink(p, text, url)

def add_status_badge(cell, status):
    colors = {
        "PASS": ("2E7D32", "PASS"),
        "FLAG": ("E65100", "FLAG"),
        "FAIL": ("C62828", "FAIL"),
        "PASS w/ minor flags": ("558B2F", "PASS*"),
        "FLAG — needs CCB discussion": ("BF360C", "FLAG"),
    }
    color_hex, label = colors.get(status, ("757575", status))
    set_cell_text(cell, label, bold=True, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor.from_string(color_hex))

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# Title
title = doc.add_heading('ITIL 4 Change Enablement Compliance Review', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('CAB Date: March 5, 2026 — 14 Change Requests')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()

# ── Executive Summary ──
doc.add_heading('Executive Summary', level=1)

summary_data = [
    ("PASS", "9", "CHG0039423, CHG0039418, CHG0039415, CHG0039403, CHG0039383, CHG0039339, CHG0039284, CHG0039414, CHG0039213"),
    ("PASS w/ minor flags", "3", "CHG0039420, CHG0039399, CHG0039283"),
    ("FLAG — needs CCB discussion", "2", "CHG0039412, CHG0039386"),
]

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 1'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
set_cell_text(hdr[0], "Status", bold=True, font_size=10)
set_cell_text(hdr[1], "Count", bold=True, font_size=10)
set_cell_text(hdr[2], "Changes", bold=True, font_size=10)
set_cell_shading(hdr[0], "D6E4F0")
set_cell_shading(hdr[1], "D6E4F0")
set_cell_shading(hdr[2], "D6E4F0")

for status, count, changes in summary_data:
    row = tbl.add_row().cells
    add_status_badge(row[0], status)
    set_cell_text(row[1], count, font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # Add changes as hyperlinks
    row[2].text = ""
    p = row[2].paragraphs[0]
    chg_list = [c.strip() for c in changes.split(",")]
    for i, chg in enumerate(chg_list):
        if i > 0:
            p.add_run(", ").font.size = Pt(9)
        add_hyperlink(p, chg, SNOW_URL + chg)

doc.add_paragraph()

# ── Key Items for CCB Attention ──
doc.add_heading('Key Items for CCB Attention', level=1)

attention_items = [
    ("CHG0039412 — ORDC Hx Firmware", "Backout plan explicitly states HXDP/ESXi downgrade is \"generally not supported\" and risks data loss. CCB should confirm acceptable risk given Cisco TAC standby. ESXi target version is missing from the version field."),
    ("CHG0039386 — trust-manager Emergency", "Already implemented 2/23. This is a retroactive authorization. CCB should confirm the emergency was justified (PDRS/Athena IDX outage) and that post-implementation review was completed."),
    ("General Observation", "Three changes have weak pre-implementation test plans (CHG0039420, CHG0039399, CHG0039283) relying on \"not applicable\" or \"historical experience.\" While acceptable for their risk levels, this is a recurring pattern worth noting."),
]

for title_text, body_text in attention_items:
    p = doc.add_paragraph()
    run = p.add_run(title_text + ": ")
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(body_text)
    run.font.size = Pt(10)

doc.add_page_break()

# ── Individual Change Reviews ──
doc.add_heading('Individual Change Reviews', level=1)

changes = [
    {
        "number": "CHG0039423",
        "title": "Expose Titan SFTP To Internet",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Daniel Anderson", "group": "Enterprise Networking",
        "schedule": "3/5 2:00–3:00 PM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal — appropriate for firewall rule change"),
            ("Justification", "PASS", "CornerStone replacement, continuous improvement"),
            ("Implementation Plan", "PASS", "Discrete steps with specific device (pdx-edgefw01, 10.100.25.103)"),
            ("Backout Plan", "PASS", "Re-add IP restrictions; clear rollback trigger"),
            ("Test Plan", "PASS", "Prior testing done; post-impl systems validation"),
            ("Risk Assessment", "PASS", "All R&I questions answered"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039420",
        "title": "Deploy Jump Server for EDI",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Michael Castro", "group": "Enterprise Systems",
        "schedule": "3/5 2:00–2:30 PM PST",
        "overall": "PASS w/ minor flags",
        "criteria": [
            ("Type Classification", "PASS", "Normal — new server deployment"),
            ("Justification", "PASS", "DR testing gap, Illumio block workaround"),
            ("Implementation Plan", "PASS", "Discrete steps: deploy VM, set networking, join domain, configure RDS, grant access"),
            ("Backout Plan", "PASS", "Delete VM if not needed"),
            ("Test Plan", "FLAG", "Pre-impl testing \"Not applicable\" — should validate template boots"),
            ("Risk Assessment", "PASS", "All questions answered"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039418",
        "title": "Abnormal - Remove Domains from Safelist",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Bill Carter", "group": "Infosec",
        "schedule": "3/10–3/27 (rolling)",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Compromised sender inbox detection improvement"),
            ("Implementation Plan", "PASS", "7 clear steps including monitoring and remediation"),
            ("Backout Plan", "PASS", "Re-add domains"),
            ("Test Plan", "PASS", "Tested with 5 domains, no issues"),
            ("Risk Assessment", "PASS", "All answered"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039415",
        "title": "Workato Recipe: Out-of-Sync Device Notification",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Pallav Malu", "group": "Enterprise Applications",
        "schedule": "3/5 8:00–10:00 PM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Compliance remediation automation"),
            ("Implementation Plan", "PASS", "4 clear steps with recipe names and documentation links"),
            ("Backout Plan", "PASS", "Stop recipes in Workato"),
            ("Test Plan", "PASS", "UAT performed, data validated by requestor"),
            ("Risk Assessment", "PASS", "All answered"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039414",
        "title": "Update Okta RADIUS Agent",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Jeffrey How", "group": "Infosec",
        "schedule": "3/6 10:00–11:00 AM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal — vulnerability remediation"),
            ("Justification", "PASS", "Log4j vulnerability remediation"),
            ("Implementation Plan", "PASS", "4 clear steps, sequential server upgrade (rcm-srvokta, ordc-srvokta)"),
            ("Backout Plan", "PASS", "Reinstall v2.24.2 or restore snapshot"),
            ("Test Plan", "PASS", "Historical success; post-impl RADIUS service verification"),
            ("Risk Assessment", "PASS", "Versions documented (2.24.2 → 2.26.0)"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039412",
        "title": "Upgrade Firmware on ORDC Hx",
        "type": "Normal", "risk": "Moderate", "impact": "2 - Medium",
        "assignee": "Philip Weiss", "group": "Enterprise Systems",
        "schedule": "3/11 6:00–11:59 PM PST",
        "overall": "FLAG — needs CCB discussion",
        "criteria": [
            ("Type Classification", "PASS", "Normal — Moderate risk, Medium impact appropriate"),
            ("Justification", "PASS", "Multiple vulnerability remediation (firmware + ESXi)"),
            ("Implementation Plan", "PASS", "4 phased steps; Step 1 already completed"),
            ("Backout Plan", "FLAG", "HXDP/ESXi downgrade \"generally not supported\" — risks data loss. Partial backout only."),
            ("Test Plan", "FLAG", "Pre-impl: \"done in the past\" only. Cisco TAC on standby is good mitigation."),
            ("Risk Assessment", "PASS", "QA: \"Not Possible: single Hx for production\" — honest/acceptable"),
            ("Schedule", "PASS", "After CCB meeting"),
            ("Version Control", "FLAG", "ESXi target version incomplete (\"7.0.3 → ???\")"),
        ]
    },
    {
        "number": "CHG0039403",
        "title": "Assign ASR Policies to PP Users (Warning Mode)",
        "type": "Normal", "risk": "Moderate", "impact": "2 - Medium",
        "assignee": "Chintan Myakal", "group": "Infosec",
        "schedule": "3/6 7:00–10:00 AM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal — Moderate risk, Medium impact appropriate for security policy rollout"),
            ("Justification", "PASS", "Security posture enhancement"),
            ("Implementation Plan", "PASS", "9 detailed Intune steps with specific group names"),
            ("Backout Plan", "PASS", "Switch from Warn to Audit mode"),
            ("Test Plan", "PASS", "Tested with PP test group; exclusions validated"),
            ("Risk Assessment", "PASS", "All answered"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039399",
        "title": "Apply Security Patches to F5 Load Balancers",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Philip Weiss", "group": "Enterprise Systems",
        "schedule": "3/6 6:00–8:00 PM PST",
        "overall": "PASS w/ minor flags",
        "criteria": [
            ("Type Classification", "PASS", "Normal — vulnerability remediation"),
            ("Justification", "PASS", "Two specific CVE references"),
            ("Implementation Plan", "PASS", "Rolling patch with failover sequence (ORDC-LB01, ORDC-LB02)"),
            ("Backout Plan", "PASS", "Fail traffic to unpatched unit"),
            ("Test Plan", "FLAG", "Pre-impl: \"None\" — relies on historical experience only"),
            ("Risk Assessment", "PASS", "Versions documented (17.5.1-0.0.7 → 17.5.1.4-0.0.20)"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039386",
        "title": "Upgrade trust-manager to v0.21.0",
        "type": "Emergency", "risk": "Low", "impact": "3 - Low",
        "assignee": "Stefan Nuxoll", "group": "Development - RCM",
        "schedule": "2/23 10:00 AM–12:00 PM (ALREADY IMPLEMENTED)",
        "overall": "FLAG — needs CCB discussion",
        "criteria": [
            ("Type Classification", "FLAG", "Emergency — retroactive authorization required"),
            ("Justification", "PASS", "PDRS/Athena IDX CA certificate issue — time-sensitive"),
            ("Implementation Plan", "FLAG", "Single step (\"bump chart version\") — minimal detail"),
            ("Backout Plan", "PASS", "Revert to v0.19.0"),
            ("Test Plan", "PASS", "Verify cert bundles + PDRS functionality"),
            ("Risk Assessment", "PASS", "QA: N/A — applied to all environments immediately"),
            ("Schedule", "FLAG", "Already executed 2/23 — retroactive CCB review"),
        ]
    },
    {
        "number": "CHG0039383",
        "title": "Vituity Stats SharePoint Page",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "DivyaRani Bhat", "group": "Enterprise Applications",
        "schedule": "3/5 3:00–5:00 PM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Executive-sponsored (Mu), read-only reporting pipeline"),
            ("Implementation Plan", "PASS", "4 stored procedures + Workato recipe + SharePoint CSVs with links"),
            ("Backout Plan", "PASS", "Disable Workato recipe; clear backout triggers"),
            ("Test Plan", "PASS", "SP validation, Workato testing, data validation, post-impl monitoring"),
            ("Risk Assessment", "PASS", "Thoroughly answered; no PHI/PII exposure"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039339",
        "title": "ServiceNow to Otto Integration - Agent Alerts",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Dan Spengler", "group": "Service Delivery Optimization",
        "schedule": "3/5 3:00–4:00 PM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Improve ticket visibility for agents"),
            ("Implementation Plan", "PASS", "Very detailed: update set, 2 flows, field addition, OAuth integration, Moveworks listener"),
            ("Backout Plan", "PASS", "Disable flows + disconnect listener; update set backout documented"),
            ("Test Plan", "PASS", "Tested in dev/sandbox with multiple agents"),
            ("Risk Assessment", "PASS", "SAR completed"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039284",
        "title": "Enable Otto to Triage to HR and Facilities",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Beth Vanderheiden", "group": "Service Delivery Optimization",
        "schedule": "3/10 9:00–10:00 AM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Faster ticket routing"),
            ("Implementation Plan", "PASS", "4 clear steps in Moveworks portal"),
            ("Backout Plan", "PASS", "Re-add blocked groups"),
            ("Test Plan", "PASS", "Tested in dev; post-impl monitoring via analytics"),
            ("Risk Assessment", "PASS", "All answered"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039283",
        "title": "Deactivate 3 ServiceNow Forms",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Beth Vanderheiden", "group": "Enterprise Applications",
        "schedule": "3/10 9:00–10:00 AM PST",
        "overall": "PASS w/ minor flags",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Streamlining; forms replaced by Otto and Staples"),
            ("Implementation Plan", "PASS", "Update set promotion from dev to prod"),
            ("Backout Plan", "PASS", "Back out update set"),
            ("Test Plan", "FLAG", "Pre-impl: \"not needed\" — should verify update set previews cleanly"),
            ("Risk Assessment", "FLAG", "\"Documented prior notes\": No — minor gap"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
    {
        "number": "CHG0039213",
        "title": "New Global ServiceNow Survey",
        "type": "Normal", "risk": "Low", "impact": "3 - Low",
        "assignee": "Beth Vanderheiden", "group": "Enterprise Applications",
        "schedule": "3/5 4:00–6:00 PM PST",
        "overall": "PASS",
        "criteria": [
            ("Type Classification", "PASS", "Normal"),
            ("Justification", "PASS", "Consolidated survey for IT + HR"),
            ("Implementation Plan", "PASS", "Detailed: update set, survey config, triggers, email notifications, Flow Designer recipe"),
            ("Backout Plan", "PASS", "Back out update set"),
            ("Test Plan", "PASS", "Tested in dev with business; post-impl manual trigger verification"),
            ("Risk Assessment", "PASS", "All answered; SAR completed"),
            ("Schedule", "PASS", "After CCB meeting"),
        ]
    },
]

for chg in changes:
    # Change header
    h = doc.add_heading(level=2)
    add_hyperlink(h, chg["number"], SNOW_URL + chg["number"])
    h.add_run(f' — {chg["title"]}')

    # Meta info
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    for label, val in [("Type", chg["type"]), ("Risk", chg["risk"]), ("Impact", chg["impact"]),
                       ("Assignee", chg["assignee"]), ("Group", chg["group"]), ("Schedule", chg["schedule"])]:
        run = meta.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(9)
        run = meta.add_run(f"{val}   ")
        run.font.size = Pt(9)

    # Overall verdict
    verdict_p = doc.add_paragraph()
    run = verdict_p.add_run("Overall Verdict: ")
    run.bold = True
    run.font.size = Pt(11)
    verdict_run = verdict_p.add_run(chg["overall"])
    verdict_run.bold = True
    verdict_run.font.size = Pt(11)
    if "FLAG" in chg["overall"]:
        verdict_run.font.color.rgb = RGBColor(0xBF, 0x36, 0x0C)
    elif "minor" in chg["overall"]:
        verdict_run.font.color.rgb = RGBColor(0x55, 0x8B, 0x2F)
    else:
        verdict_run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)

    # Criteria table
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    set_cell_text(hdr[0], "Criteria", bold=True, font_size=9)
    set_cell_text(hdr[1], "Status", bold=True, font_size=9)
    set_cell_text(hdr[2], "Notes", bold=True, font_size=9)
    set_cell_shading(hdr[0], "D6E4F0")
    set_cell_shading(hdr[1], "D6E4F0")
    set_cell_shading(hdr[2], "D6E4F0")

    for crit_name, crit_status, crit_notes in chg["criteria"]:
        row = tbl.add_row().cells
        set_cell_text(row[0], crit_name, font_size=9)
        add_status_badge(row[1], crit_status)
        set_cell_text(row[2], crit_notes, font_size=9)
        if crit_status == "FLAG":
            set_cell_shading(row[0], "FFF3E0")
            set_cell_shading(row[1], "FFF3E0")
            set_cell_shading(row[2], "FFF3E0")

    doc.add_paragraph()

# Save
output_dir = r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\CCB\2026"
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, "ITIL4 Compliance Review - CAB 2026-03-05.docx")
doc.save(output_path)
print(f"Saved to: {output_path}")
