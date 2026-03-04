"""Generate CCB Prep Report for 2026-02-26 CCB meeting (Final — 02/25 PM)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def make_url(chg):
    return f"https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number={chg}"


def styled_header_row(table, headers):
    for i, t in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(t)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")


def add_row(table, chg, rest, shade=None):
    row = table.add_row()
    cell0 = row.cells[0]
    cell0.text = ""
    p = cell0.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    add_hyperlink(p, chg, make_url(chg))
    if shade:
        set_cell_shading(cell0, shade)
    for idx, val in enumerate(rest):
        cell = row.cells[idx + 1]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        if shade:
            set_cell_shading(cell, shade)


def add_section_heading(doc, chg, title, verdict, verdict_color):
    p = doc.add_paragraph()
    add_hyperlink(p, chg, make_url(chg))
    run = p.add_run(f" \u2014 {title}  |  ")
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(verdict)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = verdict_color


def add_check_table(doc, rows_data):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    styled_header_row(tbl, ["Check", "Status"])
    for label, val in rows_data:
        row = tbl.add_row()
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.bold = True
        row.cells[1].text = ""
        p = row.cells[1].paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(val)
        run.font.size = Pt(9)
    return tbl


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def bold_para(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(10)
    run = p.add_run(normal_text)
    run.font.size = Pt(10)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    GREEN = RGBColor(0x00, 0x80, 0x00)
    AMBER = RGBColor(0xCC, 0x7A, 0x00)

    # ---- Title ----
    title = doc.add_heading("CCB Prep \u2014 2026-02-26 CCB Review", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thursday, February 26, 2026 | 1:00 PM PST")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Prepared by: Dan Fallon, Change Manager | Updated: February 25, 2026 PM"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    # ---- Executive Summary ----
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "Fourteen Normal changes have a CAB date of 02/26/2026 for Thursday CCB review. "
        "Nine are ready for approval and five have advisories (all approvable with conditions). "
        "No changes are blocked from CCB. The agenda is heavily weighted toward InfoSec "
        "hardening (4 changes), access review housekeeping (2 changes), endpoint security "
        "(2 changes), and enterprise infrastructure (2 changes). Three changes are new to "
        "the agenda since the morning pull: Delinea Secret Server HA, 2026 Annual SFDC Purge, "
        "and the SNow\u2013Power Automate connector. Two previously-blocked changes (MOOV Platform "
        "and SNow\u2013Power Automate) have been dramatically updated in response to Teams messages "
        "sent earlier today \u2014 both moved from New to Assess with Approval Requested and "
        "substantially improved content."
    )

    # ---- At-a-Glance ----
    doc.add_heading("At-a-Glance", level=1)

    glance = doc.add_table(rows=1, cols=7)
    glance.style = "Table Grid"
    glance.alignment = WD_TABLE_ALIGNMENT.CENTER
    styled_header_row(
        glance,
        [
            "CHG",
            "Short Description",
            "Group",
            "Assigned To",
            "State",
            "Planned Start",
            "Verdict",
        ],
    )

    glance_data = [
        ("CHG0039376", "Block malicious IPs in ZIA", "Infosec", "Jeffrey How", "Assess", "02/27 09:00", "READY", "E2EFDA"),
        ("CHG0039375", "Geoblock Cuba & Syria in Okta", "Infosec", "Jeffrey How", "Assess", "02/26 15:00", "READY", "E2EFDA"),
        ("CHG0039373", "Abnormal - Remove Whitelist", "Infosec", "Bill Carter", "Assess", "02/26 14:00", "READY", "E2EFDA"),
        ("CHG0039372", "O365 Impersonation Protection", "Infosec", "Bill Carter", "Assess", "03/03 10:00", "READY", "E2EFDA"),
        ("CHG0039369", "Update Zscaler Client Connector", "SDO", "Andrew Sanchez", "Assess", "03/03 15:00", "READY", "E2EFDA"),
        ("CHG0039356", "Service Account disable", "Ent. Systems", "Philip Weiss", "Assess", "02/26 14:00", "READY", "E2EFDA"),
        ("CHG0039379", "MS Defender scan frequency (Mac)", "SDO", "Jason Nguyen", "Assess", "02/26 16:00", "READY", "E2EFDA"),
        ("CHG0039317", "Workato 180-day password policy", "Ent. Apps", "Bala Tadisetty", "Assess", "02/26 14:00", "READY", "E2EFDA"),
        ("CHG0039384", "Delinea Secret Server HA", "Ent. Systems", "Melvin Mah", "Assess", "03/02 09:00", "READY", "E2EFDA"),
        ("CHG0039367", "MOOV Platform Prod Deployment", "Dev - RCM", "Stefan Nuxoll", "Assess", "03/02 00:01", "ADVISORY", "FFEB9C"),
        ("CHG0039285", "SNow-Power Automate to Prod", "Ent. Apps", "Ray Blor", "Assess", "02/26 14:00", "ADVISORY", "FFEB9C"),
        ("CHG0039351", "2026 Annual Purge SFDC", "Ent. Apps", "Yanyan Meng", "Assess", "02/26 19:00", "ADVISORY", "FFEB9C"),
        ("CHG0039377", "Disable FTP accounts", "Ent. Systems", "Philip Weiss", "Assess", "02/26 14:00", "ADVISORY", "FFEB9C"),
        ("CHG0039374", "Block Remote Desktop Software", "Infosec", "Bill Carter", "Assess", "03/03 10:00", "ADVISORY", "FFEB9C"),
    ]

    for chg, desc, grp, assign, state, sched, verdict, shade in glance_data:
        add_row(glance, chg, [desc, grp, assign, state, sched, verdict], shade)

    doc.add_paragraph()

    # ---- Summary Counts ----
    sum_tbl = doc.add_table(rows=1, cols=2)
    sum_tbl.style = "Table Grid"
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    styled_header_row(sum_tbl, ["Category", "Count"])

    for label, val, shade in [
        ("Ready for Approval", "9", "E2EFDA"),
        ("Approvable with Advisories", "5", "FFEB9C"),
        ("Not CCB-Ready", "0", None),
        ("Total on Agenda", "14", None),
    ]:
        row = sum_tbl.add_row()
        for i, v in enumerate([label, val]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            if i == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            run.font.size = Pt(9)
            run.bold = True
            if shade:
                set_cell_shading(cell, shade)

    doc.add_paragraph()

    # ==== DETAILED REVIEWS — READY ====
    doc.add_heading("Detailed Change Reviews \u2014 Ready for Approval", level=1)

    # ---- CHG0039376 ----
    add_section_heading(
        doc, "CHG0039376",
        "Enable Block malicious IPs and domains in ZIA",
        "READY", GREEN,
    )
    bold_para(doc, "Jeffrey How | Infosec | ", "Planned: 02/27 09:00\u201310:00")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/27 \u2014 day after CCB"),
        ("Implementation Plan", "4 discrete steps in ZIA Admin \u2192 Firewall Control"),
        ("Backout Plan", "Reverse steps: disable rule, activate change"),
        ("Test Plan", "Cannot pre-test (vendor limitation); Zscaler recommends; check ZIA logs post-impl"),
        ("Risk Assessment", "Complete"),
    ])
    bold_para(doc, "Assessment: ",
        "Well-prepared. Honest about testing limitations. Built-in vendor rule enabled by "
        "default for new customers. Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039375 ----
    add_section_heading(
        doc, "CHG0039375", "Geoblock Cuba & Syria in Okta", "READY", GREEN)
    bold_para(doc, "Jeffrey How | Infosec | ", "Planned: 02/26 15:00\u201316:00")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 15:00 \u2014 after CCB"),
        ("Implementation Plan", "3 steps: edit Okta geo-block zone, add countries, update Confluence docs"),
        ("Backout Plan", "Reverse steps: remove countries from zone"),
        ("Test Plan", "90-day log review: zero logons from Cuba or Syria. Post-test: verify login works"),
        ("Risk Assessment", "Complete"),
    ])
    bold_para(doc, "Assessment: ",
        "Clean, low-risk security hardening. Sanctioned country blocking is well-justified. "
        "No business presence in either country. Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039373 ---- UPGRADED from ADVISORY
    add_section_heading(
        doc, "CHG0039373",
        "Abnormal \u2014 Remove Whitelist Domains",
        "READY", GREEN,
    )
    bold_para(doc, "Bill Carter | Infosec | ", "Planned: 02/26 14:00\u201314:30")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 14:00 \u2014 after CCB"),
        ("Implementation Plan", "4 numbered steps: log into Abnormal console, navigate to whitelist, remove 5 domains, monitor"),
        ("Backout Plan", "Re-add domains to whitelist"),
        ("Test Plan", "Monitor for false positives post-implementation"),
        ("Risk Assessment", "Complete"),
        ("Domains", "tmrrg.com, appriver.com, usimmigrationlaw.com, riddlelawoffices.com, kslaw.com"),
    ])
    bold_para(doc, "Update from prior review: ", "")
    bullet(doc,
        "Implementation Plan was previously a single sentence. Bill updated it with 4 "
        "numbered steps addressing prior CCB feedback. Upgraded from Advisory to Ready.")
    bold_para(doc, "Assessment: ",
        "Low risk \u2014 removing whitelist increases security scrutiny on email. "
        "Implementation plan now meets expectations. Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039372 ----
    add_section_heading(
        doc, "CHG0039372",
        "O365 Impersonation Protection \u2014 Practice Physicians",
        "READY", GREEN,
    )
    bold_para(doc, "Bill Carter | Infosec | ", "Planned: 03/03 10:00\u201310:30")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "03/03 \u2014 next week"),
        ("Implementation Plan", "Add Practice Physicians to impersonation protection policy, enable, monitor"),
        ("Backout Plan", "Remove PP user base from policy"),
        ("Test Plan", "PM and RCM already rolled out successfully \u2014 proven phased approach"),
        ("Risk Assessment", "Complete"),
    ])
    bold_para(doc, "Assessment: ",
        "Final phase of a proven rollout. PM and RCM already live without issues. "
        "Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039369 ----
    add_section_heading(
        doc, "CHG0039369",
        "Update Zscaler Client Connector (Win & Mac)",
        "READY", GREEN,
    )
    bold_para(doc,
        "Andrew Sanchez | SDO | Requested By: Jeffrey How | ",
        "Planned: 03/03 15:00\u201317:00")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "03/03 \u2014 next week"),
        ("Implementation Plan", "Enable new versions in Zscaler console; clients self-update"),
        ("Backout Plan", "Disable versions, push previous via Intune"),
        ("Test Plan", "IT dept pilot tested for 2 weeks with no issues"),
        ("Risk Assessment", "Complete, SAR completed"),
        ("Versions", "Win: 4.8.0.88 \u2192 4.8.0.115 | Mac: 4.5.2.121 \u2192 4.5.2.153"),
    ])
    bold_para(doc, "Assessment: ",
        "Best-prepared change on the agenda. 2-week IT pilot, clear versioning, solid "
        "backout via Intune. Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039356 ----
    add_section_heading(
        doc, "CHG0039356",
        "Service Account Disable \u2014 Access Review Jan 2026",
        "READY", GREEN,
    )
    bold_para(doc,
        "Philip Weiss | Enterprise Systems | Requested By: Abdul Nazeeruddin | ",
        "Planned: 02/26 14:00\u201314:15")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 14:00 \u2014 after CCB"),
        ("Implementation Plan", "Account table with revocation audit trail for all 3 accounts"),
        ("Backout Plan", "Re-enable accounts"),
        ("Test Plan", "Verify account status as disabled"),
        ("Risk Assessment", "Complete \u2014 SAR question answered Yes"),
        ("Accounts", "PMPlus_maSync, PMPlus_mawSync, SVC_RelyHealth_Athena"),
    ])
    bold_para(doc, "Updates since initial review (02/24):", "")
    bullet(doc,
        "Justification rewritten with concise business rationale \u2014 references Jan 2026 "
        "access review and ISMS governance")
    bullet(doc, "SAR question added to Risk Assessment and answered \"Yes\"")
    bold_para(doc, "Assessment: ",
        "Substantially improved from initial review. All revocation evidence is present "
        "and clear. Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039379 ---- NEW
    add_section_heading(
        doc, "CHG0039379",
        "Update MS Defender Scanning Frequency (Mac)",
        "READY", GREEN,
    )
    bold_para(doc, "Jason Nguyen | SDO | ", "Planned: 02/26 16:00\u201316:30")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 16:00 \u2014 after CCB"),
        ("Implementation Plan",
         "Modify Defender plist config via Jamf Pro to add weekly scan cadence "
         "(Sundays 7pm local). Deploy to all managed Macs."),
        ("Backout Plan",
         "Scope devices to previous profile via Jamf Pro. APNs provides real-time "
         "deployment; offline devices default to old profile."),
        ("Test Plan",
         "Deployed to IT test group \u2014 scanning verified, no complaints. "
         "Post-impl: pull Jamf and Microsoft logs."),
        ("Risk Assessment", "Complete"),
    ])
    bold_para(doc, "Assessment: ",
        "Well-prepared. Tested with IT pilot group, clear backout via Jamf profile "
        "scoping. Low-risk configuration change. Recommend approval.")
    doc.add_paragraph()

    # ---- CHG0039317 ---- NEW
    add_section_heading(
        doc, "CHG0039317",
        "Workato \u2014 180-Day Password Policy for RCM Contractors",
        "READY", GREEN,
    )
    bold_para(doc,
        "Bala Subramanyam Tadisetty | Enterprise Applications | Requested By: Hari Krishna Pothula | ",
        "Planned: 02/26 14:00\u201315:00")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 14:00 \u2014 after CCB"),
        ("Implementation Plan",
         "4 steps: log into Workato PROD, update recipe "
         "05|ActiveDirectory|FUNC with contractor/RCM condition to assign "
         "VITUITY-180_DAYS-GSG group, start recipe, log off. Documentation linked."),
        ("Backout Plan",
         "4 steps: revert recipe to previous working version in Workato PROD"),
        ("Test Plan",
         "Validated condition in test flow. Same logic previously implemented for "
         "PP & PM. Post-impl: monitor jobs for errors."),
        ("Risk Assessment", "Complete"),
    ])
    bold_para(doc, "Context: ", "")
    bullet(doc,
        "This change remediates a security incident: RCM contractor accounts were being "
        "placed in the wrong AD group (VITUITY-RCM_NO_AGE-GSG instead of "
        "VITUITY-180_DAYS-GSG), resulting in accounts with \"password not required\" policy.")
    bullet(doc,
        "Multiple contractor accounts (AppuA, ChakramahS, DarugupalY, DhontharaV, "
        "PunnapuK) were affected. Issue was escalated by Michael Castro and Khalil.")
    bullet(doc,
        "Change has been open 14 days (created 02/11) \u2014 approval for this security "
        "fix is overdue.")
    bold_para(doc, "Assessment: ",
        "Well-prepared, addresses an active security gap. Same pattern already proven for "
        "PP & PM contractors. Recommend approval \u2014 urgent security fix.")
    doc.add_paragraph()

    # ---- CHG0039384 ---- NEW
    add_section_heading(
        doc, "CHG0039384",
        "Delinea Secret Server High Availability",
        "READY", GREEN,
    )
    bold_para(doc,
        "Melvin Mah | Enterprise Systems | Requested By: Jeffrey How | ",
        "Planned: 03/02 09:00 \u2013 03/13 17:00 (11-day vendor engagement)")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "03/02\u201303/13 \u2014 multi-day Delinea Professional Services engagement"),
        ("Implementation Plan",
         "4 steps: (1) configure new web node, (2) set up VIP on F5, "
         "(3) upgrade Secret Server 11.8\u219211.9, (4) reconfigure SAML SSO to new VIP URL"),
        ("Backout Plan",
         "Old URL (priv.medamerica.com) remains available. Delinea PS advises on backout decision."),
        ("Test Plan",
         "Under guidance of Delinea PS. Post-test: verify both backend web nodes, "
         "HA failover test (shut down each node individually, confirm availability)."),
        ("Risk Assessment", "Complete. Expected outage: 2 hours."),
        ("Versions", "Secret Server 11.8.000001 \u2192 11.9.000006"),
    ])
    bold_para(doc, "Note: ", "")
    bullet(doc,
        "11-day implementation window is appropriate for a vendor-guided professional "
        "services engagement with Delinea.")
    bullet(doc,
        "Minor typo in CR: SAML SSO URL listed as \"priv.vitiuty.com\" \u2014 should be "
        "\"priv.vituity.com\". Melvin should confirm correct URL in actual configuration.")
    bold_para(doc, "Assessment: ",
        "Well-prepared. HA failover validation plan is strong. Vendor PS engagement "
        "provides safety net. Recommend approval.")
    doc.add_paragraph()

    # ==== DETAILED REVIEWS — ADVISORY ====
    doc.add_heading("Detailed Change Reviews \u2014 Advisories", level=1)

    # ---- CHG0039367 ---- UPGRADED from NOT READY
    add_section_heading(
        doc, "CHG0039367",
        "MOOV Platform Production Deployment",
        "ADVISORY", AMBER,
    )
    bold_para(doc,
        "Stefan Nuxoll | Development - RCM | Requested By: Rosselyn Gonzalez | ",
        "Planned: 03/02 00:01\u201323:00")

    doc.add_paragraph()
    bold_para(doc, "Stefan responded to feedback. ",
        "This CR has been transformed from NOT CCB-READY to one of the most detailed "
        "changes on the agenda.")

    # Blocker resolution tracker
    blocker_tbl = doc.add_table(rows=1, cols=3)
    blocker_tbl.style = "Table Grid"
    styled_header_row(blocker_tbl, ["Prior Blocker", "Prior Status", "Current Status"])

    for blocker, prior, current, shade in [
        ("State", "New", "Assess \u2014 RESOLVED", "E2EFDA"),
        ("Approval", "Not Yet Requested", "Requested \u2014 RESOLVED", "E2EFDA"),
        ("Implementation Plan", "\"Stefan will come back in and edit 2/24\"",
         "Completely rewritten \u2014 RESOLVED", "E2EFDA"),
        ("Charm EHR Endpoint", "\"confirm production URL/IP\"",
         "Resolved \u2014 using wildcard *.charmtracker.com", "E2EFDA"),
    ]:
        row = blocker_tbl.add_row()
        for i, v in enumerate([blocker, prior, current]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(v)
            run.font.size = Pt(9)
        set_cell_shading(row.cells[2], shade)

    doc.add_paragraph()
    bold_para(doc, "Implementation Plan now covers:", "")
    bullet(doc, "Azure Firewall policy updates with specific FQDN allow rules")
    bullet(doc, "Terraform deployment via CI/CD pipeline (consumer-infra)")
    bullet(doc, "AKS cluster registration with ArgoCD")
    bullet(doc, "ACR token creation for container registry")
    bullet(doc, "ArgoCD Application deployment for consumer-prod cluster")
    bullet(doc, "PostgreSQL database setup (keycloak + moov databases, 32-char passwords)")
    bullet(doc, "Azure Key Vault secret population")
    bullet(doc, "Entra ID App Registration (\"Moov Health Admin\")")
    bullet(doc, "Cloudflare tunnel with 4 route bindings (app, auth, API, pwreset)")
    bullet(doc, "Full Keycloak bootstrap (admin accounts, Entra ID federation)")
    bullet(doc,
        "Keycloak realm configuration (Vituity SSO, Microsoft/Google social login, "
        "custom attributes, scopes, clients, password policies per InfoSec)")

    doc.add_paragraph()
    bold_para(doc, "Remaining findings:", "")
    bullet(doc,
        "Backout Plan is still \"Disable access\" \u2014 a single sentence for a change of "
        "this complexity. However, since this is a net-new greenfield deployment with "
        "nothing existing to restore, a kill switch is a reasonable approach.")
    bullet(doc,
        "Test Plan is still generic \u2014 \"Confirm login and features are available\" does "
        "not match the depth of the implementation. Hypercare period (March 3\u201313) with "
        "defined team partially mitigates this.")

    bold_para(doc, "Hypercare Team: ",
        "Camilla Simon (PM), Ari Krohnfeldt (Product), Stefan Nuxoll (DevOps), "
        "Danny Orlando (Engineering). March 3\u201313, 2026.")
    bold_para(doc, "Recommendation: ",
        "Approve. The transformation from NOT READY to this level of detail is remarkable. "
        "Backout is acceptable for a greenfield deployment. CCB should note the hypercare plan.")
    doc.add_paragraph()

    # ---- CHG0039285 ---- UPGRADED from NOT READY
    add_section_heading(
        doc, "CHG0039285",
        "Upgrade SNow\u2013Power Automate Connector to Prod",
        "ADVISORY", AMBER,
    )
    bold_para(doc,
        "Ray Blor | Enterprise Applications | Requested By: Collin Hills | ",
        "Planned: 02/26 14:00\u201315:00")

    doc.add_paragraph()
    bold_para(doc, "Ray and Collin responded to Teams messages. ",
        "Five prior blockers addressed:")

    blocker_tbl2 = doc.add_table(rows=1, cols=3)
    blocker_tbl2.style = "Table Grid"
    styled_header_row(blocker_tbl2, ["Prior Blocker", "Prior Status", "Current Status"])

    for blocker, prior, current, shade in [
        ("State", "New", "Assess \u2014 RESOLVED", "E2EFDA"),
        ("Approval", "Not Yet Requested", "Requested \u2014 RESOLVED", "E2EFDA"),
        ("Justification", "\"tbd\"",
         "Rewritten \u2014 clear ITSM automation rationale", "E2EFDA"),
        ("Description", "Placeholder email",
         "Rewritten \u2014 proper scope (read-only REST API connector, config only)", "E2EFDA"),
        ("SAR", "Blank / NA", "Yes \u2014 RESOLVED", "E2EFDA"),
    ]:
        row = blocker_tbl2.add_row()
        for i, v in enumerate([blocker, prior, current]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(v)
            run.font.size = Pt(9)
        set_cell_shading(row.cells[2], shade)

    doc.add_paragraph()
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 14:00 \u2014 after CCB"),
        ("Scope", "Connector configuration and credentials only. No application logic, "
         "business rules, or data schema changes. Read-only access."),
        ("Implementation Plan",
         "Narrative format: setup in dev by Ray, reviewed by Collin, "
         "commit update set to production by Ray, monitor."),
        ("Backout Plan", "Back out of the update set"),
        ("Test Plan", "Tested and reviewed by Collin in dev. Post-impl: Ray reviews "
         "deployment in prod; Collin monitors for issues."),
        ("Risk Assessment", "Complete. SAR: Yes."),
    ])

    doc.add_paragraph()
    bold_para(doc, "Remaining finding:", "")
    bullet(doc,
        "Implementation Plan is narrative/workflow format rather than discrete numbered "
        "steps. However, this is a low-risk configuration-only change and update set "
        "promotion is a well-understood ServiceNow admin procedure.")
    bold_para(doc, "Recommendation: ",
        "Approve. Dramatic improvement from initial review (5 blockers resolved). "
        "Low-risk configuration change with clear scope boundaries.")
    doc.add_paragraph()

    # ---- CHG0039351 ---- NEW
    add_section_heading(
        doc, "CHG0039351",
        "2026 Annual Purge SFDC",
        "ADVISORY", AMBER,
    )
    bold_para(doc,
        "Yanyan Meng | Enterprise Applications | Requested By: Racquel Llavore | ",
        "Planned: 02/26 19:00\u201321:00")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 19:00 \u2014 evening window, after CCB"),
        ("Implementation Plan",
         "Delete records across 39 Salesforce object types for 436 inactive contacts. "
         "Update Provider Status comment field with compliance note. "
         "Spreadsheet with full detail attached."),
        ("Backout Plan",
         "\"No rollback since it is data deletion\" \u2014 EXPLICITLY IRREVERSIBLE"),
        ("Test Plan",
         "Tested in Sandbox2, Racquel verified and signed off. "
         "Post-impl: Racquel verifies in prod."),
        ("Risk Assessment", "PHI/PII: Yes. Business critical: Yes."),
    ])
    doc.add_paragraph()
    bold_para(doc, "Findings:", "")
    bullet(doc,
        "IRREVERSIBLE BY DESIGN \u2014 this is a compliance-required data purge with "
        "explicitly no rollback. CCB must formally acknowledge irreversibility before approval.")
    bullet(doc,
        "PHI/PII data involved \u2014 purging credentialing data (licenses, malpractice, "
        "education, work history, etc.) for 436 providers.")
    bullet(doc,
        "Compliance-driven \u2014 per Records Retention Policy (tenure of relationship "
        "plus 10 years). This purge is required, not optional.")
    bullet(doc,
        "Description reads like an email \u2014 still contains \"Hi, I would like to "
        "request assistance...\" format.")
    bullet(doc,
        "Implementation Plan is an object type list (39 types), not discrete steps. "
        "Should describe: (1) execute deletion for 436 contacts, (2) update Provider "
        "Status comments, (3) Racquel verifies.")
    bold_para(doc, "Recommendation: ",
        "Approve with explicit CCB acknowledgment of irreversibility. This is a "
        "compliance-required data purge, tested in sandbox with business owner sign-off.")
    doc.add_paragraph()

    # ---- CHG0039377 ---- ADVISORY (unchanged)
    add_section_heading(
        doc, "CHG0039377",
        "Disable FTP accounts \u2014 Access Review Jan 2026",
        "ADVISORY", AMBER,
    )
    bold_para(doc,
        "Philip Weiss | Enterprise Systems | Requested By: Abdul Nazeeruddin | ",
        "Planned: 02/26 14:00\u201314:20")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "02/26 14:00 \u2014 after CCB"),
        ("Implementation Plan",
         "Account table with audit trail \u2014 raw data format, not discrete steps"),
        ("Backout Plan", "Re-enable accounts"),
        ("Test Plan", "Verify account unable to access FTP"),
        ("Risk Assessment",
         "SAR question STILL UNANSWERED \u2014 needs Yes/No/NA response"),
    ])
    doc.add_paragraph()
    bold_para(doc, "Findings (unchanged from prior review):", "")
    bullet(doc,
        "SAR question was added back to the Risk Assessment but left blank/unanswered "
        "\u2014 needs a Yes/No/NA response")
    bullet(doc,
        "Implementation Plan has the right content (7 accounts with revocation evidence) "
        "but formatted as raw CSV data, not actionable steps")
    bullet(doc,
        "netsuite_clarity account has NO Revoked action or date \u2014 the Action, "
        "Actioned Date, and Actioned By columns are empty. Manager is Rosalinda Rafael. "
        "This account may not have been properly revoked. Needs clarification before disabling.")
    bullet(doc,
        "Teams messages sent to Abdul Nazeeruddin and Philip Weiss on 02/25 with specific "
        "action items. No CR updates observed.")
    bold_para(doc, "Recommendation: ",
        "Approve for 6 of 7 accounts. Hold netsuite_clarity until revocation by "
        "Rosalinda Rafael is confirmed.")
    doc.add_paragraph()

    # ---- CHG0039374 ---- ADVISORY (improved)
    add_section_heading(
        doc, "CHG0039374",
        "Block Remote Desktop Software (RustDesk, Chrome RDP)",
        "ADVISORY", AMBER,
    )
    bold_para(doc, "Bill Carter | Infosec | ", "Planned: 03/03 10:00\u201310:30")
    add_check_table(doc, [
        ("State / Approval", "Assess / Requested"),
        ("Schedule", "03/03 \u2014 next week"),
        ("Implementation Plan",
         "Vague \u2014 \"add each software to a deny rule set and push via Intune.\" "
         "Does not specify AppLocker rules or Intune profiles."),
        ("Backout Plan", "Turn off AppLocker policy, revert"),
        ("Test Plan",
         "Updated \u2014 \"Testing was successful on my laptop and wanting to expand "
         "the test group with this CR.\""),
        ("Risk Assessment", "Complete"),
    ])
    doc.add_paragraph()
    bold_para(doc, "Findings:", "")
    bullet(doc,
        "Test plan contradiction from prior review has been RESOLVED \u2014 Bill clarified "
        "that testing was successful on his laptop. This is an improvement.")
    bullet(doc,
        "Implementation Plan remains vague \u2014 should specify which AppLocker rules, "
        "which Intune configuration profiles, and the scope of the IT/IS pilot group")
    bullet(doc,
        "Phased rollout (IT/IS first) is a good approach and limits blast radius, but "
        "the plan does not describe the timeline for expanding to all users")
    bold_para(doc, "Recommendation: ",
        "Approve \u2014 phased rollout to IT/IS limits risk. Bill should clarify specific "
        "AppLocker rules on record during CCB.")
    doc.add_paragraph()

    # ---- Recommended CCB Agenda ----
    doc.add_heading("Recommended CCB Agenda Order", level=1)
    doc.add_paragraph(
        "The following order groups changes by readiness to streamline the meeting:")

    agenda_tbl = doc.add_table(rows=1, cols=4)
    agenda_tbl.style = "Table Grid"
    agenda_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    styled_header_row(agenda_tbl, ["#", "CHG", "Short Description", "Expected Action"])

    agenda_items = [
        ("1", "CHG0039376", "Block malicious IPs in ZIA", "Approve"),
        ("2", "CHG0039375", "Geoblock Cuba & Syria in Okta", "Approve"),
        ("3", "CHG0039373", "Abnormal - Remove Whitelist", "Approve"),
        ("4", "CHG0039372", "O365 Impersonation Protection", "Approve"),
        ("5", "CHG0039369", "Update Zscaler Client Connector", "Approve"),
        ("6", "CHG0039356", "Service Account disable", "Approve"),
        ("7", "CHG0039379", "MS Defender scan frequency (Mac)", "Approve"),
        ("8", "CHG0039317", "Workato 180-day password policy", "Approve"),
        ("9", "CHG0039384", "Delinea Secret Server HA", "Approve; note URL typo"),
        ("10", "CHG0039367", "MOOV Platform Prod Deployment", "Approve; note hypercare plan"),
        ("11", "CHG0039285", "SNow-Power Automate to Prod", "Approve; impl plan is narrative"),
        ("12", "CHG0039351", "2026 Annual Purge SFDC", "Approve; CCB acknowledges irreversibility"),
        ("13", "CHG0039377", "Disable FTP accounts", "Approve 6/7; hold netsuite_clarity"),
        ("14", "CHG0039374", "Block Remote Desktop Software", "Approve; Bill clarifies AppLocker rules"),
    ]

    for num, chg, desc, action in agenda_items:
        row = agenda_tbl.add_row()
        row.cells[0].text = ""
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(num)
        run.font.size = Pt(9)
        cell1 = row.cells[1]
        cell1.text = ""
        p = cell1.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        add_hyperlink(p, chg, make_url(chg))
        row.cells[2].text = ""
        p = row.cells[2].paragraphs[0]
        run = p.add_run(desc)
        run.font.size = Pt(9)
        row.cells[3].text = ""
        p = row.cells[3].paragraphs[0]
        run = p.add_run(action)
        run.font.size = Pt(9)
        # Color code
        if "hold" in action.lower() or "clarif" in action.lower() or "note" in action.lower() or "acknowledges" in action.lower():
            for c in row.cells:
                set_cell_shading(c, "FFEB9C")

    doc.add_paragraph()

    # ---- Policy References ----
    doc.add_heading("Policy References", level=1)
    for ref in [
        "ISMS-STA-11.01-01: Enterprise Change Management Program",
        "ISMS-PROC-11.01-05: Procedure for Managing Change Request",
        "ISMS-WI-11.01-06: Work Instruction for Submitting Change Request",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    # ---- Save ----
    output_path = (
        r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management"
        r"\CCB\2026\CCB Prep 2026-02-26 CCB Review.docx"
    )
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
