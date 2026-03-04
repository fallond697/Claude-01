"""Generate CCB Meeting Questions for 2026-02-26 CCB."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph


def make_url(chg):
    return f"https://vituity.service-now.com/nav_to.do?uri=change_request.do?sysparm_query=number={chg}"


def styled_header_row(table, headers):
    for i, t in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(t)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "2F5496")


def chg_heading(doc, chg, title, verdict, verdict_color):
    p = doc.add_paragraph()
    add_hyperlink(p, chg, make_url(chg))
    run = p.add_run(f" \u2014 {title}  |  ")
    run.font.size = Pt(11)
    run.bold = True
    run = p.add_run(verdict)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = verdict_color


def question(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(10)


def note(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    run = p.add_run(bold_text)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)
    run = p.add_run(normal_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    GREEN = RGBColor(0x00, 0x80, 0x00)
    AMBER = RGBColor(0xCC, 0x7A, 0x00)

    # ---- Title ----
    title = doc.add_heading("CCB Meeting Questions \u2014 2026-02-26", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Thursday, February 26, 2026 | 1:00 PM PST")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("14 Changes | Prepared for: Dan Fallon, Change Manager")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph()

    # ================================================================
    doc.add_heading("Ready Changes \u2014 Quick Confirmation Questions", level=1)
    doc.add_paragraph(
        "These 9 changes are ready for approval. Questions below are for verbal "
        "confirmation on the record. Expect brief answers \u2014 keep momentum.")

    # ---- CHG0039376 ----
    chg_heading(doc, "CHG0039376",
        "Block malicious IPs in ZIA", "READY", GREEN)
    note(doc, "Ask: ", "Jeffrey How")
    question(doc,
        "Jeffrey, this rule is on by default for new Zscaler customers. Have you "
        "coordinated with SDO on the support process if users report legitimate "
        "traffic being blocked?")
    question(doc,
        "The standard block page directs users to open a VERA ticket. Is the Service "
        "Desk aware this change is going live tomorrow so they can expect potential tickets?")
    doc.add_paragraph()

    # ---- CHG0039375 ----
    chg_heading(doc, "CHG0039375",
        "Geoblock Cuba & Syria in Okta", "READY", GREEN)
    note(doc, "Ask: ", "Jeffrey How")
    question(doc,
        "Jeffrey, you reviewed 90 days of logs and found zero logons from Cuba or Syria. "
        "Have you confirmed with HR and Legal that Vituity has no employees, contractors, "
        "or locum tenens who may travel to or connect from these countries?")
    question(doc,
        "Are there any VPN or proxy considerations \u2014 could a legitimate user appear to "
        "come from a blocked country due to routing?")
    doc.add_paragraph()

    # ---- CHG0039373 ----
    chg_heading(doc, "CHG0039373",
        "Abnormal \u2014 Remove Whitelist Domains", "READY", GREEN)
    note(doc, "Ask: ", "Bill Carter")
    question(doc,
        "Bill, can you confirm the 5 domains being removed from the whitelist: "
        "tmrrg.com, appriver.com, usimmigrationlaw.com, riddlelawoffices.com, kslaw.com?")
    question(doc,
        "What is the monitoring cadence for false positives? Will you check daily, or "
        "are you relying on user reports?")
    question(doc,
        "Do any of these domains correspond to active business relationships? For example, "
        "is kslaw.com (King & Spalding LLP) or any other firm currently engaged with Vituity Legal?")
    doc.add_paragraph()

    # ---- CHG0039372 ----
    chg_heading(doc, "CHG0039372",
        "O365 Impersonation Protection \u2014 Practice Physicians", "READY", GREEN)
    note(doc, "Ask: ", "Bill Carter")
    question(doc,
        "Bill, PM and RCM have been live \u2014 what was the false positive rate during those "
        "rollouts? Any quarantine issues that needed remediation?")
    question(doc,
        "Is the PP user base significantly different in email patterns than PM/RCM? Any "
        "reason to expect a different result?")
    doc.add_paragraph()

    # ---- CHG0039369 ----
    chg_heading(doc, "CHG0039369",
        "Update Zscaler Client Connector (Win & Mac)", "READY", GREEN)
    note(doc, "Ask: ", "Andrew Sanchez / Jeffrey How")
    question(doc,
        "Andrew, the 2-week IT pilot had no issues. How many devices were in the pilot "
        "group, and what percentage of the total fleet does that represent?")
    question(doc,
        "What is the expected rollout timeline once the new versions are enabled \u2014 how "
        "quickly do clients self-update, and is there a forced update deadline?")
    doc.add_paragraph()

    # ---- CHG0039356 ----
    chg_heading(doc, "CHG0039356",
        "Service Account Disable \u2014 Access Review Jan 2026", "READY", GREEN)
    note(doc, "Ask: ", "Philip Weiss")
    question(doc,
        "Philip, all 3 accounts have revocation evidence. Can you confirm that no active "
        "processes, scheduled jobs, or integrations currently depend on PMPlus_maSync, "
        "PMPlus_mawSync, or SVC_RelyHealth_Athena?")
    question(doc,
        "How will you verify the accounts are disabled \u2014 AD console check, or will "
        "you also test that authentication attempts fail?")
    doc.add_paragraph()

    # ---- CHG0039379 ----
    chg_heading(doc, "CHG0039379",
        "MS Defender Scan Frequency (Mac)", "READY", GREEN)
    note(doc, "Ask: ", "Jason Nguyen")
    question(doc,
        "Jason, the weekly scan runs Sundays at 7pm local time. Was that timing chosen "
        "to minimize user impact? What happens if a device is actively in use during the scan?")
    question(doc,
        "How many managed Macs are in scope? Is there any performance impact observed "
        "during the IT test group scans?")
    question(doc,
        "If a device is offline on Sunday, the scan runs at next boot. Could that cause "
        "performance issues if a user powers on Monday morning and a full scan starts?")
    doc.add_paragraph()

    # ---- CHG0039317 ----
    chg_heading(doc, "CHG0039317",
        "Workato \u2014 180-Day Password Policy for RCM Contractors", "READY", GREEN)
    note(doc, "Ask: ", "Bala Tadisetty / Hari Krishna Pothula")
    question(doc,
        "Bala, this fixes a security incident where RCM contractors were placed in the "
        "wrong AD group with no password expiration. How many contractor accounts are "
        "currently affected that need remediation after the recipe is updated?")
    question(doc,
        "Will the recipe fix only apply to new hires/rehires going forward, or will it "
        "also retroactively move existing contractors (AppuA, ChakramahS, etc.) to the "
        "correct group?")
    question(doc,
        "If it's forward-only, what is the plan to remediate the existing misplaced accounts?")
    doc.add_paragraph()

    # ---- CHG0039384 ----
    chg_heading(doc, "CHG0039384",
        "Delinea Secret Server High Availability", "READY", GREEN)
    note(doc, "Ask: ", "Melvin Mah")
    question(doc,
        "Melvin, the implementation window is 03/02 through 03/13 \u2014 11 days. Can you "
        "walk through the expected timeline? Which days will have the 2-hour outage?")
    question(doc,
        "The SAML SSO URL is changing from priv.medamerica.com to priv.vituity.com. Will "
        "the old URL remain as a permanent fallback, or is there a sunset date? "
        "Note: the CR has a typo \u2014 \"priv.vitiuty.com\" \u2014 please confirm the "
        "correct URL in your actual configuration.")
    question(doc,
        "Who are the Delinea Professional Services contacts, and will they be available "
        "on-call throughout the 11-day window or only during scheduled sessions?")
    question(doc,
        "How will privileged access users be notified about the URL change and any "
        "expected downtime?")
    doc.add_paragraph()

    # ================================================================
    doc.add_heading("Advisory Changes \u2014 Discussion Questions", level=1)
    doc.add_paragraph(
        "These 5 changes have findings that require discussion or acknowledgment on "
        "the record before approval. Allow more time for these.")

    # ---- CHG0039367 ----
    chg_heading(doc, "CHG0039367",
        "MOOV Platform Production Deployment", "ADVISORY", AMBER)
    note(doc, "Ask: ", "Stefan Nuxoll")
    question(doc,
        "Stefan, the implementation plan is now very detailed \u2014 acknowledged. The "
        "backout plan is still \"Disable access.\" For a deployment of this complexity "
        "(Azure Firewall, AKS, ArgoCD, Cloudflare tunnels, Keycloak, PostgreSQL), can "
        "you describe what \"disable access\" means specifically? Is it tearing down "
        "the Cloudflare tunnel, removing DNS entries, or something else?")
    question(doc,
        "The test plan says \"Confirm login and features are available.\" Given the "
        "number of components being deployed, do you have a more specific acceptance "
        "test checklist? For example: Keycloak SSO with Entra ID, Google/Microsoft "
        "social login, patient appointment flow, HIPAA messaging?")
    question(doc,
        "Keycloak is being configured with custom password policies per InfoSec "
        "guidelines. Has InfoSec reviewed and approved the specific Keycloak realm "
        "configuration?")
    question(doc,
        "The Charm EHR integration uses a wildcard (*.charmtracker.com). Is that "
        "intentionally broad, or should it be scoped to the specific production FQDN?")
    question(doc,
        "PostgreSQL credentials are 32-character random passwords stored in Azure Key "
        "Vault. Who has access to the Key Vault, and is access audited?")
    question(doc,
        "The hypercare team is Camilla Simon, Ari Krohnfeldt, Stefan Nuxoll, and Danny "
        "Orlando from March 3\u201313. What is the on-call rotation during hypercare? "
        "Is there 24/7 coverage or business hours only?")
    doc.add_paragraph()

    # ---- CHG0039285 ----
    chg_heading(doc, "CHG0039285",
        "SNow\u2013Power Automate Connector to Prod", "ADVISORY", AMBER)
    note(doc, "Ask: ", "Ray Blor / Collin Hills")
    question(doc,
        "Ray, the implementation plan describes the workflow but not the specific steps "
        "for promoting the update set. Can you briefly walk us through: "
        "(1) export from dev, (2) import to prod, (3) preview, (4) commit, (5) verify?")
    question(doc,
        "Collin, the Risk Assessment marks PHI/PII as \"Yes\" but the description says "
        "this is read-only access to ITSM records (incidents, changes, tasks). Can you "
        "clarify \u2014 does this connector access any records containing PHI/PII, or is "
        "it strictly ITSM operational data?")
    question(doc,
        "What Power Automate flows will use this connector in production? Are they "
        "governed under the Citizen Developer Program approval process?")
    question(doc,
        "What OAuth scopes or ServiceNow ACLs are configured for this connector? "
        "Is access restricted to specific tables or is it broad REST API access?")
    doc.add_paragraph()

    # ---- CHG0039351 ----
    chg_heading(doc, "CHG0039351",
        "2026 Annual Purge SFDC", "ADVISORY", AMBER)
    note(doc, "Ask: ", "Yanyan Meng / Racquel Llavore")
    question(doc,
        "This is an IRREVERSIBLE data purge of 436 provider records across 39 Salesforce "
        "object types. Before we vote, I need the board to formally acknowledge that this "
        "change has no rollback capability. Once approved and executed, the data cannot "
        "be recovered. Does the board acknowledge?")
    question(doc,
        "Racquel, can you confirm the retention policy basis? Specifically, that all 436 "
        "contacts have exceeded the \"tenure of relationship plus 10 years\" threshold?")
    question(doc,
        "Is a backup or export of the data being taken before deletion, even if the "
        "intent is permanent removal? Some compliance frameworks require proof of "
        "what was purged.")
    question(doc,
        "Yanyan, 39 object types is a large scope. The sandbox test was signed off by "
        "Racquel \u2014 was there a reconciliation step to verify the correct records "
        "were targeted and no active provider data was accidentally included?")
    question(doc,
        "The Provider Status comment will be updated to note the purge. Is there a "
        "corresponding audit log entry or compliance record being generated for the "
        "retention policy?")
    doc.add_paragraph()

    # ---- CHG0039377 ----
    chg_heading(doc, "CHG0039377",
        "Disable FTP accounts \u2014 Access Review Jan 2026", "ADVISORY", AMBER)
    note(doc, "Ask: ", "Philip Weiss / Abdul Nazeeruddin")
    question(doc,
        "Philip, the SAR question in the Risk Assessment is still blank. Can you answer "
        "it on the record now: does this change require a Security Architecture Review? "
        "Yes, No, or N/A with explanation?")
    question(doc,
        "The netsuite_clarity account has no revocation evidence \u2014 the Action, "
        "Actioned Date, and Actioned By fields are all empty. Manager is Rosalinda "
        "Rafael. Has Rosalinda confirmed this account should be revoked?")
    question(doc,
        "If netsuite_clarity revocation is not confirmed, are you comfortable proceeding "
        "with disabling the other 6 accounts and holding netsuite_clarity until "
        "Rosalinda responds?")
    question(doc,
        "For the 6 confirmed accounts \u2014 are any of them still used by active "
        "integrations, file transfers, or scheduled jobs? How did you verify they are "
        "truly unused?")
    doc.add_paragraph()

    # ---- CHG0039374 ----
    chg_heading(doc, "CHG0039374",
        "Block Remote Desktop Software (RustDesk, Chrome RDP)", "ADVISORY", AMBER)
    note(doc, "Ask: ", "Bill Carter")
    question(doc,
        "Bill, the implementation plan says you'll \"add each software to a deny rule "
        "set and push via Intune.\" Can you be more specific: what are the exact "
        "AppLocker rule names, and which Intune configuration profile will contain them?")
    question(doc,
        "You tested on your laptop successfully. How many devices are in the IT/IS "
        "pilot group for this CR, and how long will the pilot run before expanding "
        "to the broader organization?")
    question(doc,
        "Are RustDesk and Chrome Remote Desktop currently in use by any IT support "
        "staff or end users? Have you surveyed for installed instances before blocking?")
    question(doc,
        "AppLocker deny rules can sometimes conflict with other software. Did your "
        "laptop test include verification that other applications \u2014 particularly "
        "approved remote support tools like TeamViewer or the built-in Windows Remote "
        "Desktop \u2014 still function normally?")
    doc.add_paragraph()

    # ---- Policy References ----
    doc.add_heading("Policy References", level=1)
    for ref in [
        "ISMS-STA-11.01-01: Enterprise Change Management Program",
        "ISMS-PROC-11.01-05: Procedure for Managing Change Request",
        "ISMS-WI-11.01-06: Work Instruction for Submitting Change Request",
    ]:
        doc.add_paragraph(ref, style="List Bullet")

    # ---- Save ----
    output_path = (
        r"C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management"
        r"\CCB\2026\CCB Questions 2026-02-26.docx"
    )
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
