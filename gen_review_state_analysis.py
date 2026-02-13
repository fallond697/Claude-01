from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Changes in Review State \u2014 Post-Implementation Analysis', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Enterprise Change Management\nFebruary 13, 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()

# Meta table
meta = doc.add_table(rows=4, cols=2)
meta.style = 'Light Shading Accent 1'
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Prepared by', 'Dan Fallon, Enterprise Architecture'),
    ('Source', 'ServiceNow Change Management \u2014 Review State Query'),
    ('Policy Reference', 'ISMS-STA-11.01-01 Enterprise Change Management Program'),
    ('Classification', 'Internal Use Only'),
]
for i, (k, v) in enumerate(meta_data):
    meta.rows[i].cells[0].text = k
    meta.rows[i].cells[1].text = v

doc.add_paragraph()

# Executive Summary
doc.add_heading('Executive Summary', level=1)
doc.add_paragraph(
    'There are 10 change requests currently in Review state (7 Normal, 3 Standard), all '
    'with planned implementation dates of February 12, 2026. These changes require '
    'post-implementation validation confirmation before they can be advanced to Closed. '
    'Nine changes are rated Low risk and one Moderate (CHG0039221). All 10 changes show '
    'approval status as Approved. The 7 Normal changes have individual approver records; '
    'the 3 Standard changes were auto-approved per policy.'
)

doc.add_paragraph()

# Summary table
doc.add_heading('Change Summary', level=1)

doc.add_heading('Normal Changes (7)', level=2)

summary_table = doc.add_table(rows=8, cols=7)
summary_table.style = 'Medium Shading 1 Accent 1'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Change', 'Short Description', 'Risk', 'Assigned To', 'Impl Plan', 'Backout', 'Test Plan']
for i, h in enumerate(headers):
    cell = summary_table.rows[0].cells[i]
    cell.text = h

rows_data = [
    ('CHG0039321', 'BugFix: Expiring Credentials Notifications', 'Low', 'Yang Lee', 'OK', 'OK', 'OK'),
    ('CHG0039310', 'JAMF: Deploy Tanium to Mac', 'Low', 'Jason Nguyen', 'OK', 'OK', 'OK'),
    ('CHG0039294', 'Enable Salesforce CDC for Workato', 'Low', 'Derrick Chin', 'OK', 'OK', 'OK'),
    ('CHG0039290', 'PE-1: Retire email inbox notifications', 'Low', 'Paramasivan Arunachalam', 'OK', 'OK', 'OK'),
    ('CHG0039267', 'Expirables Extract Report Scheduling', 'Low', 'Divyarani Bhat', 'OK', 'OK', 'OK'),
    ('CHG0039154', 'MOOV Email Automation in HCM', 'Low', 'Bala Tadisetty', 'OK', 'OK', 'OK'),
    ('CHG0039141', 'Convert Comments Field to Rich Text', 'Low', 'Parvathi Arun', 'OK', 'OK', 'OK'),
]
for i, row_data in enumerate(rows_data):
    row = summary_table.rows[i + 1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()

# Standard Changes table
doc.add_heading('Standard Changes (3)', level=2)

std_table = doc.add_table(rows=4, cols=7)
std_table.style = 'Medium Shading 1 Accent 1'
std_table.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(headers):
    std_table.rows[0].cells[i].text = h

std_rows = [
    ('CHG0039315', 'New Configuration Items in ServiceNow', 'Low', 'Ray Blor', 'OK', 'OK', 'OK'),
    ('CHG0039286', 'Create AI Analysis fields on NewHire/Recred', 'Low', 'Niraj Ganani', 'OK', 'OK', 'OK'),
    ('CHG0039221', 'Add survey responses to Survey Sender report', 'Moderate', 'Niraj Ganani', 'OK', 'OK', 'OK'),
]
for i, row_data in enumerate(std_rows):
    row = std_table.rows[i + 1]
    for j, val in enumerate(row_data):
        row.cells[j].text = val

doc.add_paragraph()

# Findings
doc.add_heading('Findings & Observations', level=1)

# Approval Status
doc.add_heading('Approval Status', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'All 10 changes show approval status as Approved. All changes were approved and added to '
    'Release Req RLSE0012885 (02/12 EA Release window). Approval data sourced from the '
    'sysapproval_approver table, Notes tab (work_notes/comments), and sys_history_line audit trail.'
)
doc.add_paragraph()

# Full approval table - all 10 changes
approval_table = doc.add_table(rows=11, cols=5)
approval_table.style = 'Medium Shading 1 Accent 1'
approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
appr_headers = ['Change', 'Type', 'Approved By', 'Date', 'Notes']
for i, h in enumerate(appr_headers):
    approval_table.rows[0].cells[i].text = h
appr_data = [
    ('CHG0039321', 'Normal', 'Brian Ouderkirk', '02/12/2026', 'Added to RLSE0012885'),
    ('CHG0039310', 'Normal', 'Dan Fallon', '02/12/2026', 'CCB Approved per work note'),
    ('CHG0039294', 'Normal', 'Brian Ouderkirk', '02/11/2026', 'Added to RLSE0012885'),
    ('CHG0039290', 'Normal', 'Brian Ouderkirk', '02/11/2026', 'Rejected 02/10, re-approved 02/11'),
    ('CHG0039267', 'Normal', 'Brian Ouderkirk', '02/11/2026', 'Hunter Dix noted "Approved" in work notes'),
    ('CHG0039154', 'Normal', 'Dan Fallon', '02/12/2026', 'CCB Approved per work note'),
    ('CHG0039141', 'Normal', 'Brian Ouderkirk', '02/11/2026', 'Added to RLSE0012885'),
    ('CHG0039315', 'Standard', 'Ray Blor (submitter)', '02/10/2026', 'Brian Ouderkirk approved in comments'),
    ('CHG0039286', 'Standard', 'Hunter Dix (submitter)', '02/06/2026', 'Brian Ouderkirk approved in comments'),
    ('CHG0039221', 'Standard', 'Kelly Tremper (submitter)', '01/27/2026', 'Brian Ouderkirk approved in comments'),
]
for i, (chg, ctype, approver, date, notes) in enumerate(appr_data):
    row = approval_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = ctype
    row.cells[2].text = approver
    row.cells[3].text = date
    row.cells[4].text = notes
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('CCB Approver Pool: ').bold = True
p.add_run(
    'Brian Ouderkirk (5 Normal approvals, 3 Standard approvals in comments), '
    'Dan Fallon (2 Normal approvals). Derrick Chin served as alternate approver '
    '(set to "No Longer Required" once primary approval was granted).'
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('CHG0039290 Rejection History: ').bold = True
p.add_run(
    'Initially approved by Sara Jew (02/09), reset by Brian Ouderkirk for corrections. '
    'Submitted by Paramasivan (02/10), rejected by Brian Ouderkirk citing incomplete test plan. '
    'Re-submitted by Hunter Dix (02/11), approved by Brian Ouderkirk (02/11). '
    'Backout plan was updated between submissions.'
)

doc.add_paragraph()

# Post-Deploy Validation from Notes Tab
doc.add_heading('Post-Deployment Validation (from Notes Tab)', level=2)
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'Several changes have post-deployment validation comments in the Notes tab (Additional Comments). '
    'The following changes have received validation confirmation from stakeholders:'
)
doc.add_paragraph()

val_table = doc.add_table(rows=11, cols=4)
val_table.style = 'Medium Shading 1 Accent 1'
val_table.alignment = WD_TABLE_ALIGNMENT.CENTER
val_headers = ['Change', 'Validator', 'Date', 'Validation Status']
for i, h in enumerate(val_headers):
    val_table.rows[0].cells[i].text = h
val_data = [
    ('CHG0039321', 'Jill Truesdale', '02/13/2026', 'Validated \u2014 emails sending correctly'),
    ('CHG0039310', 'Jason Nguyen', '02/12/2026', '70% deployed, devices trickling in'),
    ('CHG0039294', 'Anjali Mewara', '02/12/2026', 'CDC triggers configured, awaiting test'),
    ('CHG0039290', 'Paramasivan Arunachalam', '02/12/2026', 'Deployed to prod, awaiting feedback'),
    ('CHG0039267', 'Divyarani Bhat', '02/12/2026', 'Implemented; pending month-end comparison'),
    ('CHG0039154', 'Bala Tadisetty', '02/12/2026', 'Deployed, monitoring jobs'),
    ('CHG0039141', 'Parvathi Arun', '02/12/2026', 'Implemented, asked Taylor to validate'),
    ('CHG0039315', 'Ray Blor', '02/12/2026', 'Reviewed in production'),
    ('CHG0039286', 'Racquel Llavore', '02/12/2026', 'Validated \u2014 "looks good in prod"'),
    ('CHG0039221', 'Savannah Lee', '02/13/2026', 'Validated \u2014 "Good to close"'),
]
for i, (chg, validator, date, status) in enumerate(val_data):
    row = val_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = validator
    row.cells[2].text = date
    row.cells[3].text = status
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Summary: ').bold = True
p.add_run(
    '3 changes have explicit closure-ready validation (CHG0039321, CHG0039286, CHG0039221). '
    '6 changes have deployment confirmation but are awaiting final validation. '
    'CHG0039267 cannot be fully validated until the first month-end run.'
)

doc.add_paragraph()

# Finding 2 - CHG0039290 backout
doc.add_heading('CHG0039290 \u2014 Incomplete Backout Plan', level=2)
p = doc.add_paragraph()
p.add_run('Risk: ').bold = True
p.add_run('Low | ')
p.add_run('Impact: ').bold = True
p.add_run('3 - Low')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'The PE-1 production release involves a Salesforce changeset with permission sets, '
    'record types, page layouts, and an activated flow. The backout plan only addresses '
    'removing the permission set ("Remove PE IEC Task Set"), but does not cover reverting '
    'the deployed changeset, deactivating the flow, or restoring prior page layout assignments. '
    'If a rollback were needed, removing the permission set alone would not fully revert the '
    'system to its pre-implementation state.'
)
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
p.add_run(
    'Paramasivan Arunachalam has been contacted and has updated the backout plan in the CR '
    'to include changeset rollback steps. No further action needed.'
)

doc.add_paragraph()

# Finding 3 - CHG0039310 blank field
doc.add_heading('CHG0039310 \u2014 Sites/Business Units Field Updated', level=2)
p = doc.add_paragraph()
p.add_run('Risk: ').bold = True
p.add_run('Low | ')
p.add_run('Impact: ').bold = True
p.add_run('3 - Low')
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Finding: ').bold = True
p.add_run(
    'The JAMF/Tanium Mac deployment affects all managed Mac devices across the organization, '
    'but the "What sites or business units are impacted?" field in the Risk & Impact Analysis '
    'is blank. This was previously flagged during CCB prep for the February 12 meeting.'
)
p = doc.add_paragraph()
p.add_run('Resolution: ').bold = True
p.add_run(
    'Jason Nguyen has been contacted and the Sites/BU field has been updated in the CR. '
    'No further action needed.'
)

doc.add_paragraph()

# Detailed Change Cards
doc.add_heading('Detailed Change Review', level=1)

changes = [
    {
        'number': 'CHG0039321',
        'title': 'BugFix: Expiring Credentials Notification Emails',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Yang Lee',
        'requested': 'Jill Truesdale',
        'planned': '02/12/2026 09:00 \u2013 11:00',
        'ccb_approved': 'CCB Approved 02/12/2026',
        'description': (
            'Bug fix for notification service failing with "too many parameters" error '
            'when provider count exceeds 2,100. Fix updates SQL query in the '
            'Portal-Credentialing Notifications exe on ordc-prdweb3.'
        ),
        'impl_rating': 'OK',
        'impl_notes': 'Replace exe build files on ordc-prdweb3 with latest build from repo.',
        'backout_rating': 'OK',
        'backout_notes': 'Re-enable backed up exe. Backout trigger: error when building exe.',
        'test_rating': 'OK',
        'test_notes': 'Tested in dev. Jill validates post-deploy; INC created if issues.',
        'review_action': 'Confirm with Jill Truesdale that notification emails are sending correctly.',
    },
    {
        'number': 'CHG0039310',
        'title': 'JAMF: Deploy Tanium Client to Mac Devices',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Service Delivery Optimization',
        'assigned': 'Jason Nguyen',
        'requested': 'Andrew Sanchez',
        'planned': '02/12/2026 16:00 \u2013 16:30',
        'ccb_approved': 'CCB Approved 02/12/2026',
        'description': (
            'Deploy Tanium client (v7.8.2.1111) to all managed Mac devices via Jamf Pro '
            'policy/configuration profile. Silent background install, no user interaction.'
        ),
        'impl_rating': 'OK',
        'impl_notes': 'Deploy via Jamf Pro policy/config profile scoped to all managed Macs.',
        'backout_rating': 'OK',
        'backout_notes': 'Revoke config profile by adjusting scope; near-instant via APNS.',
        'test_rating': 'OK',
        'test_notes': 'Tested on IT test group. Verified via Jamf logs and Tanium dashboard.',
        'review_action': 'Confirm Tanium client reporting from Mac endpoints. Sites/BU field has been updated.',
    },
    {
        'number': 'CHG0039294',
        'title': 'Enable Salesforce Change Data Capture (CDC) for Workato',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Derrick Chin',
        'requested': 'Anjali Mewara',
        'planned': '02/12/2026 09:00 \u2013 10:00',
        'ccb_approved': 'CCB Approved 02/11/2026',
        'description': (
            'Enable CDC on Contact, Contract, and Contract Implementation objects so '
            'Workato can receive real-time change events for data synchronization.'
        ),
        'impl_rating': 'OK',
        'impl_notes': 'Step-by-step: Setup > CDC > select objects > Save.',
        'backout_rating': 'OK',
        'backout_notes': 'Remove objects from CDC tracking.',
        'test_rating': 'OK',
        'test_notes': 'Tested in Sandbox by Rosie and Anjali. Savannah validates post-deploy.',
        'review_action': 'Confirm with Savannah that Workato is receiving CDC events.',
    },
    {
        'number': 'CHG0039290',
        'title': 'PE-1: Retire Email Inbox Notifications, Replace with System Notifications',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Paramasivan Arunachalam',
        'requested': 'Nancy Nair',
        'planned': '02/12/2026 09:00 \u2013 11:00',
        'ccb_approved': 'CCB Approved 02/11/2026',
        'description': (
            'Deploy Salesforce changeset PE-1 to replace email inbox-based PE intake '
            'with system-level notifications and a task queue for IEC creation. Includes '
            'permission sets, record types, page layouts, and a flow trigger.'
        ),
        'impl_rating': 'OK',
        'impl_notes': '6 detailed steps: deploy changeset, assign permissions, verify record types, layouts, flow, login validation.',
        'backout_rating': 'OK',
        'backout_notes': 'Updated by Paramasivan to include changeset rollback steps.',
        'test_rating': 'OK',
        'test_notes': 'Tested in SB2 by Nancy Nair and team. Sara Jew/Nancy Nair validate post-deploy.',
        'review_action': 'Confirm with Nancy Nair/Sara Jew that PE notification queue is functional.',
    },
    {
        'number': 'CHG0039267',
        'title': 'Expirables Extract Report \u2014 Production Scheduling',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Divyarani Bhat',
        'requested': 'Jennifer Dorais',
        'planned': '02/12/2026 09:00 \u2013 11:00',
        'ccb_approved': 'CCB Approved 02/11/2026',
        'description': (
            'Schedule Workato recipes to generate Expirables Extract report 5 and 3 '
            'business days before month-end. Includes SQL data prep, notifications, '
            'and daily summary.'
        ),
        'impl_rating': 'OK',
        'impl_notes': '5 steps: configure SQL, build recipes, configure notifications, test, deploy. Recipe links provided.',
        'backout_rating': 'OK',
        'backout_notes': 'Disable Workato recipes immediately; no downtime.',
        'test_rating': 'OK',
        'test_notes': 'Validated SQL logic, executed test runs with controlled recipients.',
        'review_action': 'Confirm first scheduled run executes correctly. Side-by-side comparison with existing reports at month-end.',
    },
    {
        'number': 'CHG0039154',
        'title': 'MOOV Email Automation in HCM',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Bala Subramanyam Tadisetty',
        'requested': 'Melinda Chiem',
        'planned': '02/12/2026 14:00 \u2013 16:00',
        'ccb_approved': 'CCB Approved 02/12/2026',
        'description': (
            'Automate MOOV employee email domain setup: set primary email to @moov.health, '
            'add @vituity.com alias, add to MOOV employee list. Workato recipe update with '
            'department and job ID filtering.'
        ),
        'impl_rating': 'OK',
        'impl_notes': '4 steps: log in to Workato, update recipe with MOOV domain logic, start recipe, log off.',
        'backout_rating': 'OK',
        'backout_notes': 'Revert to previous working version of recipe.',
        'test_rating': 'OK',
        'test_notes': 'Validated and tested email generation in lower env.',
        'review_action': 'Confirm Bala has monitored jobs with no errors. Verify MOOV email routing is working.',
    },
    {
        'number': 'CHG0039141',
        'title': 'Convert Comments Field to Rich Text',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Parvathi Arun',
        'requested': 'Taylor Heberling',
        'planned': '02/12/2026 09:00 \u2013 11:00',
        'ccb_approved': 'CCB Approved 02/11/2026',
        'description': (
            'Convert the Comments field on the Line of Business object from Long Text Area '
            'to Rich Text Area in Salesforce. Requires updating dependent Apex classes and '
            'a flow before and after the data type change.'
        ),
        'impl_rating': 'OK',
        'impl_notes': 'Comprehensive: changeset with 5 components, test class validation, step-by-step deploy.',
        'backout_rating': 'OK',
        'backout_notes': 'Detailed rollback: disable dependencies, revert field type, redeploy prior versions, reactivate, regression test.',
        'test_rating': 'OK',
        'test_notes': 'Impact analysis, unit tests, functional tests, deployment validation all completed in SB2.',
        'review_action': 'Confirm with Taylor Heberling that rich text formatting is working as expected.',
    },
    {
        'number': 'CHG0039315',
        'title': 'New Configuration Items in ServiceNow',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Ray Blor',
        'requested': 'Rosalinda Rafael',
        'planned': '02/12/2026 09:00 \u2013 11:00',
        'type': 'Standard',
        'ccb_approved': 'Approved 02/12/2026',
        'description': (
            'Add 6 new Configuration Items to ServiceNow CMDB: Novarad PACS, NovaRis, '
            'Mmodal, Mirth Connect, Rscriptor (AI), Vendor Neutral Archive (VNA). '
            'Deployed via update set RB-CHG0039315.'
        ),
        'impl_rating': 'OK',
        'impl_notes': 'Deploy update set from dev to production. Tested and reviewed by Vinod.',
        'backout_rating': 'OK',
        'backout_notes': 'Back out the update set. Trigger: update set errors out.',
        'test_rating': 'OK',
        'test_notes': 'Tested and reviewed by Vinod in dev. Ray reviews deployment in prod.',
        'review_action': 'Confirm CIs are visible in CMDB and available for incident/CR assignment.',
    },
    {
        'number': 'CHG0039286',
        'title': 'Create AI Analysis Fields on NewHire/Recred Object',
        'risk': 'Low',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Niraj Ganani',
        'requested': 'Racquel Llavore',
        'planned': '02/12/2026 10:00 \u2013 10:30',
        'type': 'Standard',
        'ccb_approved': 'Approved 02/11/2026',
        'description': (
            'Create 3 new AI Analysis fields (Licensure & Privileges, Education, Work History) '
            'on the NewHire/Recred object in Salesforce to enable Workato Recredentialing '
            'integration. Preboarding integration already working; these fields unblock Recred.'
        ),
        'impl_rating': 'OK',
        'impl_notes': '2 steps: create fields via Object Manager, update VF pages (NewHireRecredVF, NewHireRecredEdit). Changeset from SB2.',
        'backout_rating': 'OK',
        'backout_notes': 'Restore VF page backups, remove FLS for new fields, deploy changeset if required.',
        'test_rating': 'OK',
        'test_notes': 'Tested in Sandbox with Racquel Llavore.',
        'review_action': 'Confirm with Racquel that AI fields are visible and Workato can write Recred analysis results.',
    },
    {
        'number': 'CHG0039221',
        'title': 'Add New Survey Responses to Survey Sender Report',
        'risk': 'Moderate',
        'impact': '3 - Low',
        'group': 'Enterprise Applications',
        'assigned': 'Niraj Ganani',
        'requested': 'Kelly Tremper',
        'planned': '02/12/2026 10:00 \u2013 10:30',
        'type': 'Standard',
        'ccb_approved': 'Approved 02/12/2026',
        'description': (
            'Add SMS and Email survey responses to the existing Survey Sender report in '
            'Salesforce. Create new Survey, Survey Questions, and Notes records. No changeset '
            'needed \u2014 direct production data entry. Tracks NPS metric reviewed by the board.'
        ),
        'impl_rating': 'OK',
        'impl_notes': 'Create records directly in Prod: App Launcher > Survey Tab > New. Standard data entry, no changeset.',
        'backout_rating': 'OK',
        'backout_notes': 'Delete newly created records.',
        'test_rating': 'OK',
        'test_notes': 'Tested in Sandbox with Savannah.',
        'review_action': 'Confirm with Savannah that survey responses appear in report with correct source (SMS/Email) attribution.',
    },
]

for chg in changes:
    doc.add_heading(f"{chg['number']} \u2014 {chg['title']}", level=2)

    # Metadata card
    card = doc.add_table(rows=5, cols=4)
    card.style = 'Light List Accent 1'
    card.alignment = WD_TABLE_ALIGNMENT.CENTER

    card.rows[0].cells[0].text = 'Risk'
    card.rows[0].cells[1].text = chg['risk']
    card.rows[0].cells[2].text = 'Impact'
    card.rows[0].cells[3].text = chg['impact']

    card.rows[1].cells[0].text = 'Assignment Group'
    card.rows[1].cells[1].text = chg['group']
    card.rows[1].cells[2].text = 'Assigned To'
    card.rows[1].cells[3].text = chg['assigned']

    card.rows[2].cells[0].text = 'Requested By'
    card.rows[2].cells[1].text = chg['requested']
    card.rows[2].cells[2].text = 'Planned Window'
    card.rows[2].cells[3].text = chg['planned']

    card.rows[3].cells[0].text = 'Impl Plan'
    card.rows[3].cells[1].text = chg['impl_rating']
    card.rows[3].cells[2].text = 'Backout Plan'
    card.rows[3].cells[3].text = chg['backout_rating']

    card.rows[4].cells[0].text = 'Test Plan'
    card.rows[4].cells[1].text = chg['test_rating']
    card.rows[4].cells[2].text = 'Approval'
    card.rows[4].cells[3].text = chg.get('ccb_approved', '')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Description: ').bold = True
    p.add_run(chg['description'])

    p = doc.add_paragraph()
    p.add_run('Implementation: ').bold = True
    p.add_run(f"[{chg['impl_rating']}] {chg['impl_notes']}")

    p = doc.add_paragraph()
    p.add_run('Backout: ').bold = True
    p.add_run(f"[{chg['backout_rating']}] {chg['backout_notes']}")

    p = doc.add_paragraph()
    p.add_run('Testing: ').bold = True
    p.add_run(f"[{chg['test_rating']}] {chg['test_notes']}")

    p = doc.add_paragraph()
    p.add_run('Review Action: ').bold = True
    p.add_run(chg['review_action'])

    doc.add_paragraph()

# Closure readiness
doc.add_heading('Closure Readiness Assessment', level=1)

closure_table = doc.add_table(rows=11, cols=4)
closure_table.style = 'Medium Shading 1 Accent 1'
closure_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cl_headers = ['Change', 'Post-Deploy Validated?', 'Outstanding Items', 'Ready to Close?']
for i, h in enumerate(cl_headers):
    closure_table.rows[0].cells[i].text = h

cl_data = [
    ('CHG0039321', 'Validated \u2014 Jill Truesdale 02/13', 'None', 'Yes \u2014 ready to close'),
    ('CHG0039310', 'Partial \u2014 70% deployed', 'Full deployment pending', 'Yes, after full rollout'),
    ('CHG0039294', 'Deployed \u2014 Anjali awaiting test', 'CDC trigger test pending', 'Yes, pending validation'),
    ('CHG0039290', 'Deployed \u2014 awaiting Nancy/Sara', 'None', 'Yes, pending validation'),
    ('CHG0039267', 'Implemented \u2014 Divyarani', 'Month-end comparison needed', 'No \u2014 after month-end'),
    ('CHG0039154', 'Deployed \u2014 Bala monitoring', 'Job monitoring in progress', 'Yes, pending validation'),
    ('CHG0039141', 'Deployed \u2014 Taylor to validate', 'Taylor validation pending', 'Yes, pending validation'),
    ('CHG0039315', 'Reviewed \u2014 Ray Blor', 'None', 'Yes \u2014 ready to close'),
    ('CHG0039286', 'Validated \u2014 Racquel 02/12', 'None', 'Yes \u2014 ready to close'),
    ('CHG0039221', 'Validated \u2014 Savannah 02/13', 'None', 'Yes \u2014 ready to close'),
]
for i, (chg, validated, outstanding, ready) in enumerate(cl_data):
    row = closure_table.rows[i + 1]
    row.cells[0].text = chg
    row.cells[1].text = validated
    row.cells[2].text = outstanding
    row.cells[3].text = ready

doc.add_paragraph()

# Recommendations
doc.add_heading('Recommendations', level=1)

doc.add_paragraph(
    'Change owners should add post-deployment validation comments to each CR so '
    'the Change Manager can advance the record to Closed.',
    style='List Number'
)
doc.add_paragraph(
    'CHG0039310: Sites/BU impacted field has been updated by Jason Nguyen. Ready for closure pending validation.',
    style='List Number'
)
doc.add_paragraph(
    'CHG0039290: Backout plan has been updated by Paramasivan to include changeset rollback. '
    'Initially rejected by Brian Ouderkirk on 02/10, re-submitted and approved 02/11. Ready for closure.',
    style='List Number'
)
doc.add_paragraph(
    'CHG0039267: This change cannot be fully validated until the first month-end run. '
    'Consider leaving in Review until the side-by-side comparison is complete.',
    style='List Number'
)

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Enterprise Architecture | Vituity Information Technology Services')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

import shutil
temp_path = r'C:\Users\FallonD\Code\Claude-01\Changes in Review State Analysis 2026-02-13.docx'
final_path = r'C:\Users\FallonD\OneDrive - Vituity\Documents\Change Management\CCB\2026\Changes in Review State Analysis 2026-02-13.docx'
doc.save(temp_path)
try:
    shutil.copy2(temp_path, final_path)
    print(f'Saved to: {final_path}')
except PermissionError:
    print(f'Saved to temp: {temp_path}')
    print(f'OneDrive copy locked. Close the file in Word/browser, then copy manually.')
